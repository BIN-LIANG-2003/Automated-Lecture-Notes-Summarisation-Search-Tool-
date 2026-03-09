import os
import sqlite3
import uuid
import io
import sys
import json
import hashlib
import mimetypes
import re
import shutil
import subprocess
import tempfile
import requests
from html import escape as html_escape
from datetime import datetime, timedelta, timezone
from flask import Flask, request, jsonify, send_from_directory, redirect, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from sklearn.feature_extraction.text import TfidfVectorizer
from itsdangerous import URLSafeTimedSerializer, BadSignature, BadTimeSignature, SignatureExpired

# --- Google 登录库 ---
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests

# --- 文本提取库 ---
import docx
import PyPDF2
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from lxml import etree, html as lxml_html

try:
    import fitz  # PyMuPDF
    FITZ_IMPORT_ERROR = ''
except Exception as e:
    fitz = None
    FITZ_IMPORT_ERROR = str(e)
    print(f"⚠️ PyMuPDF unavailable: {FITZ_IMPORT_ERROR}")

try:
    import cv2
    import numpy as np
    CV2_IMPORT_ERROR = ''
except Exception as e:
    cv2 = None
    np = None
    CV2_IMPORT_ERROR = str(e)
    print(f"⚠️ OpenCV unavailable: {CV2_IMPORT_ERROR}")

try:
    from rapidocr_onnxruntime import RapidOCR
    RAPIDOCR_IMPORT_ERROR = ''
except Exception as e:
    RapidOCR = None
    RAPIDOCR_IMPORT_ERROR = str(e)
    print(f"⚠️ RapidOCR unavailable: {RAPIDOCR_IMPORT_ERROR}")

# --- PostgreSQL 驱动 ---
import psycopg2
from psycopg2.extras import RealDictCursor

# --- AWS S3 库 (新增) ---
import boto3
from botocore.exceptions import NoCredentialsError

try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    # Keep running even when python-dotenv is not installed.
    pass

# ================= 配置部分 =================
app = Flask(__name__, static_folder='dist', static_url_path='')
CORS(app)

UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20MB
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'docx', 'webp'}
CANVAS_DEFAULT_DOMAIN = (os.environ.get('CANVAS_DEFAULT_DOMAIN') or 'canvas.instructure.com').strip().lower()
try:
    CANVAS_MAX_LIST_PAGES = max(1, min(20, int((os.environ.get('CANVAS_MAX_LIST_PAGES') or '5').strip())))
except Exception:
    CANVAS_MAX_LIST_PAGES = 5
GOOGLE_CLIENT_ID = "1076922320508-6jdkr9v6g7rku2dipd6kr3n3thojdvn4.apps.googleusercontent.com"
AUTH_TOKEN_SECRET = (
    (os.environ.get("AUTH_TOKEN_SECRET") or "").strip()
    or (os.environ.get("FLASK_SECRET_KEY") or "").strip()
    or "studyhub-dev-secret-change-me"
)
AUTH_TOKEN_SALT = "studyhub-auth-token-v1"
try:
    AUTH_TOKEN_TTL_SECONDS = max(3600, int((os.environ.get("AUTH_TOKEN_TTL_SECONDS") or "604800").strip()))
except Exception:
    AUTH_TOKEN_TTL_SECONDS = 604800
AUTH_BYPASS_ENDPOINTS = {
    "register",
    "login",
    "google_login",
    "get_document_by_share_token",
    "get_invitation_by_token",
    "ocr_health",
}

# ================= Hugging Face AI 服务配置 =================
HF_TOKEN = (os.environ.get("HF_API_TOKEN") or "").strip()
HF_MODEL_BASE_URL = (os.environ.get("HF_MODEL_BASE_URL") or "https://router.huggingface.co/hf-inference/models").rstrip("/")
# 外部 OCR 服务地址（可选）。仅在配置后才会调用云端/公网算力中心。
EXTERNAL_OCR_SERVICE_URL = (os.getenv("EXTERNAL_OCR_SERVICE_URL") or "").strip()
OCR_MODEL_ID = os.environ.get("HF_OCR_MODEL") or "lbin2021/my-lecture-ocr"
SUMMARIZER_MODEL_ID = os.environ.get("HF_SUMMARIZER_MODEL") or "facebook/bart-large-cnn"
OCRMYPDF_BINARY = (os.getenv("OCRMYPDF_BINARY") or "ocrmypdf").strip() or "ocrmypdf"
OCRMYPDF_LANGUAGE = (os.getenv("OCRMYPDF_LANGUAGE") or "eng").strip() or "eng"
_pdf_ocr_enabled_raw = str(os.getenv("ENABLE_PDF_OCR_FALLBACK") or "1").strip().lower()
ENABLE_PDF_OCR_FALLBACK = _pdf_ocr_enabled_raw not in ("0", "false", "no", "off")
try:
    OCRMYPDF_TIMEOUT_SECONDS = max(15, int((os.getenv("OCRMYPDF_TIMEOUT_SECONDS") or "180").strip()))
except Exception:
    OCRMYPDF_TIMEOUT_SECONDS = 180
try:
    TRASH_RETENTION_DAYS = max(1, min(365, int((os.getenv("TRASH_RETENTION_DAYS") or "30").strip())))
except Exception:
    TRASH_RETENTION_DAYS = 30
DEFAULT_DOCUMENT_CATEGORY = "Uncategorized"
CATEGORY_KEYWORDS = {
    'Computer Science': (
        'computer', 'algorithm', 'network', 'database', 'data structure', 'python', 'java', 'c++',
        'operating system', 'os', 'software', 'machine learning', 'deep learning', 'programming'
    ),
    'Mathematics': (
        'math', 'algebra', 'calculus', 'geometry', 'equation', 'probability', 'statistics', 'linear algebra'
    ),
    'Physics': ('physics', 'mechanics', 'thermodynamics', 'quantum', 'electromagnetic', 'optics'),
    'Chemistry': ('chemistry', 'organic', 'inorganic', 'molecule', 'reaction', 'chemical'),
    'Biology': ('biology', 'cell', 'genetics', 'ecology', 'anatomy', 'physiology'),
    'Economics': ('economics', 'microeconomics', 'macroeconomics', 'market', 'inflation', 'gdp'),
    'Business': ('business', 'management', 'marketing', 'finance', 'accounting', 'strategy'),
    'Language': ('english', 'language', 'vocabulary', 'grammar', 'literature', 'essay'),
}
WORKSPACE_SUMMARY_LENGTH_LEVELS = {'short', 'medium', 'long'}
WORKSPACE_LINK_SHARING_MODES = {'restricted', 'workspace', 'public'}
WORKSPACE_HOME_TABS = {'home', 'files', 'ai'}
DEFAULT_WORKSPACE_SETTINGS = {
    'workspace_icon': '📚',
    'description': '',
    'default_category': DEFAULT_DOCUMENT_CATEGORY,
    'auto_categorize': True,
    'default_home_tab': 'home',
    'recent_items_limit': 10,
    'allow_uploads': True,
    'allow_note_editing': True,
    'allow_ai_tools': True,
    'allow_ocr': True,
    'summary_length': 'medium',
    'keyword_limit': 5,
    'allow_member_invites': False,
    'default_invite_expiry_days': 7,
    'default_share_expiry_days': 7,
    'link_sharing_mode': 'workspace',
    'allow_member_share_management': False,
    'max_active_share_links_per_document': 5,
    'auto_revoke_previous_share_links': False,
    'allow_export': True,
}
SUMMARY_CACHE_VERSION = 'v2'

_rapid_ocr_engine = None
_auth_token_serializer = URLSafeTimedSerializer(AUTH_TOKEN_SECRET)

MIME_BY_EXT = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'txt': 'text/plain',
    'png': 'image/png',
    'jpg': 'image/jpeg',
    'jpeg': 'image/jpeg',
    'gif': 'image/gif',
    'webp': 'image/webp',
}

EDITOR_ALLOWED_TAGS = {
    'p', 'br', 'strong', 'b', 'em', 'i', 'u', 's', 'strike', 'sub', 'sup', 'mark', 'span', 'div',
    'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li', 'blockquote', 'pre',
    'code', 'a', 'table', 'thead', 'tbody', 'tr', 'th', 'td', 'colgroup', 'col', 'img', 'hr'
}
EDITOR_ALLOWED_STYLE_PROPS = {
    'font-weight', 'font-style', 'text-decoration', 'color', 'background-color',
    'text-align', 'font-size', 'font-family', 'vertical-align', 'margin-left',
    'width', 'height', 'border', 'border-collapse'
}
BLOCK_TAGS = {
    'p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote', 'pre',
    'table', 'thead', 'tbody', 'tr', 'th', 'td', 'hr'
}

NAMED_COLORS = {
    'black': (0, 0, 0),
    'white': (255, 255, 255),
    'red': (255, 0, 0),
    'green': (0, 128, 0),
    'blue': (0, 0, 255),
    'yellow': (255, 255, 0),
    'gray': (128, 128, 128),
    'grey': (128, 128, 128),
    'orange': (255, 165, 0),
    'purple': (128, 0, 128),
    'brown': (165, 42, 42),
}

HIGHLIGHT_RGB_BY_INDEX = {
    WD_COLOR_INDEX.YELLOW: (255, 255, 0),
    WD_COLOR_INDEX.BRIGHT_GREEN: (0, 255, 0),
    WD_COLOR_INDEX.TURQUOISE: (0, 255, 255),
    WD_COLOR_INDEX.PINK: (255, 192, 203),
    WD_COLOR_INDEX.BLUE: (0, 0, 255),
    WD_COLOR_INDEX.RED: (255, 0, 0),
    WD_COLOR_INDEX.DARK_BLUE: (0, 0, 139),
    WD_COLOR_INDEX.TEAL: (0, 128, 128),
    WD_COLOR_INDEX.GREEN: (0, 128, 0),
    WD_COLOR_INDEX.VIOLET: (238, 130, 238),
    WD_COLOR_INDEX.DARK_RED: (139, 0, 0),
    WD_COLOR_INDEX.DARK_YELLOW: (128, 128, 0),
    WD_COLOR_INDEX.GRAY_50: (128, 128, 128),
    WD_COLOR_INDEX.GRAY_25: (192, 192, 192),
    WD_COLOR_INDEX.BLACK: (0, 0, 0),
    WD_COLOR_INDEX.WHITE: (255, 255, 255),
}

# ================= AWS S3 配置 (新增) =================
# 从环境变量获取密钥
S3_BUCKET = os.environ.get('S3_BUCKET_NAME')
S3_KEY = os.environ.get('AWS_ACCESS_KEY_ID')
S3_SECRET = os.environ.get('AWS_SECRET_ACCESS_KEY')
S3_REGION = os.environ.get('AWS_REGION', 'us-west-2')

# ================= 工作空间邀请配置 =================
DEFAULT_INVITE_BASE_URL = 'https://automated-lecture-notes-summarisation.onrender.com'
INVITE_BASE_URL = (os.environ.get('APP_BASE_URL') or DEFAULT_INVITE_BASE_URL).rstrip('/')
RESEND_API_KEY = (os.environ.get('RESEND_API_KEY') or '').strip()
RESEND_FROM_EMAIL = (os.environ.get('RESEND_FROM_EMAIL') or 'StudyHub <onboarding@resend.dev>').strip()
INVITE_EXPIRY_DAYS = 7

# 创建 S3 客户端
try:
    s3_client = boto3.client(
        's3',
        aws_access_key_id=S3_KEY,
        aws_secret_access_key=S3_SECRET,
        region_name=S3_REGION
    )
    print("✅ AWS S3 Client initialized.")
except Exception as e:
    print(f"⚠️ AWS S3 Client failed to initialize: {e}")
    s3_client = None

# ================= 数据库智能兼容层 (DBWrapper) =================
class DBWrapper:
    """
    这个类用于屏蔽 SQLite 和 PostgreSQL 的语法差异。
    Render 使用 PostgreSQL (%s 占位符)，本地开发使用 SQLite (? 占位符)。
    """
    def __init__(self, conn, db_type):
        self.conn = conn
        self.db_type = db_type

    def execute(self, query, params=()):
        # 1. 自动转换占位符：如果是 Postgres，把 SQL 里的 '?' 替换为 '%s'
        if self.db_type == 'postgres':
            query = query.replace('?', '%s')
        
        # 2. 执行查询
        try:
            if self.db_type == 'postgres':
                cursor = self.conn.cursor()
                cursor.execute(query, params)
                return cursor
            else:
                # SQLite
                return self.conn.execute(query, params)
        except Exception as e:
            print(f"Database Execution Error: {e}")
            raise e

    def commit(self):
        self.conn.commit()

    def close(self):
        self.conn.close()

# ================= 数据库连接函数 =================
def get_db_connection():
    # 1. 尝试从 Render 环境变量获取 PostgreSQL 地址
    database_url = os.environ.get('DATABASE_URL')
    
    if database_url:
        # === 生产环境: PostgreSQL ===
        try:
            # 修正 URL 格式 (SQLAlchemy/Psycopg2 需要 postgresql:// 开头)
            if database_url.startswith("postgres://"):
                database_url = database_url.replace("postgres://", "postgresql://", 1)
            
            conn = psycopg2.connect(database_url, cursor_factory=RealDictCursor)
            return DBWrapper(conn, 'postgres')
        except Exception as e:
            print(f"❌ PostgreSQL connection failed: {e}")
            return None
    else:
        # === 本地开发环境: SQLite ===
        conn = sqlite3.connect('database.db')
        conn.row_factory = sqlite3.Row
        return DBWrapper(conn, 'sqlite')

# ================= 初始化数据库表 =================
def table_column_exists(conn, table_name, column_name):
    safe_table = str(table_name or '').strip()
    safe_column = str(column_name or '').strip()
    if not safe_table or not safe_column:
        return False

    if conn.db_type == 'sqlite':
        cursor = conn.execute(f'PRAGMA table_info({safe_table})')
        rows = cursor.fetchall()
        return any((row['name'] if hasattr(row, 'keys') else row[1]) == safe_column for row in rows)

    cursor = conn.execute(
        '''
        SELECT column_name
        FROM information_schema.columns
        WHERE table_name = ? AND column_name = ?
        ''',
        (safe_table, safe_column)
    )
    return cursor.fetchone() is not None


def documents_column_exists(conn, column_name):
    return table_column_exists(conn, 'documents', column_name)


def ensure_documents_column(conn, column_name, column_type='TEXT'):
    safe_column = str(column_name or '').strip()
    safe_type = str(column_type or 'TEXT').strip().upper()
    if not safe_column:
        return
    if documents_column_exists(conn, safe_column):
        return
    conn.execute(f'ALTER TABLE documents ADD COLUMN {safe_column} {safe_type}')


def ensure_documents_columns(conn):
    ensure_documents_column(conn, 'content_html', 'TEXT')
    ensure_documents_column(conn, 'category', 'TEXT')
    ensure_documents_column(conn, 'workspace_id', 'TEXT')
    ensure_documents_column(conn, 'deleted_at', 'TEXT')


def ensure_workspaces_column(conn, column_name, column_type='TEXT'):
    safe_column = str(column_name or '').strip()
    safe_type = str(column_type or 'TEXT').strip().upper()
    if not safe_column:
        return
    if table_column_exists(conn, 'workspaces', safe_column):
        return
    conn.execute(f'ALTER TABLE workspaces ADD COLUMN {safe_column} {safe_type}')


def ensure_workspaces_columns(conn):
    ensure_workspaces_column(conn, 'settings_json', 'TEXT')


def init_db():
    conn = get_db_connection()
    if not conn:
        print("⚠️ Warning: Could not connect to database for initialization.")
        return

    print(f"✅ Connected to database type: {conn.db_type}")

    # 根据数据库类型选择不同的建表语法
    if conn.db_type == 'postgres':
        # Postgres 使用 SERIAL 自增
        id_type = "SERIAL PRIMARY KEY"
        timestamp_type = "TIMESTAMP DEFAULT CURRENT_TIMESTAMP"
    else:
        # SQLite 使用 INTEGER PRIMARY KEY AUTOINCREMENT
        id_type = "INTEGER PRIMARY KEY AUTOINCREMENT"
        timestamp_type = "TEXT" 

    # 创建用户表
    users_sql = f'''
        CREATE TABLE IF NOT EXISTS users (
            id {id_type},
            username TEXT UNIQUE NOT NULL,
            email TEXT UNIQUE,
            password_hash TEXT NOT NULL
        );
    '''
    
    # 创建文档表
    docs_sql = f'''
        CREATE TABLE IF NOT EXISTS documents (
            id {id_type},
            filename TEXT NOT NULL,
            title TEXT,
            uploaded_at {timestamp_type},
            file_type TEXT,
            content TEXT,
            content_html TEXT,
            tags TEXT,
            category TEXT,
            workspace_id TEXT,
            username TEXT,
            last_access_at {timestamp_type},
            deleted_at {timestamp_type}
        );
    '''

    workspaces_sql = f'''
        CREATE TABLE IF NOT EXISTS workspaces (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            plan TEXT NOT NULL DEFAULT 'Free',
            owner_username TEXT NOT NULL,
            settings_json TEXT,
            created_at {timestamp_type},
            updated_at {timestamp_type}
        );
    '''

    workspace_members_sql = f'''
        CREATE TABLE IF NOT EXISTS workspace_members (
            id {id_type},
            workspace_id TEXT NOT NULL,
            username TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'member',
            status TEXT NOT NULL DEFAULT 'active',
            created_at {timestamp_type}
        );
    '''

    workspace_invitations_sql = f'''
        CREATE TABLE IF NOT EXISTS workspace_invitations (
            id {id_type},
            workspace_id TEXT NOT NULL,
            email TEXT NOT NULL,
            token TEXT UNIQUE NOT NULL,
            status TEXT NOT NULL DEFAULT 'pending',
            expires_at {timestamp_type},
            created_at {timestamp_type},
            requested_username TEXT,
            requested_at {timestamp_type},
            reviewed_by TEXT,
            reviewed_at {timestamp_type},
            review_note TEXT
        );
    '''

    document_share_links_sql = f'''
        CREATE TABLE IF NOT EXISTS document_share_links (
            id {id_type},
            document_id INTEGER NOT NULL,
            workspace_id TEXT,
            token TEXT UNIQUE NOT NULL,
            created_by TEXT,
            status TEXT NOT NULL DEFAULT 'active',
            expires_at {timestamp_type},
            created_at {timestamp_type},
            last_access_at {timestamp_type}
        );
    '''

    document_summary_cache_sql = f'''
        CREATE TABLE IF NOT EXISTS document_summary_cache (
            id {id_type},
            document_id INTEGER NOT NULL,
            workspace_id TEXT,
            username TEXT,
            content_hash TEXT NOT NULL,
            summary_length TEXT NOT NULL,
            keyword_limit INTEGER NOT NULL DEFAULT 5,
            summary_json TEXT NOT NULL,
            summary_source TEXT,
            summary_note TEXT,
            created_at {timestamp_type},
            updated_at {timestamp_type}
        );
    '''

    workspace_members_unique_sql = '''
        CREATE UNIQUE INDEX IF NOT EXISTS idx_workspace_members_workspace_user
        ON workspace_members(workspace_id, username);
    '''

    workspace_owner_idx_sql = '''
        CREATE INDEX IF NOT EXISTS idx_workspaces_owner_username
        ON workspaces(owner_username);
    '''

    workspace_invitation_lookup_sql = '''
        CREATE INDEX IF NOT EXISTS idx_workspace_invitations_workspace_status
        ON workspace_invitations(workspace_id, status);
    '''

    document_share_links_doc_idx_sql = '''
        CREATE INDEX IF NOT EXISTS idx_document_share_links_doc_status
        ON document_share_links(document_id, status);
    '''

    document_summary_cache_lookup_idx_sql = '''
        CREATE UNIQUE INDEX IF NOT EXISTS idx_document_summary_cache_lookup
        ON document_summary_cache(document_id, content_hash, summary_length, keyword_limit);
    '''

    document_summary_cache_recent_idx_sql = '''
        CREATE INDEX IF NOT EXISTS idx_document_summary_cache_updated
        ON document_summary_cache(updated_at);
    '''

    try:
        conn.execute(users_sql)
        conn.execute(docs_sql)
        conn.execute(workspaces_sql)
        conn.execute(workspace_members_sql)
        conn.execute(workspace_invitations_sql)
        conn.execute(document_share_links_sql)
        conn.execute(document_summary_cache_sql)
        conn.execute(workspace_members_unique_sql)
        conn.execute(workspace_owner_idx_sql)
        conn.execute(workspace_invitation_lookup_sql)
        conn.execute(document_share_links_doc_idx_sql)
        conn.execute(document_summary_cache_lookup_idx_sql)
        conn.execute(document_summary_cache_recent_idx_sql)
        ensure_documents_columns(conn)
        ensure_workspaces_columns(conn)
        backfill_documents_workspace_ids(conn)
        conn.commit()
        print("✅ Database tables initialized successfully.")
    except Exception as e:
        print(f"❌ Error initializing tables: {e}")
    finally:
        conn.close()

# ================= 辅助函数 =================
def utcnow_iso():
    return datetime.utcnow().isoformat()


def row_to_dict(row):
    if row is None:
        return None
    if isinstance(row, dict):
        return dict(row)
    if hasattr(row, 'keys'):
        return {key: row[key] for key in row.keys()}
    return dict(row)


def parse_iso_datetime(value):
    raw = str(value or '').strip()
    if not raw:
        return None
    normalized = raw.replace('Z', '+00:00')
    try:
        dt = datetime.fromisoformat(normalized)
    except Exception:
        return None
    if dt.tzinfo is not None:
        dt = dt.astimezone(timezone.utc).replace(tzinfo=None)
    return dt


def invitation_is_expired(expires_at):
    dt = parse_iso_datetime(expires_at)
    if dt is None:
        return True
    return dt < datetime.utcnow()


def normalize_email(value):
    return str(value or '').strip().lower()


def create_auth_token(username):
    safe_username = str(username or '').strip()
    if not safe_username:
        return ''
    payload = {
        'username': safe_username,
        'issued_at': utcnow_iso(),
    }
    return _auth_token_serializer.dumps(payload, salt=AUTH_TOKEN_SALT)


def decode_auth_token(token):
    safe_token = str(token or '').strip()
    if not safe_token:
        return False, '', 'Missing auth token'
    try:
        payload = _auth_token_serializer.loads(
            safe_token,
            salt=AUTH_TOKEN_SALT,
            max_age=AUTH_TOKEN_TTL_SECONDS
        )
    except SignatureExpired:
        return False, '', 'Auth token expired, please sign in again'
    except (BadSignature, BadTimeSignature):
        return False, '', 'Invalid auth token'
    except Exception:
        return False, '', 'Invalid auth token'

    if not isinstance(payload, dict):
        return False, '', 'Invalid auth token payload'
    username = str(payload.get('username') or '').strip()
    if not username:
        return False, '', 'Invalid auth token payload'
    return True, username, ''


def get_bearer_token():
    auth_header = str(request.headers.get('Authorization') or '').strip()
    if not auth_header:
        return ''
    if not auth_header.lower().startswith('bearer '):
        return ''
    return auth_header[7:].strip()


def extract_request_username():
    query_username = (request.args.get('username') or '').strip()
    if query_username:
        return query_username

    form_username = (request.form.get('username') or '').strip()
    if form_username:
        return form_username

    if request.is_json:
        data = request.get_json(silent=True) or {}
        if isinstance(data, dict):
            json_username = (data.get('username') or '').strip()
            if json_username:
                return json_username

    value_username = (request.values.get('username') or '').strip()
    if value_username:
        return value_username
    return ''


@app.before_request
def enforce_auth_token_middleware():
    path = str(request.path or '')
    if not path.startswith('/api/'):
        return None
    if request.method == 'OPTIONS':
        return None

    endpoint = str(request.endpoint or '')
    if endpoint in AUTH_BYPASS_ENDPOINTS:
        return None

    username = extract_request_username()
    if not username:
        return None

    bearer_token = get_bearer_token()
    if not bearer_token:
        return jsonify({'error': 'Auth token is required'}), 401

    token_ok, token_username, token_error = decode_auth_token(bearer_token)
    if not token_ok:
        return jsonify({'error': token_error or 'Invalid auth token'}), 401
    if token_username != username:
        return jsonify({'error': 'Auth token does not match username'}), 403
    return None


def normalize_workspace_name(name, owner_username=''):
    raw = str(name or '').strip()
    owner = str(owner_username or '').strip()
    if not raw:
        return f"{owner}'s Workspace" if owner else 'Untitled Workspace'
    if raw.endswith(' 的工作空间'):
        base = raw[:-len(' 的工作空间')].strip()
        if base:
            return f"{base}'s Workspace"
        return f"{owner}'s Workspace" if owner else 'Workspace'
    if raw == '未命名空间':
        return 'Untitled Workspace'
    return raw


def parse_bool(value, default=False):
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    if isinstance(value, str):
        lowered = value.strip().lower()
        if lowered in ('1', 'true', 'yes', 'on'):
            return True
        if lowered in ('0', 'false', 'no', 'off'):
            return False
    return bool(default)


def parse_int(value, default_value, min_value=None, max_value=None):
    try:
        parsed = int(value)
    except Exception:
        parsed = int(default_value)
    if min_value is not None:
        parsed = max(min_value, parsed)
    if max_value is not None:
        parsed = min(max_value, parsed)
    return parsed


def parse_float(value, default_value, min_value=None, max_value=None):
    try:
        parsed = float(value)
    except Exception:
        parsed = float(default_value)
    if min_value is not None:
        parsed = max(float(min_value), parsed)
    if max_value is not None:
        parsed = min(float(max_value), parsed)
    return parsed


def is_document_soft_deleted(doc_row):
    doc = row_to_dict(doc_row) or {}
    return bool(str(doc.get('deleted_at') or '').strip())


def remove_document_file_from_storage(filename):
    safe_filename = str(filename or '').strip()
    if not safe_filename:
        return ''
    try:
        if S3_BUCKET and s3_client:
            s3_client.delete_object(Bucket=S3_BUCKET, Key=safe_filename)
        else:
            local_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
            if os.path.exists(local_path):
                os.remove(local_path)
        return ''
    except Exception as e:
        warning = f'File cleanup failed: {e}'
        print(f"⚠️ {warning}")
        return warning


def hard_delete_document_record(conn, doc_id):
    safe_doc_id = parse_int(doc_id, 0, 0)
    if safe_doc_id <= 0:
        return None
    cursor = conn.execute(
        'SELECT id, filename, username, workspace_id, deleted_at FROM documents WHERE id = ?',
        (safe_doc_id,)
    )
    doc_row = cursor.fetchone()
    if not doc_row:
        return None
    doc = row_to_dict(doc_row) or {}
    conn.execute('DELETE FROM document_share_links WHERE document_id = ?', (safe_doc_id,))
    conn.execute('DELETE FROM document_summary_cache WHERE document_id = ?', (safe_doc_id,))
    conn.execute('DELETE FROM documents WHERE id = ?', (safe_doc_id,))
    return doc


def purge_expired_trashed_documents(conn, username='', workspace_id=''):
    safe_username = str(username or '').strip()
    safe_workspace_id = str(workspace_id or '').strip()
    cutoff = (datetime.utcnow() - timedelta(days=TRASH_RETENTION_DAYS)).isoformat()
    where_parts = [
        "COALESCE(deleted_at, '') <> ''",
        "COALESCE(deleted_at, '') <= ?",
    ]
    params = [cutoff]
    if safe_username:
        where_parts.append('username = ?')
        params.append(safe_username)
    if safe_workspace_id:
        where_parts.append('workspace_id = ?')
        params.append(safe_workspace_id)

    where_sql = ' AND '.join(where_parts)
    cursor = conn.execute(
        f'''
        SELECT id, filename
        FROM documents
        WHERE {where_sql}
        ORDER BY deleted_at ASC, id ASC
        ''',
        tuple(params)
    )
    stale_rows = [row_to_dict(item) for item in cursor.fetchall()]
    if not stale_rows:
        return {'purged_count': 0, 'warnings': []}

    purged_files = []
    for row in stale_rows:
        doc_id = parse_int((row or {}).get('id'), 0, 0)
        if doc_id <= 0:
            continue
        deleted = hard_delete_document_record(conn, doc_id)
        if deleted:
            purged_files.append(str((deleted or {}).get('filename') or '').strip())
    conn.commit()

    warnings = []
    for filename in purged_files:
        warning = remove_document_file_from_storage(filename)
        if warning:
            warnings.append(f'{filename}: {warning}')
    return {'purged_count': len(purged_files), 'warnings': warnings}


def build_summary_cache_text_hash(text):
    normalized = re.sub(r'\s+', ' ', str(text or '').strip())
    if not normalized:
        return ''
    payload = f"{SUMMARY_CACHE_VERSION}:{normalized}"
    return hashlib.sha256(payload.encode('utf-8')).hexdigest()


def load_document_summary_cache(conn, document_id, content_hash, summary_length, keyword_limit):
    safe_doc_id = parse_int(document_id, 0, 0)
    safe_hash = str(content_hash or '').strip()
    safe_summary_length = str(summary_length or '').strip().lower()
    safe_keyword_limit = parse_int(keyword_limit, 5, 1)
    if safe_doc_id <= 0 or not safe_hash or not safe_summary_length:
        return None
    try:
        cursor = conn.execute(
            '''
            SELECT summary_json, summary_source, summary_note, created_at, updated_at
            FROM document_summary_cache
            WHERE document_id = ?
              AND content_hash = ?
              AND summary_length = ?
              AND keyword_limit = ?
            ORDER BY updated_at DESC, id DESC
            LIMIT 1
            ''',
            (safe_doc_id, safe_hash, safe_summary_length, safe_keyword_limit)
        )
        row = row_to_dict(cursor.fetchone())
        if not row:
            return None
        raw_json = row.get('summary_json')
        try:
            payload = json.loads(raw_json) if isinstance(raw_json, str) else {}
        except Exception:
            payload = {}
        if not isinstance(payload, dict):
            payload = {}
        payload['summary_source'] = str(
            payload.get('summary_source') or row.get('summary_source') or 'cache'
        ).strip().lower() or 'cache'
        payload['summary_note'] = str(payload.get('summary_note') or row.get('summary_note') or '').strip()
        payload['cached_at'] = row.get('updated_at') or row.get('created_at') or utcnow_iso()
        return payload
    except Exception as e:
        print(f"⚠️ Summary cache read failed: {e}")
        return None


def save_document_summary_cache(
    conn,
    document_id,
    workspace_id,
    username,
    content_hash,
    summary_length,
    keyword_limit,
    payload
):
    safe_doc_id = parse_int(document_id, 0, 0)
    safe_hash = str(content_hash or '').strip()
    safe_summary_length = str(summary_length or '').strip().lower()
    safe_keyword_limit = parse_int(keyword_limit, 5, 1)
    if safe_doc_id <= 0 or not safe_hash or not safe_summary_length:
        return False
    if not isinstance(payload, dict) or not str(payload.get('summary') or '').strip():
        return False

    safe_payload = {
        'summary': str(payload.get('summary') or '').strip(),
        'keywords': payload.get('keywords') if isinstance(payload.get('keywords'), list) else [],
        'key_sentences': payload.get('key_sentences') if isinstance(payload.get('key_sentences'), list) else [],
        'summary_source': str(payload.get('summary_source') or '').strip().lower() or 'fallback',
        'summary_note': str(payload.get('summary_note') or '').strip(),
    }
    summary_json = json.dumps(safe_payload, ensure_ascii=False)
    now_iso = utcnow_iso()
    safe_workspace_id = str(workspace_id or '').strip()
    safe_username = str(username or '').strip()
    try:
        conn.execute(
            '''
            DELETE FROM document_summary_cache
            WHERE document_id = ?
              AND content_hash = ?
              AND summary_length = ?
              AND keyword_limit = ?
            ''',
            (safe_doc_id, safe_hash, safe_summary_length, safe_keyword_limit)
        )
        conn.execute(
            '''
            INSERT INTO document_summary_cache (
                document_id,
                workspace_id,
                username,
                content_hash,
                summary_length,
                keyword_limit,
                summary_json,
                summary_source,
                summary_note,
                created_at,
                updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''',
            (
                safe_doc_id,
                safe_workspace_id,
                safe_username,
                safe_hash,
                safe_summary_length,
                safe_keyword_limit,
                summary_json,
                safe_payload.get('summary_source') or '',
                safe_payload.get('summary_note') or '',
                now_iso,
                now_iso,
            )
        )
        conn.commit()
        return True
    except Exception as e:
        print(f"⚠️ Summary cache write failed: {e}")
        return False


def normalize_workspace_settings(raw_settings):
    if isinstance(raw_settings, str):
        try:
            source = json.loads(raw_settings)
        except Exception:
            source = {}
    elif isinstance(raw_settings, dict):
        source = raw_settings
    else:
        source = {}

    base = dict(DEFAULT_WORKSPACE_SETTINGS)
    workspace_icon = str(source.get('workspace_icon', base['workspace_icon']) or '').strip()
    base['workspace_icon'] = workspace_icon[:2] or DEFAULT_WORKSPACE_SETTINGS['workspace_icon']

    description = str(source.get('description', base['description']) or '').strip()
    base['description'] = re.sub(r'\s+', ' ', description)[:220]

    default_category = normalize_document_category(source.get('default_category', base['default_category']))
    base['default_category'] = default_category or DEFAULT_DOCUMENT_CATEGORY

    base['auto_categorize'] = parse_bool(source.get('auto_categorize', base['auto_categorize']), True)

    default_home_tab = str(source.get('default_home_tab', base['default_home_tab']) or '').strip().lower()
    if default_home_tab not in WORKSPACE_HOME_TABS:
        default_home_tab = DEFAULT_WORKSPACE_SETTINGS['default_home_tab']
    base['default_home_tab'] = default_home_tab

    base['recent_items_limit'] = parse_int(
        source.get('recent_items_limit', base['recent_items_limit']),
        10,
        5,
        20
    )
    base['allow_uploads'] = parse_bool(source.get('allow_uploads', base['allow_uploads']), True)
    base['allow_note_editing'] = parse_bool(
        source.get('allow_note_editing', base['allow_note_editing']),
        True
    )
    base['allow_ai_tools'] = parse_bool(source.get('allow_ai_tools', base['allow_ai_tools']), True)
    base['allow_ocr'] = parse_bool(source.get('allow_ocr', base['allow_ocr']), True)

    summary_length = str(source.get('summary_length', base['summary_length']) or '').strip().lower()
    if summary_length not in WORKSPACE_SUMMARY_LENGTH_LEVELS:
        summary_length = DEFAULT_WORKSPACE_SETTINGS['summary_length']
    base['summary_length'] = summary_length

    base['keyword_limit'] = parse_int(source.get('keyword_limit', base['keyword_limit']), 5, 3, 12)
    base['allow_member_invites'] = parse_bool(
        source.get('allow_member_invites', base['allow_member_invites']),
        False
    )
    base['default_invite_expiry_days'] = parse_int(
        source.get('default_invite_expiry_days', base['default_invite_expiry_days']),
        7,
        1,
        30
    )
    base['default_share_expiry_days'] = parse_int(
        source.get('default_share_expiry_days', base['default_share_expiry_days']),
        7,
        1,
        30
    )
    base['max_active_share_links_per_document'] = parse_int(
        source.get('max_active_share_links_per_document', base['max_active_share_links_per_document']),
        5,
        1,
        20
    )
    base['auto_revoke_previous_share_links'] = parse_bool(
        source.get('auto_revoke_previous_share_links', base['auto_revoke_previous_share_links']),
        False
    )

    link_sharing_mode = str(source.get('link_sharing_mode', base['link_sharing_mode']) or '').strip().lower()
    if link_sharing_mode not in WORKSPACE_LINK_SHARING_MODES:
        link_sharing_mode = DEFAULT_WORKSPACE_SETTINGS['link_sharing_mode']
    base['link_sharing_mode'] = link_sharing_mode
    base['allow_member_share_management'] = parse_bool(
        source.get('allow_member_share_management', base['allow_member_share_management']),
        False
    )

    base['allow_export'] = parse_bool(source.get('allow_export', base['allow_export']), True)
    return base


def workspace_settings_to_json(value):
    normalized = normalize_workspace_settings(value)
    return json.dumps(normalized, ensure_ascii=False)


def is_valid_email(value):
    email = normalize_email(value)
    return bool(re.fullmatch(r'[^@\s]+@[^@\s]+\.[^@\s]+', email))


def expires_at_for_days(days):
    safe_days = parse_int(days, INVITE_EXPIRY_DAYS, 1, 30)
    return (datetime.utcnow() + timedelta(days=safe_days)).isoformat()


def create_invite_token():
    return f'{uuid.uuid4().hex}{uuid.uuid4().hex}'


def build_invite_url(token):
    safe_token = str(token or '').strip()
    if not safe_token:
        return ''
    return f'{INVITE_BASE_URL}/#/invite/{safe_token}'


def create_document_share_token():
    return f'{uuid.uuid4().hex}{uuid.uuid4().hex}'


def build_document_share_url(token):
    safe_token = str(token or '').strip()
    if not safe_token:
        return ''
    return f'{INVITE_BASE_URL}/#/shared/{safe_token}'


def expire_document_share_links(conn, document_id=0):
    now_iso = utcnow_iso()
    safe_doc_id = parse_int(document_id, 0, 0)
    if safe_doc_id > 0:
        conn.execute(
            '''
            UPDATE document_share_links
            SET status = 'expired'
            WHERE document_id = ?
              AND status = 'active'
              AND expires_at < ?
            ''',
            (safe_doc_id, now_iso)
        )
    else:
        conn.execute(
            '''
            UPDATE document_share_links
            SET status = 'expired'
            WHERE status = 'active'
              AND expires_at < ?
            ''',
            (now_iso,)
        )


def serialize_document_share_link_row(row):
    data = row_to_dict(row) or {}
    return {
        'id': data.get('id'),
        'document_id': data.get('document_id'),
        'workspace_id': data.get('workspace_id', ''),
        'token': data.get('token', ''),
        'status': data.get('status', ''),
        'expires_at': data.get('expires_at', ''),
        'created_at': data.get('created_at', ''),
        'created_by': data.get('created_by', ''),
        'last_access_at': data.get('last_access_at', ''),
        'share_url': build_document_share_url(data.get('token', '')),
    }


def to_document_share_link_payload(row):
    payload = serialize_document_share_link_row(row)
    status = str(payload.get('status') or '').strip().lower()
    expired = status == 'expired' or invitation_is_expired(payload.get('expires_at'))
    payload['is_expired'] = bool(expired)
    payload['is_accessible'] = status == 'active' and not expired
    return payload


def list_document_share_link_payloads(conn, document_id, limit=20):
    safe_doc_id = parse_int(document_id, 0, 0)
    safe_limit = parse_int(limit, 20, 1, 100)
    if safe_doc_id <= 0:
        return []

    expire_document_share_links(conn, safe_doc_id)
    conn.commit()
    cursor = conn.execute(
        '''
        SELECT *
        FROM document_share_links
        WHERE document_id = ?
        ORDER BY created_at DESC, id DESC
        LIMIT ?
        ''',
        (safe_doc_id, safe_limit)
    )
    return [to_document_share_link_payload(item) for item in cursor.fetchall()]


def validate_document_share_token(conn, document_id, token, mark_access=False):
    safe_token = str(token or '').strip()
    doc_id = parse_int(document_id, 0, 0)
    if not safe_token or doc_id <= 0:
        return False, None, 'Invalid share token'

    expire_document_share_links(conn, doc_id)
    conn.commit()

    cursor = conn.execute(
        '''
        SELECT *
        FROM document_share_links
        WHERE token = ? AND document_id = ?
        ORDER BY id DESC
        LIMIT 1
        ''',
        (safe_token, doc_id)
    )
    share_row = row_to_dict(cursor.fetchone())
    if not share_row:
        return False, None, 'Invalid share link'

    status = str(share_row.get('status') or '').strip().lower()
    if status == 'expired':
        return False, share_row, 'Share link has expired'
    if status and status != 'active':
        return False, share_row, 'Share link is no longer active'

    expires_at = share_row.get('expires_at')
    if invitation_is_expired(expires_at):
        conn.execute(
            "UPDATE document_share_links SET status = 'expired' WHERE id = ?",
            (share_row.get('id'),)
        )
        conn.commit()
        share_row['status'] = 'expired'
        return False, share_row, 'Share link has expired'

    if mark_access:
        now_iso = utcnow_iso()
        conn.execute(
            'UPDATE document_share_links SET last_access_at = ? WHERE id = ?',
            (now_iso, share_row.get('id'))
        )
        conn.commit()
        share_row['last_access_at'] = now_iso

    return True, share_row, ''


def send_workspace_invite_email(to_email, workspace_name, inviter_username, invite_url, expires_at):
    recipient = normalize_email(to_email)
    if not recipient:
        return False, 'Missing recipient email'
    if not RESEND_API_KEY:
        return False, 'RESEND_API_KEY is not configured'

    expiry_label = expires_at or ''
    subject = f'Workspace invitation: {workspace_name}'
    html = f'''
        <div style="font-family: Arial, sans-serif; line-height: 1.5; color: #1f2937;">
          <h2 style="margin-bottom: 12px;">You received a StudyHub workspace invitation</h2>
          <p><strong>{inviter_username}</strong> invited you to join <strong>{workspace_name}</strong>.</p>
          <p>This invitation becomes active only after the workspace owner approves your request.</p>
          <p style="margin: 18px 0;">
            <a href="{invite_url}" style="display: inline-block; padding: 10px 14px; border-radius: 8px; text-decoration: none; background: #2563eb; color: #ffffff;">
              View invitation and request access
            </a>
          </p>
          <p>Invitation expires at: {expiry_label}</p>
        </div>
    '''
    payload = {
        'from': RESEND_FROM_EMAIL,
        'to': [recipient],
        'subject': subject,
        'html': html,
    }
    try:
        response = requests.post(
            'https://api.resend.com/emails',
            headers={
                'Authorization': f'Bearer {RESEND_API_KEY}',
                'Content-Type': 'application/json',
            },
            json=payload,
            timeout=15,
        )
        if response.status_code >= 400:
            return False, f'Resend failed ({response.status_code}): {response.text[:220]}'
        return True, ''
    except Exception as e:
        return False, f'Resend request error: {e}'


def expire_workspace_invitations(conn, workspace_id=''):
    now_iso = utcnow_iso()
    if workspace_id:
        conn.execute(
            '''
            UPDATE workspace_invitations
            SET status = 'expired'
            WHERE workspace_id = ?
              AND status IN ('pending', 'requested')
              AND expires_at < ?
            ''',
            (workspace_id, now_iso)
        )
    else:
        conn.execute(
            '''
            UPDATE workspace_invitations
            SET status = 'expired'
            WHERE status IN ('pending', 'requested')
              AND expires_at < ?
            ''',
            (now_iso,)
        )


def serialize_invitation_row(row):
    data = row_to_dict(row) or {}
    return {
        'id': data.get('id'),
        'workspace_id': data.get('workspace_id', ''),
        'email': data.get('email', ''),
        'token': data.get('token', ''),
        'status': data.get('status', ''),
        'expires_at': data.get('expires_at', ''),
        'created_at': data.get('created_at', ''),
        'requested_username': data.get('requested_username', ''),
        'requested_at': data.get('requested_at', ''),
        'reviewed_by': data.get('reviewed_by', ''),
        'reviewed_at': data.get('reviewed_at', ''),
        'review_note': data.get('review_note', ''),
        'invite_url': build_invite_url(data.get('token', '')),
    }


def ensure_owner_membership(conn, workspace_id, owner_username):
    cursor = conn.execute(
        'SELECT id FROM workspace_members WHERE workspace_id = ? AND username = ?',
        (workspace_id, owner_username)
    )
    existing = cursor.fetchone()
    if existing:
        conn.execute(
            '''
            UPDATE workspace_members
            SET role = ?, status = ?
            WHERE workspace_id = ? AND username = ?
            ''',
            ('owner', 'active', workspace_id, owner_username)
        )
    else:
        conn.execute(
            '''
            INSERT INTO workspace_members (workspace_id, username, role, status, created_at)
            VALUES (?, ?, ?, ?, ?)
            ''',
            (workspace_id, owner_username, 'owner', 'active', utcnow_iso())
        )


def get_workspace_record(conn, workspace_id):
    cursor = conn.execute('SELECT * FROM workspaces WHERE id = ?', (workspace_id,))
    return row_to_dict(cursor.fetchone())


def workspace_belongs_to_user(conn, workspace_id, username):
    if not username:
        return False
    owner_cursor = conn.execute(
        'SELECT 1 FROM workspaces WHERE id = ? AND owner_username = ?',
        (workspace_id, username)
    )
    if owner_cursor.fetchone() is not None:
        return True
    cursor = conn.execute(
        '''
        SELECT 1
        FROM workspace_members
        WHERE workspace_id = ? AND username = ? AND status = 'active'
        ''',
        (workspace_id, username)
    )
    return cursor.fetchone() is not None


def get_or_create_default_workspace_id(conn, username):
    owner = str(username or '').strip()
    if not owner:
        return ''

    cursor = conn.execute(
        '''
        SELECT id
        FROM workspaces
        WHERE owner_username = ?
        ORDER BY created_at ASC, id ASC
        LIMIT 1
        ''',
        (owner,)
    )
    row = cursor.fetchone()
    row_data = row_to_dict(row)
    workspace_id = row_data.get('id') if row_data else ''
    if workspace_id:
        ensure_owner_membership(conn, workspace_id, owner)
        return workspace_id

    workspace_id = f'ws-{uuid.uuid4().hex[:12]}'
    now_iso = utcnow_iso()
    conn.execute(
        '''
        INSERT INTO workspaces (id, name, plan, owner_username, settings_json, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        ''',
        (
            workspace_id,
            f"{owner}'s Workspace",
            'Free',
            owner,
            workspace_settings_to_json(DEFAULT_WORKSPACE_SETTINGS),
            now_iso,
            now_iso
        )
    )
    ensure_owner_membership(conn, workspace_id, owner)
    return workspace_id


def backfill_documents_workspace_ids(conn):
    if not documents_column_exists(conn, 'workspace_id'):
        return

    cursor = conn.execute(
        '''
        SELECT DISTINCT username
        FROM documents
        WHERE username IS NOT NULL
          AND TRIM(username) <> ''
          AND (workspace_id IS NULL OR TRIM(workspace_id) = '')
        '''
    )
    usernames = [
        (row_to_dict(item).get('username') or '').strip()
        for item in cursor.fetchall()
    ]
    usernames = [item for item in usernames if item]

    for owner in usernames:
        workspace_id = get_or_create_default_workspace_id(conn, owner)
        if not workspace_id:
            continue
        conn.execute(
            '''
            UPDATE documents
            SET workspace_id = ?
            WHERE username = ?
              AND (workspace_id IS NULL OR TRIM(workspace_id) = '')
            ''',
            (workspace_id, owner)
        )


def get_workspace_details(conn, workspace_row, for_username=''):
    workspace = row_to_dict(workspace_row) or {}
    workspace_id = workspace.get('id', '')
    owner_username = workspace.get('owner_username', '')
    is_owner = bool(for_username and for_username == owner_username)
    settings = normalize_workspace_settings(workspace.get('settings_json'))

    members_cursor = conn.execute(
        '''
        SELECT username, role, status, created_at
        FROM workspace_members
        WHERE workspace_id = ? AND status = 'active'
        ORDER BY created_at ASC
        ''',
        (workspace_id,)
    )
    members = [row_to_dict(item) for item in members_cursor.fetchall()]

    pending_requests = []
    invitations = []
    if is_owner:
        invite_cursor = conn.execute(
            '''
            SELECT *
            FROM workspace_invitations
            WHERE workspace_id = ?
              AND status IN ('pending', 'requested')
            ORDER BY created_at DESC
            ''',
            (workspace_id,)
        )
        invitations = [serialize_invitation_row(item) for item in invite_cursor.fetchall()]
        pending_requests = [item for item in invitations if item.get('status') == 'requested']

    return {
        'id': workspace_id,
        'name': normalize_workspace_name(workspace.get('name', ''), owner_username),
        'plan': 'Free' if workspace.get('plan', '') in ('', '免费版') else workspace.get('plan', ''),
        'owner_username': owner_username,
        'settings': settings,
        'created_at': workspace.get('created_at', ''),
        'updated_at': workspace.get('updated_at', ''),
        'is_owner': is_owner,
        'members_count': len(members),
        'members': members if is_owner else [],
        'invites': invitations,
        'pending_requests': pending_requests,
    }


def check_document_access(conn, doc_row, username='', share_token=''):
    doc = row_to_dict(doc_row) or {}
    if is_document_soft_deleted(doc):
        return False, 'Document is in Trash'
    viewer = str(username or '').strip()
    safe_share_token = str(share_token or '').strip()
    doc_id = parse_int(doc.get('id', 0), 0, 0)
    workspace_id = str(doc.get('workspace_id') or '').strip()
    owner_username = str(doc.get('username') or '').strip()

    # Backward compatibility for legacy rows without workspace binding.
    if not workspace_id:
        if owner_username and viewer == owner_username:
            return True, ''
        if safe_share_token:
            token_ok, _, token_reason = validate_document_share_token(
                conn,
                doc_id,
                safe_share_token,
                mark_access=True
            )
            if token_ok:
                return True, ''
            if token_reason:
                return False, token_reason
        if owner_username and viewer and viewer != owner_username:
            return False, 'You do not have access to this document'
        return True, ''

    workspace_row = get_workspace_record(conn, workspace_id)
    settings = normalize_workspace_settings((workspace_row or {}).get('settings_json'))
    link_mode = settings.get('link_sharing_mode', DEFAULT_WORKSPACE_SETTINGS['link_sharing_mode'])

    if viewer and workspace_belongs_to_user(conn, workspace_id, viewer):
        return True, ''

    if safe_share_token and link_mode != 'restricted':
        token_ok, _, token_reason = validate_document_share_token(
            conn,
            doc_id,
            safe_share_token,
            mark_access=True
        )
        if token_ok:
            return True, ''
        if token_reason and link_mode == 'workspace':
            return False, token_reason

    if link_mode == 'public':
        return True, ''
    if link_mode == 'workspace':
        return False, 'This link is only available to workspace members'
    return False, 'Link sharing is restricted by workspace settings'


def get_document_link_sharing_mode(conn, doc_row):
    doc = row_to_dict(doc_row) or {}
    workspace_id = str(doc.get('workspace_id') or '').strip()
    if not workspace_id:
        return DEFAULT_WORKSPACE_SETTINGS['link_sharing_mode']
    workspace_row = get_workspace_record(conn, workspace_id)
    settings = normalize_workspace_settings((workspace_row or {}).get('settings_json'))
    return settings.get('link_sharing_mode', DEFAULT_WORKSPACE_SETTINGS['link_sharing_mode'])


def get_workspace_settings(conn, workspace_id):
    target_id = str(workspace_id or '').strip()
    if not target_id:
        return dict(DEFAULT_WORKSPACE_SETTINGS)
    workspace_row = get_workspace_record(conn, target_id)
    return normalize_workspace_settings((workspace_row or {}).get('settings_json'))


def can_user_manage_workspace_share_links(conn, workspace_id, username=''):
    safe_workspace_id = str(workspace_id or '').strip()
    actor = str(username or '').strip()
    if not safe_workspace_id or not actor:
        return False

    workspace_row = get_workspace_record(conn, safe_workspace_id)
    if not workspace_row:
        return False

    owner_username = str((workspace_row or {}).get('owner_username') or '').strip()
    if owner_username and owner_username == actor:
        return True
    if not workspace_belongs_to_user(conn, safe_workspace_id, actor):
        return False

    workspace_settings = normalize_workspace_settings((workspace_row or {}).get('settings_json'))
    return parse_bool(workspace_settings.get('allow_member_share_management'), False)


def user_can_manage_document_share_links(conn, doc_row, username=''):
    doc = row_to_dict(doc_row) or {}
    actor = str(username or '').strip()
    if not actor:
        return False

    workspace_id = str(doc.get('workspace_id') or '').strip()
    owner_username = str(doc.get('username') or '').strip()
    if workspace_id:
        return can_user_manage_workspace_share_links(conn, workspace_id, actor)
    if owner_username:
        return owner_username == actor
    return False


def count_active_document_share_links(conn, document_id):
    safe_doc_id = parse_int(document_id, 0, 0)
    if safe_doc_id <= 0:
        return 0

    expire_document_share_links(conn, safe_doc_id)
    conn.commit()
    cursor = conn.execute(
        '''
        SELECT COUNT(1) AS total
        FROM document_share_links
        WHERE document_id = ? AND status = 'active'
        ''',
        (safe_doc_id,)
    )
    row = row_to_dict(cursor.fetchone()) or {}
    return parse_int(row.get('total', 0), 0, 0)


def user_can_edit_document(conn, doc_row, username=''):
    doc = row_to_dict(doc_row) or {}
    editor = str(username or '').strip()
    if not editor:
        return False

    workspace_id = str(doc.get('workspace_id') or '').strip()
    owner_username = str(doc.get('username') or '').strip()
    if workspace_id:
        return workspace_belongs_to_user(conn, workspace_id, editor)
    if owner_username:
        return owner_username == editor
    return False


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def detect_mimetype(filename, file_ext=''):
    ext = (file_ext or '').lower().strip('.')
    if ext in MIME_BY_EXT:
        return MIME_BY_EXT[ext]
    guessed = mimetypes.guess_type(filename)[0]
    return guessed or 'application/octet-stream'


def normalize_canvas_domain(value):
    raw = str(value or '').strip()
    if not raw:
        raw = CANVAS_DEFAULT_DOMAIN
    if raw.startswith('http://') or raw.startswith('https://'):
        raw = raw.split('://', 1)[1]
    raw = raw.split('/', 1)[0].strip().lower()
    if ':' in raw:
        host, port = raw.split(':', 1)
        if port.isdigit():
            raw = host
    if not raw or '.' not in raw:
        return ''
    if not re.fullmatch(r'[a-z0-9.-]{3,255}', raw):
        return ''
    return raw


def canvas_headers(token):
    safe_token = str(token or '').strip()
    if not safe_token:
        return {}
    return {'Authorization': f'Bearer {safe_token}'}


def list_canvas_user_files(domain, headers, max_pages=5, per_page=100):
    safe_domain = normalize_canvas_domain(domain)
    if not safe_domain:
        raise ValueError('Invalid Canvas domain')

    safe_pages = max(1, min(20, int(max_pages or 5)))
    safe_per_page = max(1, min(100, int(per_page or 100)))
    next_url = f'https://{safe_domain}/api/v1/users/self/files?per_page={safe_per_page}'
    files = []
    page_count = 0

    while next_url and page_count < safe_pages:
        response = requests.get(next_url, headers=headers, timeout=20)
        response.raise_for_status()
        payload = response.json()
        if isinstance(payload, list):
            files.extend(payload)
        page_count += 1
        links = response.links if isinstance(response.links, dict) else {}
        next_url = str((links.get('next') or {}).get('url') or '').strip()

    return files


def build_canvas_file_item(raw):
    item = raw if isinstance(raw, dict) else {}
    file_id = item.get('id')
    filename = str(
        item.get('filename')
        or item.get('display_name')
        or item.get('name')
        or ''
    ).strip()
    if not filename or not allowed_file(filename):
        return None

    safe_size = 0
    try:
        safe_size = max(0, int(item.get('size') or 0))
    except Exception:
        safe_size = 0

    file_type = ''
    if '.' in filename:
        file_type = filename.rsplit('.', 1)[1].lower().strip('.')

    return {
        'id': file_id,
        'filename': filename,
        'size': safe_size,
        'url': str(item.get('url') or '').strip(),
        'updated_at': str(item.get('updated_at') or '').strip(),
        'content_type': str(item.get('content-type') or item.get('content_type') or '').strip(),
        'file_type': file_type,
    }


def normalize_newlines(value):
    text = value if isinstance(value, str) else str(value or '')
    return text.replace('\r\n', '\n').replace('\r', '\n')


def normalize_document_category(value):
    category = str(value or '').strip()
    if not category:
        return ''
    category = re.sub(r'\s+', ' ', category)
    return category[:80]


def infer_document_category(title, text_content=''):
    title_text = str(title or '')
    body_text = str(text_content or '')[:5000]
    source = f"{title_text}\n{body_text}".lower()
    if not source.strip():
        return DEFAULT_DOCUMENT_CATEGORY

    for category, keywords in CATEGORY_KEYWORDS.items():
        if any(keyword in source for keyword in keywords):
            return category
    return DEFAULT_DOCUMENT_CATEGORY


def sanitize_style_declarations(style_text):
    style_map = {}
    if not isinstance(style_text, str):
        return style_map

    for declaration in style_text.split(';'):
        if ':' not in declaration:
            continue
        prop, val = declaration.split(':', 1)
        prop = prop.strip().lower()
        val = val.strip()
        if prop not in EDITOR_ALLOWED_STYLE_PROPS or not val:
            continue

        lower_val = val.lower()
        if 'expression(' in lower_val or 'javascript:' in lower_val or 'url(' in lower_val:
            continue

        if prop == 'font-family':
            cleaned_parts = [
                part.strip().strip('"').strip("'")
                for part in val.split(',')
                if part.strip()
            ]
            if not cleaned_parts:
                continue
            val = ', '.join(cleaned_parts[:3])
        elif prop in ('width', 'height', 'margin-left'):
            if not re.fullmatch(r'\d+(?:\.\d+)?(px|pt|em|rem|%)', lower_val):
                continue
        elif prop == 'border-collapse':
            if lower_val not in ('collapse', 'separate'):
                continue
        elif prop == 'border':
            if not re.fullmatch(r'[\w\s.#()-]+', val):
                continue

        style_map[prop] = val

    return style_map


def style_map_to_inline(style_map):
    if not style_map:
        return ''
    return '; '.join(f'{k}: {v}' for k, v in style_map.items())


def parse_css_color(color_value):
    if not isinstance(color_value, str):
        return None
    value = color_value.strip().lower()
    if not value:
        return None

    if value in NAMED_COLORS:
        return NAMED_COLORS[value]

    if re.fullmatch(r'#?[0-9a-f]{6}', value):
        hex_color = value.lstrip('#')
        return tuple(int(hex_color[idx:idx + 2], 16) for idx in (0, 2, 4))

    rgb_match = re.fullmatch(
        r'rgb\(\s*(\d{1,3})\s*,\s*(\d{1,3})\s*,\s*(\d{1,3})\s*\)',
        value
    )
    if rgb_match:
        channels = [max(0, min(255, int(item))) for item in rgb_match.groups()]
        return tuple(channels)

    return None


def sanitize_int_attr(value, min_value=1, max_value=10000):
    raw = str(value or '').strip()
    if not raw or not raw.isdigit():
        return None
    number = int(raw)
    if number < min_value:
        return None
    return min(number, max_value)


def is_safe_link_href(href):
    value = str(href or '').strip()
    if not value:
        return False
    lower = value.lower()
    if lower.startswith(('javascript:', 'data:', 'vbscript:')):
        return False
    return True


def is_safe_image_src(src):
    value = str(src or '').strip()
    if not value:
        return False
    lower = value.lower()
    if lower.startswith(('javascript:', 'vbscript:')):
        return False
    if lower.startswith('data:'):
        return lower.startswith('data:image/')
    return True


def sanitize_colwidth_attr(value):
    raw = str(value or '').strip()
    if not raw:
        return None
    if re.fullmatch(r'\d+(,\d+)*', raw):
        return raw
    return None


def pick_highlight_index_from_css(color_value):
    rgb = parse_css_color(color_value)
    if not rgb:
        return None

    closest_index = None
    closest_distance = None
    for index, target_rgb in HIGHLIGHT_RGB_BY_INDEX.items():
        distance = (
            (rgb[0] - target_rgb[0]) ** 2
            + (rgb[1] - target_rgb[1]) ** 2
            + (rgb[2] - target_rgb[2]) ** 2
        )
        if closest_distance is None or distance < closest_distance:
            closest_distance = distance
            closest_index = index

    return closest_index


def parse_css_font_size_pt(size_value):
    if not isinstance(size_value, str):
        return None
    value = size_value.strip().lower()
    if not value:
        return None
    match = re.fullmatch(r'([0-9]+(?:\.[0-9]+)?)(pt|px|em|rem)?', value)
    if not match:
        return None

    amount = float(match.group(1))
    unit = (match.group(2) or 'pt')
    if unit == 'pt':
        return amount
    if unit == 'px':
        return amount * 0.75
    if unit in ('em', 'rem'):
        return amount * 12.0
    return amount


def css_alignment_from_docx_alignment(alignment):
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return 'center'
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return 'right'
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return 'justify'
    return ''


def docx_alignment_from_css(style_map):
    alignment = (style_map.get('text-align') or '').strip().lower()
    if alignment == 'center':
        return WD_ALIGN_PARAGRAPH.CENTER
    if alignment == 'right':
        return WD_ALIGN_PARAGRAPH.RIGHT
    if alignment == 'justify':
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def apply_block_style_to_paragraph(paragraph, style_map):
    paragraph.alignment = docx_alignment_from_css(style_map)
    margin_left = parse_css_font_size_pt(style_map.get('margin-left'))
    if isinstance(margin_left, (int, float)) and margin_left > 0:
        paragraph.paragraph_format.left_indent = Pt(margin_left)


def plaintext_to_html(content):
    text = normalize_newlines(content)
    lines = text.split('\n')
    if not lines:
        return '<p><br></p>'

    blocks = []
    for line in lines:
        if line == '':
            blocks.append('<p><br></p>')
        else:
            blocks.append(f'<p>{html_escape(line)}</p>')
    return ''.join(blocks) or '<p><br></p>'


def html_to_plaintext(content_html):
    if not isinstance(content_html, str) or not content_html.strip():
        return ''

    normalized_html = re.sub(r'(?i)<br\s*/?>', '\n', content_html)
    normalized_html = re.sub(
        r'(?i)</(p|div|li|h[1-6]|blockquote|pre|ul|ol|table|thead|tbody|tr|th|td)>',
        '\n',
        normalized_html
    )
    normalized_html = re.sub(r'(?i)<hr\s*/?>', '\n', normalized_html)

    try:
        root = lxml_html.fragment_fromstring(normalized_html, create_parent='div')
        text = root.text_content()
    except Exception:
        text = re.sub(r'<[^>]+>', '', normalized_html)

    text = text.replace('\xa0', ' ')
    text = normalize_newlines(text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def sanitize_editor_html(raw_html):
    if raw_html is None:
        return '<p><br></p>'

    source_html = str(raw_html).strip()
    if not source_html:
        return '<p><br></p>'

    source_html = re.sub(r'(?is)<(script|style)[^>]*>.*?</\1>', '', source_html)
    source_html = re.sub(r'(?is)<!--.*?-->', '', source_html)

    try:
        root = lxml_html.fragment_fromstring(source_html, create_parent='div')
    except Exception:
        return plaintext_to_html(source_html)

    present_tags = {el.tag.lower() for el in root.iter() if isinstance(el.tag, str)}
    strip_tags = [tag for tag in present_tags if tag not in EDITOR_ALLOWED_TAGS]
    if strip_tags:
        etree.strip_tags(root, *strip_tags)

    for el in root.iter():
        if not isinstance(el.tag, str):
            continue
        tag = el.tag.lower()
        style_map = sanitize_style_declarations(el.attrib.get('style', ''))

        sanitized_attrs = {}
        if style_map:
            sanitized_attrs['style'] = style_map_to_inline(style_map)

        if tag == 'a':
            href = (el.attrib.get('href') or '').strip()
            if is_safe_link_href(href):
                sanitized_attrs['href'] = href
                sanitized_attrs['target'] = '_blank'
                sanitized_attrs['rel'] = 'noopener noreferrer'
        elif tag == 'img':
            src = (el.attrib.get('src') or '').strip()
            if is_safe_image_src(src):
                sanitized_attrs['src'] = src

            alt = str(el.attrib.get('alt') or '').strip()
            if alt:
                sanitized_attrs['alt'] = alt[:200]

            title = str(el.attrib.get('title') or '').strip()
            if title:
                sanitized_attrs['title'] = title[:200]

            width = sanitize_int_attr(el.attrib.get('width'), 1, 4000)
            if width:
                sanitized_attrs['width'] = str(width)

            height = sanitize_int_attr(el.attrib.get('height'), 1, 4000)
            if height:
                sanitized_attrs['height'] = str(height)
        elif tag in ('th', 'td'):
            colspan = sanitize_int_attr(el.attrib.get('colspan'), 1, 20)
            if colspan and colspan > 1:
                sanitized_attrs['colspan'] = str(colspan)

            rowspan = sanitize_int_attr(el.attrib.get('rowspan'), 1, 100)
            if rowspan and rowspan > 1:
                sanitized_attrs['rowspan'] = str(rowspan)

            colwidth = sanitize_colwidth_attr(el.attrib.get('colwidth'))
            if colwidth:
                sanitized_attrs['colwidth'] = colwidth
        elif tag == 'col':
            span = sanitize_int_attr(el.attrib.get('span'), 1, 20)
            if span and span > 1:
                sanitized_attrs['span'] = str(span)

            width = sanitize_int_attr(el.attrib.get('width'), 1, 4000)
            if width:
                sanitized_attrs['width'] = str(width)

        el.attrib.clear()
        if sanitized_attrs:
            el.attrib.update(sanitized_attrs)

    serialized_parts = []
    if root.text and root.text.strip():
        serialized_parts.append(f'<p>{html_escape(root.text)}</p>')
    for child in root:
        serialized_parts.append(lxml_html.tostring(child, encoding='unicode', method='html'))

    sanitized_html = ''.join(serialized_parts).strip()
    return sanitized_html or '<p><br></p>'


def apply_run_style(run, style_ctx):
    run.bold = bool(style_ctx.get('bold'))
    run.italic = bool(style_ctx.get('italic'))
    run.underline = bool(style_ctx.get('underline'))
    run.font.strike = bool(style_ctx.get('strike'))
    has_subscript = bool(style_ctx.get('subscript'))
    has_superscript = bool(style_ctx.get('superscript'))
    if has_subscript and has_superscript:
        has_subscript = False
    run.font.subscript = has_subscript
    run.font.superscript = has_superscript

    font_size_pt = style_ctx.get('font_size_pt')
    if isinstance(font_size_pt, (int, float)) and font_size_pt > 0:
        run.font.size = Pt(font_size_pt)

    font_name = style_ctx.get('font_name')
    if isinstance(font_name, str) and font_name.strip():
        run.font.name = font_name.strip()

    rgb = style_ctx.get('color_rgb')
    if isinstance(rgb, tuple) and len(rgb) == 3:
        run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])

    highlight_index = style_ctx.get('highlight_index')
    if highlight_index in HIGHLIGHT_RGB_BY_INDEX:
        run.font.highlight_color = highlight_index


def add_text_to_paragraph(paragraph, text, style_ctx):
    if not text:
        return
    parts = text.split('\n')
    for idx, part in enumerate(parts):
        run = paragraph.add_run(part)
        apply_run_style(run, style_ctx)
        if idx < len(parts) - 1:
            run.add_break()


def merge_inline_style(base_style, node):
    style = dict(base_style or {})
    tag = node.tag.lower() if isinstance(node.tag, str) else ''

    if tag in ('strong', 'b'):
        style['bold'] = True
    elif tag in ('em', 'i'):
        style['italic'] = True
    elif tag == 'u':
        style['underline'] = True
    elif tag in ('s', 'strike', 'del'):
        style['strike'] = True
    elif tag == 'sub':
        style['subscript'] = True
        style['superscript'] = False
    elif tag == 'sup':
        style['superscript'] = True
        style['subscript'] = False
    elif tag == 'mark':
        style['highlight_index'] = WD_COLOR_INDEX.YELLOW
    elif tag == 'code':
        style['font_name'] = 'Courier New'
    elif tag == 'a':
        style['underline'] = True
        style['color_rgb'] = (29, 78, 216)

    style_map = sanitize_style_declarations(node.attrib.get('style', ''))
    font_weight = (style_map.get('font-weight') or '').lower()
    if font_weight == 'bold':
        style['bold'] = True
    elif font_weight.isdigit():
        style['bold'] = int(font_weight) >= 600

    if (style_map.get('font-style') or '').lower() == 'italic':
        style['italic'] = True

    decoration = (style_map.get('text-decoration') or '').lower()
    if 'underline' in decoration:
        style['underline'] = True
    if 'line-through' in decoration:
        style['strike'] = True

    color_rgb = parse_css_color(style_map.get('color'))
    if color_rgb:
        style['color_rgb'] = color_rgb

    highlight_index = pick_highlight_index_from_css(style_map.get('background-color'))
    if highlight_index:
        style['highlight_index'] = highlight_index

    vertical_align = (style_map.get('vertical-align') or '').strip().lower()
    if vertical_align == 'sub':
        style['subscript'] = True
        style['superscript'] = False
    elif vertical_align in ('super', 'sup'):
        style['superscript'] = True
        style['subscript'] = False

    font_size_pt = parse_css_font_size_pt(style_map.get('font-size'))
    if font_size_pt:
        style['font_size_pt'] = font_size_pt

    font_family = style_map.get('font-family')
    if font_family:
        style['font_name'] = font_family.split(',')[0].strip()

    return style


def add_inline_node_to_paragraph(paragraph, node, inherited_style=None):
    current_style = merge_inline_style(inherited_style or {}, node)

    if node.text:
        add_text_to_paragraph(paragraph, node.text, current_style)

    for child in node:
        child_tag = child.tag.lower() if isinstance(child.tag, str) else ''
        if child_tag == 'br':
            paragraph.add_run().add_break()
        elif child_tag == 'img':
            alt = (child.attrib.get('alt') or '').strip()
            src = (child.attrib.get('src') or '').strip()
            placeholder = alt or src or 'Image'
            add_text_to_paragraph(paragraph, f'[Image] {placeholder}', current_style)
        elif child_tag == 'hr':
            add_text_to_paragraph(paragraph, '------------------------------', current_style)
        elif child_tag in BLOCK_TAGS.union({'ul', 'ol'}):
            nested_text = html_to_plaintext(lxml_html.tostring(child, encoding='unicode', method='html'))
            if nested_text:
                paragraph.add_run().add_break()
                add_text_to_paragraph(paragraph, nested_text, current_style)
        else:
            add_inline_node_to_paragraph(paragraph, child, current_style)

        if child.tail:
            add_text_to_paragraph(paragraph, child.tail, current_style)


def append_html_element_to_docx(document, element):
    if not isinstance(element.tag, str):
        return

    tag = element.tag.lower()
    style_map = sanitize_style_declarations(element.attrib.get('style', ''))

    if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(tag[1])
        paragraph = document.add_paragraph(style=f'Heading {min(level, 6)}')
        apply_block_style_to_paragraph(paragraph, style_map)
        add_inline_node_to_paragraph(paragraph, element, {})
        return

    if tag in ('ul', 'ol'):
        list_style = 'List Number' if tag == 'ol' else 'List Bullet'
        items = [child for child in element if isinstance(child.tag, str) and child.tag.lower() == 'li']
        if not items and (element.text or '').strip():
            items = [element]
        for item in items:
            paragraph = document.add_paragraph(style=list_style)
            add_inline_node_to_paragraph(paragraph, item, {})
        return

    if tag == 'blockquote':
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.space_after = Pt(6)
        apply_block_style_to_paragraph(paragraph, style_map)
        add_inline_node_to_paragraph(paragraph, element, {})
        return

    if tag == 'pre':
        text = normalize_newlines(element.text_content())
        lines = text.split('\n')
        if not lines:
            lines = ['']
        for line in lines:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = 'Courier New'
        return

    if tag == 'li':
        paragraph = document.add_paragraph(style='List Bullet')
        add_inline_node_to_paragraph(paragraph, element, {})
        return

    if tag == 'table':
        row_elements = []
        for child in element.iter():
            child_tag = child.tag.lower() if isinstance(child.tag, str) else ''
            if child_tag == 'tr':
                row_elements.append(child)

        if not row_elements:
            return

        max_cols = 1
        parsed_rows = []
        for row_el in row_elements:
            cells = [
                cell for cell in row_el
                if isinstance(cell.tag, str) and cell.tag.lower() in ('td', 'th')
            ]
            if not cells:
                continue
            parsed_rows.append(cells)
            max_cols = max(max_cols, len(cells))

        if not parsed_rows:
            return

        table = document.add_table(rows=len(parsed_rows), cols=max_cols)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass

        for row_idx, row_cells in enumerate(parsed_rows):
            for col_idx in range(max_cols):
                target_cell = table.cell(row_idx, col_idx)
                if col_idx >= len(row_cells):
                    target_cell.text = ''
                    continue
                cell_text = normalize_newlines(row_cells[col_idx].text_content()).strip()
                target_cell.text = cell_text
        return

    if tag == 'img':
        alt = (element.attrib.get('alt') or '').strip()
        src = (element.attrib.get('src') or '').strip()
        placeholder = alt or src or 'Image'
        document.add_paragraph(f'[Image] {placeholder}')
        return

    if tag == 'hr':
        document.add_paragraph('------------------------------')
        return

    paragraph = document.add_paragraph()
    apply_block_style_to_paragraph(paragraph, style_map)
    add_inline_node_to_paragraph(paragraph, element, {})


def create_docx_bytes_from_html(content_html, fallback_text=''):
    html_content = sanitize_editor_html(content_html)
    document = docx.Document()

    try:
        root = lxml_html.fragment_fromstring(html_content, create_parent='div')
    except Exception:
        root = None

    if root is None:
        lines = normalize_newlines(fallback_text).split('\n')
        for line in lines or ['']:
            document.add_paragraph(line)
    else:
        if root.text and root.text.strip():
            document.add_paragraph(root.text.strip())

        for child in root:
            append_html_element_to_docx(document, child)
            if child.tail and child.tail.strip():
                document.add_paragraph(child.tail.strip())

        if not document.paragraphs:
            fallback_lines = normalize_newlines(fallback_text).split('\n')
            for line in fallback_lines or ['']:
                document.add_paragraph(line)

    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream.read()


def run_to_html(run):
    raw_text = run.text or ''
    if raw_text == '':
        return ''

    chunk = html_escape(raw_text).replace('\n', '<br/>')
    style_parts = []

    font_size = run.font.size.pt if run.font and run.font.size else None
    if font_size:
        style_parts.append(f'font-size: {round(font_size, 2)}pt')

    font_name = run.font.name if run.font else ''
    if font_name:
        style_parts.append(f'font-family: {html_escape(font_name)}')

    font_color = run.font.color.rgb if run.font and run.font.color else None
    if font_color:
        color_hex = str(font_color)
        if re.fullmatch(r'[0-9A-Fa-f]{6}', color_hex):
            style_parts.append(f'color: #{color_hex}')

    highlight_color = run.font.highlight_color if run.font else None
    if highlight_color in HIGHLIGHT_RGB_BY_INDEX:
        rgb = HIGHLIGHT_RGB_BY_INDEX[highlight_color]
        style_parts.append(f'background-color: #{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}')

    if style_parts:
        chunk = f'<span style="{"; ".join(style_parts)}">{chunk}</span>'

    if run.bold:
        chunk = f'<strong>{chunk}</strong>'
    if run.italic:
        chunk = f'<em>{chunk}</em>'
    if run.underline:
        chunk = f'<u>{chunk}</u>'
    if run.font and run.font.strike:
        chunk = f'<s>{chunk}</s>'
    if run.font and run.font.subscript:
        chunk = f'<sub>{chunk}</sub>'
    if run.font and run.font.superscript:
        chunk = f'<sup>{chunk}</sup>'

    return chunk


def paragraph_to_html_block(paragraph):
    style_name = (paragraph.style.name if paragraph.style else '').strip().lower()
    alignment = css_alignment_from_docx_alignment(paragraph.alignment)
    style_attr = f' style="text-align: {alignment}"' if alignment else ''

    inline_html = ''.join(run_to_html(run) for run in paragraph.runs)
    if not inline_html:
        inline_html = html_escape(paragraph.text or '').replace('\n', '<br/>')
    if not inline_html:
        inline_html = '<br/>'

    if 'list bullet' in style_name:
        return 'ul', f'<li{style_attr}>{inline_html}</li>'
    if 'list number' in style_name:
        return 'ol', f'<li{style_attr}>{inline_html}</li>'

    heading_match = re.search(r'heading\s*([1-6])', style_name)
    if heading_match:
        level = heading_match.group(1)
        return '', f'<h{level}{style_attr}>{inline_html}</h{level}>'

    return '', f'<p{style_attr}>{inline_html}</p>'


def extract_docx_content(filepath):
    document = docx.Document(filepath)
    plain_lines = []
    html_parts = []
    list_type = ''
    list_items = []

    def flush_list():
        nonlocal list_type, list_items
        if list_type and list_items:
            html_parts.append(f'<{list_type}>{"".join(list_items)}</{list_type}>')
        list_type = ''
        list_items = []

    for paragraph in document.paragraphs:
        plain_lines.append(paragraph.text or '')
        block_type, block_html = paragraph_to_html_block(paragraph)
        if block_type in ('ul', 'ol'):
            if list_type and list_type != block_type:
                flush_list()
            list_type = block_type
            list_items.append(block_html)
        else:
            flush_list()
            html_parts.append(block_html)

    flush_list()

    for table in document.tables:
        row_html = []
        for row in table.rows:
            cell_html = []
            cell_text_parts = []
            for cell in row.cells:
                text_value = normalize_newlines(cell.text or '').strip()
                cell_text_parts.append(text_value)
                cell_html.append(f'<td>{html_escape(text_value) if text_value else "<br/>"}</td>')
            plain_lines.append(' | '.join(cell_text_parts).strip())
            row_html.append(f'<tr>{"".join(cell_html)}</tr>')
        if row_html:
            html_parts.append(f'<table><tbody>{"".join(row_html)}</tbody></table>')

    plain_text = '\n'.join(plain_lines)
    html_content = ''.join(html_parts).strip() or plaintext_to_html(plain_text)
    return plain_text, sanitize_editor_html(html_content)


def extract_text_content(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def extract_document_content(filepath, ext):
    file_ext = (ext or '').lower().strip('.')
    text = ''
    content_html = ''

    try:
        if file_ext == 'docx':
            text, content_html = extract_docx_content(filepath)
        elif file_ext == 'pdf':
            with open(filepath, 'rb') as f:
                file_bytes = f.read()
            text = extract_text_from_pdf_bytes(file_bytes)
        elif file_ext == 'txt':
            text = extract_text_content(filepath)
            content_html = plaintext_to_html(text)
    except Exception as e:
        print(f"Error extracting content: {e}")
        text = "Text extraction failed."
        content_html = plaintext_to_html(text)

    if file_ext in ('docx', 'txt'):
        content_html = sanitize_editor_html(content_html or plaintext_to_html(text))

    return text, content_html


def extract_text(filepath, ext):
    text, _ = extract_document_content(filepath, ext)
    return text


def normalize_pdf_text(text):
    value = normalize_newlines(text or '')
    value = value.replace('\x00', ' ')
    value = re.sub(r'-\n(?=[A-Za-z])', '', value)
    value = re.sub(r'[ \t]+', ' ', value)
    # Recover common missing spaces between latin tokens.
    value = re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', value)
    value = re.sub(r'(?<=[A-Za-z])(?=\d)', ' ', value)
    value = re.sub(r'(?<=\d)(?=[A-Za-z])', ' ', value)
    value = re.sub(r'([。！？.!?])(?=[A-Za-z0-9])', r'\1 ', value)
    value = re.sub(r'\n{3,}', '\n\n', value)
    return value.strip()


def compute_pdf_text_quality_metrics(text):
    normalized = normalize_pdf_text(text or '')
    tokens = re.findall(r'\S+', normalized)
    token_count = len(tokens)
    char_count = len(normalized)
    if token_count <= 0:
        return {
            'char_count': char_count,
            'token_count': 0,
            'avg_token_len': 0.0,
            'long_token_ratio': 0.0,
            'cjk_ratio': 0.0,
            'line_count': 0,
        }

    avg_token_len = sum(len(item) for item in tokens) / max(1, token_count)
    long_token_count = sum(1 for item in tokens if len(item) >= 18)
    long_token_ratio = long_token_count / max(1, token_count)
    cjk_chars = re.findall(r'[\u3400-\u9fff]', normalized)
    cjk_ratio = len(cjk_chars) / max(1, char_count)
    line_count = len([line for line in normalized.split('\n') if line.strip()])
    return {
        'char_count': char_count,
        'token_count': token_count,
        'avg_token_len': avg_token_len,
        'long_token_ratio': long_token_ratio,
        'cjk_ratio': cjk_ratio,
        'line_count': line_count,
    }


def score_pdf_text_quality(text):
    metrics = compute_pdf_text_quality_metrics(text)
    char_count = metrics['char_count']
    token_count = metrics['token_count']
    avg_token_len = metrics['avg_token_len']
    long_token_ratio = metrics['long_token_ratio']
    cjk_ratio = metrics['cjk_ratio']
    line_count = metrics['line_count']

    if char_count <= 0 or token_count <= 0:
        return 0.0, metrics

    score = 0.0
    score += min(char_count / 650.0, 28.0)
    score += min(token_count / 90.0, 20.0)
    score += min(line_count / 80.0, 6.0)
    if cjk_ratio < 0.2:
        score -= max(0.0, (avg_token_len - 8.8)) * 2.8
        score -= long_token_ratio * 36.0
    else:
        score -= max(0.0, (avg_token_len - 12.0)) * 1.2
        score -= long_token_ratio * 10.0

    return round(score, 3), metrics


def should_try_pdf_ocr_fallback(text):
    score, metrics = score_pdf_text_quality(text)
    char_count = metrics['char_count']
    token_count = metrics['token_count']
    avg_token_len = metrics['avg_token_len']
    long_token_ratio = metrics['long_token_ratio']
    cjk_ratio = metrics['cjk_ratio']

    if char_count < 320 or token_count < 60:
        return True, score, metrics
    if cjk_ratio < 0.2:
        if avg_token_len >= 10.5:
            return True, score, metrics
        if long_token_ratio >= 0.12:
            return True, score, metrics
    if score < 16.0:
        return True, score, metrics
    return False, score, metrics


def extract_text_from_pdf_bytes_pymupdf(file_bytes):
    if fitz is None:
        return ''

    try:
        doc = fitz.open(stream=file_bytes, filetype='pdf')
    except Exception as e:
        print(f"PyMuPDF open failed: {e}")
        return ''

    page_outputs = []
    try:
        for page in doc:
            page_lines = []
            words = page.get_text('words') or []
            if words:
                grouped = {}
                for item in words:
                    if not isinstance(item, (list, tuple)) or len(item) < 8:
                        continue
                    token = str(item[4] or '').strip()
                    if not token:
                        continue
                    block_no = int(item[5])
                    line_no = int(item[6])
                    word_no = int(item[7])
                    key = (block_no, line_no)
                    grouped.setdefault(key, []).append((word_no, token))

                for key in sorted(grouped.keys()):
                    tokens = [token for _, token in sorted(grouped[key], key=lambda pair: pair[0])]
                    line = ' '.join(tokens).strip()
                    if line:
                        page_lines.append(line)

            if not page_lines:
                raw_text = page.get_text('text') or ''
                raw_text = normalize_newlines(raw_text).strip()
                if raw_text:
                    page_lines.append(raw_text)

            page_outputs.append('\n'.join(page_lines).strip())
    except Exception as e:
        print(f"PyMuPDF extraction failed: {e}")
        return ''
    finally:
        doc.close()

    result = '\n\n'.join([chunk for chunk in page_outputs if chunk]).strip()
    return normalize_pdf_text(result)


def extract_text_from_pdf_bytes_pypdf2(file_bytes):
    text = ""
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
    except Exception as e:
        print(f"PyPDF2 extraction failed: {e}")
        return ""
    return normalize_pdf_text(text)


def run_ocrmypdf_on_pdf_bytes(file_bytes):
    if not ENABLE_PDF_OCR_FALLBACK:
        return b'', 'ocrmypdf fallback disabled by ENABLE_PDF_OCR_FALLBACK'
    if not file_bytes:
        return b'', 'Empty PDF bytes'

    ocrmypdf_path = shutil.which(OCRMYPDF_BINARY)
    if not ocrmypdf_path:
        return b'', f'ocrmypdf binary not found: {OCRMYPDF_BINARY}'

    try:
        with tempfile.TemporaryDirectory(prefix='studyhub-pdfocr-') as tmp_dir:
            input_path = os.path.join(tmp_dir, 'input.pdf')
            output_path = os.path.join(tmp_dir, 'output.pdf')
            with open(input_path, 'wb') as f:
                f.write(file_bytes)

            cmd = [
                ocrmypdf_path,
                '--force-ocr',
                '--output-type', 'pdf',
                '--optimize', '0',
                '--quiet',
                '-l', OCRMYPDF_LANGUAGE,
                input_path,
                output_path,
            ]
            proc = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=OCRMYPDF_TIMEOUT_SECONDS,
                check=False,
            )
            if proc.returncode != 0:
                stderr = (proc.stderr or b'').decode('utf-8', errors='ignore').strip()
                stdout = (proc.stdout or b'').decode('utf-8', errors='ignore').strip()
                details = stderr or stdout or f'return code {proc.returncode}'
                return b'', f'ocrmypdf failed: {details[:260]}'

            if not os.path.exists(output_path):
                return b'', 'ocrmypdf did not produce output file'
            with open(output_path, 'rb') as f:
                return f.read(), ''
    except subprocess.TimeoutExpired:
        return b'', f'ocrmypdf timeout after {OCRMYPDF_TIMEOUT_SECONDS}s'
    except Exception as e:
        return b'', f'ocrmypdf runtime error: {e}'


def extract_text_from_pdf_bytes_with_meta(file_bytes):
    meta = {
        'extractor': 'none',
        'ocr_attempted': False,
        'ocr_used': False,
        'quality_score_before': 0.0,
        'quality_score_after': 0.0,
        'quality_metrics_before': {},
        'quality_metrics_after': {},
        'note': '',
    }
    if not file_bytes:
        meta['note'] = 'Empty PDF bytes'
        return "", meta

    primary_text = extract_text_from_pdf_bytes_pymupdf(file_bytes)
    extractor = 'pymupdf'
    if not primary_text:
        primary_text = extract_text_from_pdf_bytes_pypdf2(file_bytes)
        extractor = 'pypdf2'

    if not primary_text:
        meta['note'] = 'Primary PDF text extractors returned empty text'
        return "Text extraction failed.", meta

    before_score, before_metrics = score_pdf_text_quality(primary_text)
    meta['extractor'] = extractor
    meta['quality_score_before'] = before_score
    meta['quality_metrics_before'] = before_metrics

    should_ocr, _, _ = should_try_pdf_ocr_fallback(primary_text)
    if not should_ocr:
        meta['quality_score_after'] = before_score
        meta['quality_metrics_after'] = before_metrics
        return primary_text, meta

    ocr_pdf_bytes, ocr_error = run_ocrmypdf_on_pdf_bytes(file_bytes)
    meta['ocr_attempted'] = True
    if not ocr_pdf_bytes:
        meta['quality_score_after'] = before_score
        meta['quality_metrics_after'] = before_metrics
        meta['note'] = ocr_error
        return primary_text, meta

    ocr_text = extract_text_from_pdf_bytes_pymupdf(ocr_pdf_bytes) or extract_text_from_pdf_bytes_pypdf2(ocr_pdf_bytes)
    if not ocr_text:
        meta['quality_score_after'] = before_score
        meta['quality_metrics_after'] = before_metrics
        meta['note'] = 'ocrmypdf produced file but text extraction stayed empty'
        return primary_text, meta

    after_score, after_metrics = score_pdf_text_quality(ocr_text)
    meta['quality_score_after'] = after_score
    meta['quality_metrics_after'] = after_metrics
    if after_score >= before_score + 1.0:
        meta['ocr_used'] = True
        if ocr_error:
            meta['note'] = ocr_error
        return ocr_text, meta

    meta['note'] = 'OCR output not better than primary extraction'
    return primary_text, meta


def extract_text_from_pdf_bytes(file_bytes):
    text, _ = extract_text_from_pdf_bytes_with_meta(file_bytes)
    return text


def create_pdf_bytes_from_text(content):
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    except Exception as e:
        raise RuntimeError(
            f'PDF editing requires reportlab in current interpreter: {sys.executable}. '
            f'Install it with "{sys.executable} -m pip install reportlab".'
        ) from e

    text = content if isinstance(content, str) else str(content or '')
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    stream = io.BytesIO()
    pdf_canvas = canvas.Canvas(stream, pagesize=A4)
    page_width, page_height = A4

    font_name = 'Helvetica'
    font_size = 11
    try:
        pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
        font_name = 'STSong-Light'
    except Exception:
        font_name = 'Helvetica'

    pdf_canvas.setFont(font_name, font_size)
    left_margin = 48
    top_margin = 56
    bottom_margin = 56
    line_height = 16
    max_line_width = page_width - (left_margin * 2)

    def wrap_paragraph(paragraph):
        if paragraph == '':
            return ['']

        wrapped = []
        current = ''
        for char in paragraph:
            candidate = current + char
            try:
                text_width = pdfmetrics.stringWidth(candidate, font_name, font_size)
            except Exception:
                text_width = len(candidate) * font_size * 0.6

            if text_width <= max_line_width or not current:
                current = candidate
            else:
                wrapped.append(current)
                current = char

        if current or not wrapped:
            wrapped.append(current)
        return wrapped

    y = page_height - top_margin
    paragraphs = text.split('\n')
    if not paragraphs:
        paragraphs = ['']

    for paragraph in paragraphs:
        for line in wrap_paragraph(paragraph):
            if y < bottom_margin:
                pdf_canvas.showPage()
                pdf_canvas.setFont(font_name, font_size)
                y = page_height - top_margin

            draw_text = line
            if font_name == 'Helvetica':
                draw_text = line.encode('latin-1', 'replace').decode('latin-1')

            pdf_canvas.drawString(left_margin, y, draw_text)
            y -= line_height

    pdf_canvas.save()
    stream.seek(0)
    return stream.read()


def build_editable_file_bytes(file_ext, content, content_html=''):
    ext = (file_ext or '').lower().strip('.')
    text = normalize_newlines(content if isinstance(content, str) else str(content or ''))

    if ext == 'txt':
        return text.encode('utf-8'), 'text/plain'

    if ext == 'docx':
        return create_docx_bytes_from_html(content_html, text), MIME_BY_EXT['docx']

    if ext == 'pdf':
        return create_pdf_bytes_from_text(text), MIME_BY_EXT['pdf']

    raise ValueError('Only txt, docx and pdf support direct source-file update right now.')


def write_file_bytes_to_storage(filename, file_bytes, mimetype='application/octet-stream'):
    if not filename:
        raise ValueError('filename is required')

    if S3_BUCKET and s3_client:
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=filename,
            Body=file_bytes,
            ContentType=mimetype
        )
        return

    local_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    with open(local_path, 'wb') as f:
        f.write(file_bytes)


def read_file_bytes_from_storage(filename):
    safe_filename = str(filename or '').strip()
    if not safe_filename:
        raise ValueError('filename is required')

    if S3_BUCKET and s3_client:
        s3_obj = s3_client.get_object(Bucket=S3_BUCKET, Key=safe_filename)
        return s3_obj['Body'].read()

    local_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
    with open(local_path, 'rb') as f:
        return f.read()


# 启动时运行初始化（放在辅助函数定义之后，避免前置调用未定义函数）
init_db()

# ================= API 路由接口 =================

@app.route('/api/auth/register', methods=['POST'])
def register():
    data = request.get_json()
    username = data.get('username')
    email = data.get('email')
    password = data.get('password')

    if not username or not password:
        return jsonify({'error': 'Missing fields'}), 400

    hashed_pw = generate_password_hash(password, method='pbkdf2:sha256')
    conn = get_db_connection()
    try:
        # DBWrapper 会自动处理占位符 ?
        conn.execute('INSERT INTO users (username, email, password_hash) VALUES (?, ?, ?)',
                     (username, email, hashed_pw))
        conn.commit()
        auth_token = create_auth_token(username)
        return jsonify({
            'message': 'User created successfully',
            'username': username,
            'email': email,
            'auth_token': auth_token,
        }), 201
    except Exception as e:
        return jsonify({'error': f'Registration failed (User may exist): {str(e)}'}), 409
    finally:
        conn.close()

@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.get_json()
    username_or_email = data.get('username')
    password = data.get('password')

    conn = get_db_connection()
    cursor = conn.execute('SELECT * FROM users WHERE username = ? OR email = ?', 
                        (username_or_email, username_or_email))
    user = cursor.fetchone()
    conn.close()

    if user and check_password_hash(user['password_hash'], password):
        user_email = user.get('email') if hasattr(user, 'get') else user['email']
        auth_token = create_auth_token(user['username'])
        return jsonify({
            'message': 'Login successful',
            'username': user['username'],
            'email': user_email,
            'auth_token': auth_token,
        }), 200
    else:
        return jsonify({'error': 'Invalid credentials'}), 401

@app.route('/api/auth/google', methods=['POST'])
def google_login():
    try:
        data = request.get_json()
        token = data.get('token')
        
        id_info = id_token.verify_oauth2_token(token, google_requests.Request(), GOOGLE_CLIENT_ID)
        email = id_info['email']
        name = id_info.get('name', email.split('@')[0])
        
        conn = get_db_connection()
        cursor = conn.execute('SELECT * FROM users WHERE email = ?', (email,))
        user = cursor.fetchone()
        
        if user is None:
            username = f"{name.split()[0]}_{uuid.uuid4().hex[:4]}"
            random_password = uuid.uuid4().hex
            hashed_password = generate_password_hash(random_password, method='pbkdf2:sha256')
            try:
                conn.execute('INSERT INTO users (username, email, password_hash) VALUES (?, ?, ?)',
                             (username, email, hashed_password))
                conn.commit()
                cursor = conn.execute('SELECT * FROM users WHERE email = ?', (email,))
                user = cursor.fetchone()
            except Exception as e:
                conn.close()
                return jsonify({'error': f'Register failed: {str(e)}'}), 500
        conn.close()
        user_email = user.get('email') if hasattr(user, 'get') else user['email']
        auth_token = create_auth_token(user['username'])
        return jsonify({
            'message': 'Login successful',
            'username': user['username'],
            'email': user_email,
            'auth_token': auth_token,
        }), 200
    except ValueError:
        return jsonify({'error': 'Invalid Google token'}), 401
    except Exception as e:
        print(f"Google login error: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/workspaces', methods=['GET'])
def get_workspaces():
    username = (request.args.get('username') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500

    try:
        expire_workspace_invitations(conn)

        owned_cursor = conn.execute(
            'SELECT * FROM workspaces WHERE owner_username = ? ORDER BY created_at DESC',
            (username,)
        )
        owned_rows = [row_to_dict(item) for item in owned_cursor.fetchall()]

        member_cursor = conn.execute(
            '''
            SELECT workspace_id
            FROM workspace_members
            WHERE username = ? AND status = 'active'
            ''',
            (username,)
        )
        member_workspace_ids = [
            row_to_dict(item).get('workspace_id')
            for item in member_cursor.fetchall()
            if row_to_dict(item).get('workspace_id')
        ]
        owned_ids = {item.get('id') for item in owned_rows}
        extra_ids = [workspace_id for workspace_id in member_workspace_ids if workspace_id not in owned_ids]

        extra_rows = []
        if extra_ids:
            placeholders = ','.join(['?'] * len(extra_ids))
            query = f'SELECT * FROM workspaces WHERE id IN ({placeholders}) ORDER BY created_at DESC'
            extra_cursor = conn.execute(query, tuple(extra_ids))
            extra_rows = [row_to_dict(item) for item in extra_cursor.fetchall()]

        workspace_rows = [*owned_rows, *extra_rows]
        if not workspace_rows:
            # Automatically create a default workspace on first login.
            now_iso = utcnow_iso()
            workspace_id = f'ws-{uuid.uuid4().hex[:12]}'
            conn.execute(
                '''
                INSERT INTO workspaces (id, name, plan, owner_username, settings_json, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    workspace_id,
                    f"{username}'s Workspace",
                    'Free',
                    username,
                    workspace_settings_to_json(DEFAULT_WORKSPACE_SETTINGS),
                    now_iso,
                    now_iso
                )
            )
            ensure_owner_membership(conn, workspace_id, username)
            conn.commit()
            workspace_rows = [{
                'id': workspace_id,
                'name': f"{username}'s Workspace",
                'plan': 'Free',
                'owner_username': username,
                'settings_json': workspace_settings_to_json(DEFAULT_WORKSPACE_SETTINGS),
                'created_at': now_iso,
                'updated_at': now_iso,
            }]
        else:
            conn.commit()

        payload = [get_workspace_details(conn, item, username) for item in workspace_rows]
        return jsonify(payload), 200
    finally:
        conn.close()


@app.route('/api/workspaces', methods=['POST'])
def create_workspace():
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()
    name = (data.get('name') or '').strip() or f"{username}'s Workspace"
    plan = (data.get('plan') or '').strip() or 'Free'
    settings_json = workspace_settings_to_json(data.get('settings') or DEFAULT_WORKSPACE_SETTINGS)
    if not username:
        return jsonify({'error': 'username is required'}), 400

    workspace_id = f'ws-{uuid.uuid4().hex[:12]}'
    now_iso = utcnow_iso()

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        conn.execute(
            '''
            INSERT INTO workspaces (id, name, plan, owner_username, settings_json, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            ''',
            (
                workspace_id,
                name,
                plan,
                username,
                settings_json,
                now_iso,
                now_iso
            )
        )
        ensure_owner_membership(conn, workspace_id, username)
        conn.commit()
        workspace_row = get_workspace_record(conn, workspace_id)
        return jsonify(get_workspace_details(conn, workspace_row, username)), 201
    finally:
        conn.close()


@app.route('/api/workspaces/<workspace_id>', methods=['PUT'])
def update_workspace(workspace_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()
    name = (data.get('name') or '').strip()
    incoming_settings = data.get('settings')
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        workspace_row = get_workspace_record(conn, workspace_id)
        if not workspace_row:
            return jsonify({'error': 'Workspace not found'}), 404
        if workspace_row.get('owner_username') != username:
            return jsonify({'error': 'Only workspace owner can update settings'}), 403
        if not name and incoming_settings is None:
            return jsonify({'error': 'name or settings is required'}), 400

        next_name = name or normalize_workspace_name(workspace_row.get('name', ''), username)

        existing_settings = normalize_workspace_settings(workspace_row.get('settings_json'))
        if isinstance(incoming_settings, dict):
            merged_settings = {**existing_settings, **incoming_settings}
        else:
            merged_settings = existing_settings
        settings_json = workspace_settings_to_json(merged_settings)

        conn.execute(
            'UPDATE workspaces SET name = ?, settings_json = ?, updated_at = ? WHERE id = ?',
            (next_name, settings_json, utcnow_iso(), workspace_id)
        )
        conn.commit()
        updated = get_workspace_record(conn, workspace_id)
        return jsonify(get_workspace_details(conn, updated, username)), 200
    finally:
        conn.close()


@app.route('/api/workspaces/<workspace_id>/invitations', methods=['GET'])
def list_workspace_invitations(workspace_id):
    username = (request.args.get('username') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        workspace_row = get_workspace_record(conn, workspace_id)
        if not workspace_row:
            return jsonify({'error': 'Workspace not found'}), 404
        if workspace_row.get('owner_username') != username:
            return jsonify({'error': 'Only workspace owner can view invitations'}), 403

        expire_workspace_invitations(conn, workspace_id)
        conn.commit()
        cursor = conn.execute(
            '''
            SELECT *
            FROM workspace_invitations
            WHERE workspace_id = ?
            ORDER BY created_at DESC
            LIMIT 200
            ''',
            (workspace_id,)
        )
        invitations = [serialize_invitation_row(item) for item in cursor.fetchall()]
        return jsonify(invitations), 200
    finally:
        conn.close()


@app.route('/api/workspaces/<workspace_id>/invitations', methods=['POST'])
def create_workspace_invitations(workspace_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()
    raw_emails = data.get('emails', [])
    expiry_days = data.get('expiry_days', INVITE_EXPIRY_DAYS)

    if not username:
        return jsonify({'error': 'username is required'}), 400

    if isinstance(raw_emails, str):
        candidates = [item.strip() for item in re.split(r'[\n,;]+', raw_emails) if item.strip()]
    elif isinstance(raw_emails, list):
        candidates = [str(item).strip() for item in raw_emails if str(item).strip()]
    else:
        return jsonify({'error': 'emails must be an array or a comma-separated string'}), 400

    normalized_emails = []
    invalid_emails = []
    for item in candidates:
        email = normalize_email(item)
        if is_valid_email(email):
            normalized_emails.append(email)
        else:
            invalid_emails.append(item)
    normalized_emails = list(dict.fromkeys(normalized_emails))

    if not normalized_emails:
        return jsonify({'error': 'No valid email addresses provided', 'invalid_emails': invalid_emails}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        workspace_row = get_workspace_record(conn, workspace_id)
        if not workspace_row:
            return jsonify({'error': 'Workspace not found'}), 404
        workspace_settings = normalize_workspace_settings(workspace_row.get('settings_json'))
        is_owner = workspace_row.get('owner_username') == username
        can_invite_as_member = workspace_settings.get('allow_member_invites') and workspace_belongs_to_user(
            conn, workspace_id, username
        )
        if not is_owner and not can_invite_as_member:
            return jsonify({'error': 'Only workspace owner can invite members'}), 403

        requested_expiry = data.get('expiry_days', None)
        if requested_expiry is None or str(requested_expiry).strip() == '':
            expiry_days = workspace_settings.get('default_invite_expiry_days', INVITE_EXPIRY_DAYS)
        else:
            expiry_days = requested_expiry

        expire_workspace_invitations(conn, workspace_id)
        now_iso = utcnow_iso()
        expires_at = expires_at_for_days(expiry_days)
        created_items = []
        send_errors = []

        for email in normalized_emails:
            conn.execute(
                '''
                UPDATE workspace_invitations
                SET status = 'cancelled', reviewed_by = ?, reviewed_at = ?, review_note = ?
                WHERE workspace_id = ?
                  AND email = ?
                  AND status IN ('pending', 'requested')
                ''',
                (username, now_iso, 'Replaced by newer invitation', workspace_id, email)
            )

            token = create_invite_token()
            conn.execute(
                '''
                INSERT INTO workspace_invitations (
                    workspace_id, email, token, status, expires_at, created_at
                ) VALUES (?, ?, ?, ?, ?, ?)
                ''',
                (workspace_id, email, token, 'pending', expires_at, now_iso)
            )

            invite_row_cursor = conn.execute(
                '''
                SELECT *
                FROM workspace_invitations
                WHERE token = ?
                ''',
                (token,)
            )
            invite_row = row_to_dict(invite_row_cursor.fetchone())
            invite_payload = serialize_invitation_row(invite_row)

            ok, send_error = send_workspace_invite_email(
                email,
                workspace_row.get('name', ''),
                username,
                invite_payload.get('invite_url', ''),
                expires_at
            )
            invite_payload['email_sent'] = bool(ok)
            if not ok:
                invite_payload['email_error'] = send_error
                send_errors.append({'email': email, 'error': send_error})

            created_items.append(invite_payload)

        conn.commit()
        return jsonify({
            'workspace_id': workspace_id,
            'created': created_items,
            'invalid_emails': invalid_emails,
            'send_errors': send_errors,
            'requires_owner_confirmation': True,
        }), 201
    finally:
        conn.close()


@app.route('/api/workspaces/<workspace_id>/invitations/<int:invitation_id>', methods=['DELETE'])
def cancel_workspace_invitation(workspace_id, invitation_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or request.args.get('username') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        workspace_row = get_workspace_record(conn, workspace_id)
        if not workspace_row:
            return jsonify({'error': 'Workspace not found'}), 404
        if workspace_row.get('owner_username') != username:
            return jsonify({'error': 'Only workspace owner can cancel invitations'}), 403

        cursor = conn.execute(
            '''
            SELECT *
            FROM workspace_invitations
            WHERE id = ? AND workspace_id = ?
            ''',
            (invitation_id, workspace_id)
        )
        invitation = row_to_dict(cursor.fetchone())
        if not invitation:
            return jsonify({'error': 'Invitation not found'}), 404

        if invitation.get('status') in ('approved', 'rejected', 'expired', 'cancelled'):
            return jsonify({'error': f'Invitation is already {invitation.get("status")}'}), 400

        conn.execute(
            '''
            UPDATE workspace_invitations
            SET status = 'cancelled', reviewed_by = ?, reviewed_at = ?, review_note = ?
            WHERE id = ?
            ''',
            (username, utcnow_iso(), 'Cancelled by owner', invitation_id)
        )
        conn.commit()

        refreshed_cursor = conn.execute(
            'SELECT * FROM workspace_invitations WHERE id = ?',
            (invitation_id,)
        )
        refreshed = serialize_invitation_row(refreshed_cursor.fetchone())
        return jsonify(refreshed), 200
    finally:
        conn.close()


@app.route('/api/workspaces/<workspace_id>/invitations/<int:invitation_id>/review', methods=['POST'])
def review_workspace_invitation(workspace_id, invitation_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()
    action = (data.get('action') or '').strip().lower()
    note = (data.get('note') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400
    if action not in ('approve', 'reject'):
        return jsonify({'error': 'action must be approve or reject'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        workspace_row = get_workspace_record(conn, workspace_id)
        if not workspace_row:
            return jsonify({'error': 'Workspace not found'}), 404
        if workspace_row.get('owner_username') != username:
            return jsonify({'error': 'Only workspace owner can review requests'}), 403

        expire_workspace_invitations(conn, workspace_id)
        cursor = conn.execute(
            '''
            SELECT *
            FROM workspace_invitations
            WHERE id = ? AND workspace_id = ?
            ''',
            (invitation_id, workspace_id)
        )
        invitation = row_to_dict(cursor.fetchone())
        if not invitation:
            return jsonify({'error': 'Invitation not found'}), 404
        if invitation.get('status') != 'requested':
            return jsonify({'error': f'Invitation status must be requested, current: {invitation.get("status")}'}), 400

        requested_username = (invitation.get('requested_username') or '').strip()
        if action == 'approve':
            if not requested_username:
                return jsonify({'error': 'No applicant found for this invitation'}), 400
            user_cursor = conn.execute('SELECT username FROM users WHERE username = ?', (requested_username,))
            applicant = user_cursor.fetchone()
            if not applicant:
                return jsonify({'error': 'Applicant account not found'}), 404

            ensure_owner_membership(conn, workspace_id, workspace_row.get('owner_username', ''))
            member_cursor = conn.execute(
                '''
                SELECT id
                FROM workspace_members
                WHERE workspace_id = ? AND username = ?
                ''',
                (workspace_id, requested_username)
            )
            existing_member = member_cursor.fetchone()
            if existing_member:
                conn.execute(
                    '''
                    UPDATE workspace_members
                    SET status = 'active', role = 'member'
                    WHERE workspace_id = ? AND username = ?
                    ''',
                    (workspace_id, requested_username)
                )
            else:
                conn.execute(
                    '''
                    INSERT INTO workspace_members (workspace_id, username, role, status, created_at)
                    VALUES (?, ?, ?, ?, ?)
                    ''',
                    (workspace_id, requested_username, 'member', 'active', utcnow_iso())
                )

            next_status = 'approved'
        else:
            next_status = 'rejected'

        conn.execute(
            '''
            UPDATE workspace_invitations
            SET status = ?, reviewed_by = ?, reviewed_at = ?, review_note = ?
            WHERE id = ?
            ''',
            (next_status, username, utcnow_iso(), note, invitation_id)
        )
        conn.commit()

        updated_cursor = conn.execute('SELECT * FROM workspace_invitations WHERE id = ?', (invitation_id,))
        updated_invitation = serialize_invitation_row(updated_cursor.fetchone())
        return jsonify(updated_invitation), 200
    finally:
        conn.close()


@app.route('/api/invitations/<token>', methods=['GET'])
def get_invitation_by_token(token):
    safe_token = (token or '').strip()
    username = (request.args.get('username') or '').strip()
    if not safe_token:
        return jsonify({'error': 'token is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.execute(
            '''
            SELECT inv.*, ws.name AS workspace_name, ws.owner_username
            FROM workspace_invitations inv
            JOIN workspaces ws ON ws.id = inv.workspace_id
            WHERE inv.token = ?
            ''',
            (safe_token,)
        )
        invitation = row_to_dict(cursor.fetchone())
        if not invitation:
            return jsonify({'error': 'Invitation not found'}), 404

        if invitation.get('status') in ('pending', 'requested') and invitation_is_expired(invitation.get('expires_at')):
            conn.execute(
                'UPDATE workspace_invitations SET status = ? WHERE token = ?',
                ('expired', safe_token)
            )
            conn.commit()
            invitation['status'] = 'expired'

        can_request = False
        mismatch_reason = ''
        user_email = ''
        if username:
            user_cursor = conn.execute('SELECT email FROM users WHERE username = ?', (username,))
            user_row = row_to_dict(user_cursor.fetchone()) or {}
            user_email = normalize_email(user_row.get('email', ''))
            invite_email = normalize_email(invitation.get('email', ''))
            if not user_email:
                mismatch_reason = 'The current account has no bound email, so invitation ownership cannot be verified'
            elif user_email != invite_email:
                mismatch_reason = 'The current account email does not match the invited email'
            elif invitation.get('status') == 'pending':
                can_request = True
            elif invitation.get('status') == 'requested':
                can_request = invitation.get('requested_username') == username

        payload = serialize_invitation_row(invitation)
        payload.update({
            'workspace_name': normalize_workspace_name(invitation.get('workspace_name', ''), invitation.get('owner_username', '')),
            'owner_username': invitation.get('owner_username', ''),
            'requires_owner_confirmation': True,
            'can_request': can_request,
            'mismatch_reason': mismatch_reason,
            'viewer_username': username,
            'viewer_email': user_email,
        })
        return jsonify(payload), 200
    finally:
        conn.close()


@app.route('/api/invitations/<token>/request-join', methods=['POST'])
def request_join_by_invitation(token):
    safe_token = (token or '').strip()
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()
    if not safe_token:
        return jsonify({'error': 'token is required'}), 400
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.execute(
            '''
            SELECT inv.*, ws.name AS workspace_name, ws.owner_username
            FROM workspace_invitations inv
            JOIN workspaces ws ON ws.id = inv.workspace_id
            WHERE inv.token = ?
            ''',
            (safe_token,)
        )
        invitation = row_to_dict(cursor.fetchone())
        if not invitation:
            return jsonify({'error': 'Invitation not found'}), 404

        if invitation.get('status') in ('approved', 'rejected', 'expired', 'cancelled'):
            return jsonify({'error': f'Invitation is {invitation.get("status")}'}), 400
        if invitation_is_expired(invitation.get('expires_at')):
            conn.execute(
                'UPDATE workspace_invitations SET status = ? WHERE token = ?',
                ('expired', safe_token)
            )
            conn.commit()
            return jsonify({'error': 'Invitation has expired'}), 400

        user_cursor = conn.execute('SELECT email FROM users WHERE username = ?', (username,))
        user_row = row_to_dict(user_cursor.fetchone()) or {}
        user_email = normalize_email(user_row.get('email', ''))
        invite_email = normalize_email(invitation.get('email', ''))
        if not user_email:
            return jsonify({'error': 'Your account has no email address; cannot match invitation'}), 400
        if user_email != invite_email:
            return jsonify({'error': 'Your account email does not match this invitation'}), 403

        if invitation.get('status') == 'requested':
            if invitation.get('requested_username') == username:
                return jsonify({'message': 'Join request already submitted and pending owner approval'}), 200
            return jsonify({'error': 'This invitation is already requested by another account'}), 409

        conn.execute(
            '''
            UPDATE workspace_invitations
            SET status = 'requested', requested_username = ?, requested_at = ?, reviewed_by = NULL, reviewed_at = NULL, review_note = ''
            WHERE token = ?
            ''',
            (username, utcnow_iso(), safe_token)
        )
        conn.commit()

        refreshed_cursor = conn.execute(
            '''
            SELECT inv.*, ws.name AS workspace_name, ws.owner_username
            FROM workspace_invitations inv
            JOIN workspaces ws ON ws.id = inv.workspace_id
            WHERE inv.token = ?
            ''',
            (safe_token,)
        )
        refreshed = row_to_dict(refreshed_cursor.fetchone())
        payload = serialize_invitation_row(refreshed)
        payload.update({
            'workspace_name': refreshed.get('workspace_name', ''),
            'owner_username': refreshed.get('owner_username', ''),
            'requires_owner_confirmation': True,
        })
        return jsonify(payload), 200
    finally:
        conn.close()


@app.route('/api/documents', methods=['GET'])
def get_documents():
    username = (request.args.get('username') or '').strip()
    workspace_id = (request.args.get('workspace_id') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    query = (request.args.get('q') or '').strip().lower()
    tag_filter = (request.args.get('tag') or '').strip().lower()
    category_filter = (request.args.get('category') or '').strip().lower()
    start_date = (request.args.get('start_date') or '').strip()
    end_date = (request.args.get('end_date') or '').strip()
    file_type_filter = (request.args.get('file_type') or '').strip().lower().lstrip('.')
    include_meta = parse_bool(request.args.get('include_meta') or request.args.get('meta'), False)
    include_facets = parse_bool(request.args.get('include_facets'), False)
    limit = parse_int(request.args.get('limit'), 20, 1, 100)
    offset = parse_int(request.args.get('offset'), 0, 0)
    sort_key = (request.args.get('sort') or 'newest').strip().lower()
    order_by_map = {
        'newest': "uploaded_at DESC, id DESC",
        'oldest': "uploaded_at ASC, id ASC",
        'title_asc': "LOWER(COALESCE(title, '')) ASC, id ASC",
        'title_desc': "LOWER(COALESCE(title, '')) DESC, id DESC",
        'category_asc': "LOWER(COALESCE(category, '')) ASC, LOWER(COALESCE(title, '')) ASC, id ASC",
    }
    order_by_sql = order_by_map.get(sort_key, order_by_map['newest'])

    conn = get_db_connection()
    try:
        if workspace_id and not workspace_belongs_to_user(conn, workspace_id, username):
            return jsonify({'error': 'No access to this workspace'}), 403

        where_parts = [
            'username = ?',
            "COALESCE(deleted_at, '') = ''",
        ]
        params = [username]
        if workspace_id:
            where_parts.append('workspace_id = ?')
            params.append(workspace_id)
        if category_filter:
            where_parts.append("LOWER(COALESCE(category, '')) = ?")
            params.append(category_filter)
        if tag_filter:
            where_parts.append("(',' || LOWER(COALESCE(tags, '')) || ',') LIKE ?")
            params.append(f'%,{tag_filter},%')
        if query:
            where_parts.append(
                "("
                "LOWER(COALESCE(title, '')) LIKE ? OR "
                "LOWER(COALESCE(category, '')) LIKE ? OR "
                "LOWER(COALESCE(content, '')) LIKE ? OR "
                "LOWER(COALESCE(tags, '')) LIKE ?"
                ")"
            )
            like_query = f'%{query}%'
            params.extend([like_query, like_query, like_query, like_query])
        if start_date:
            where_parts.append("COALESCE(uploaded_at, '') >= ?")
            params.append(f'{start_date}T00:00:00')
        if end_date:
            where_parts.append("COALESCE(uploaded_at, '') <= ?")
            params.append(f'{end_date}T23:59:59')
        if file_type_filter:
            if file_type_filter in ('image', 'images'):
                image_exts = ('png', 'jpg', 'jpeg', 'webp', 'gif')
                placeholders = ','.join('?' for _ in image_exts)
                where_parts.append(f"LOWER(COALESCE(file_type, '')) IN ({placeholders})")
                params.extend(image_exts)
            elif file_type_filter in ('editable', 'editables'):
                editable_exts = ('txt', 'docx')
                placeholders = ','.join('?' for _ in editable_exts)
                where_parts.append(f"LOWER(COALESCE(file_type, '')) IN ({placeholders})")
                params.extend(editable_exts)
            elif re.fullmatch(r'[a-z0-9]{1,12}', file_type_filter):
                where_parts.append("LOWER(COALESCE(file_type, '')) = ?")
                params.append(file_type_filter)

        where_sql = ' AND '.join(where_parts) if where_parts else '1=1'

        total_cursor = conn.execute(
            f'''
            SELECT COUNT(1) AS total
            FROM documents
            WHERE {where_sql}
            ''',
            params
        )
        total_row = row_to_dict(total_cursor.fetchone()) or {}
        total = parse_int(total_row.get('total', 0), 0, 0)

        list_params = [*params, limit, offset]
        cursor = conn.execute(
            f'''
            SELECT *
            FROM documents
            WHERE {where_sql}
            ORDER BY {order_by_sql}
            LIMIT ? OFFSET ?
            ''',
            list_params
        )
        docs = [dict(doc) for doc in cursor.fetchall()]
        if include_meta:
            payload = {
                'items': docs,
                'total': total,
                'limit': limit,
                'offset': offset,
                'has_more': (offset + len(docs)) < total,
            }
            if include_facets:
                facet_cursor = conn.execute(
                    f'''
                    SELECT category, tags, file_type
                    FROM documents
                    WHERE {where_sql}
                    ''',
                    params
                )
                tag_set = set()
                category_set = set()
                file_type_counts = {}
                for row in facet_cursor.fetchall():
                    item = row_to_dict(row)
                    category = normalize_document_category((item or {}).get('category', ''))
                    category_set.add(category or DEFAULT_DOCUMENT_CATEGORY)
                    raw_tags = str((item or {}).get('tags') or '')
                    for raw_tag in raw_tags.split(','):
                        safe_tag = raw_tag.strip()
                        if safe_tag:
                            tag_set.add(safe_tag)
                    ext = str((item or {}).get('file_type') or '').strip().lower().strip('.')
                    if ext:
                        file_type_counts[ext] = file_type_counts.get(ext, 0) + 1
                        if ext in ('png', 'jpg', 'jpeg', 'webp', 'gif'):
                            file_type_counts['image'] = file_type_counts.get('image', 0) + 1
                        if ext in ('txt', 'docx'):
                            file_type_counts['editable'] = file_type_counts.get('editable', 0) + 1
                payload['facets'] = {
                    'tags': sorted(tag_set, key=lambda value: value.lower()),
                    'categories': sorted(category_set, key=lambda value: value.lower()),
                    'file_types': file_type_counts,
                }
            return jsonify(payload), 200
        return jsonify(docs), 200
    finally:
        conn.close()


@app.route('/api/documents/trash', methods=['GET'])
def get_trashed_documents():
    username = (request.args.get('username') or '').strip()
    workspace_id = (request.args.get('workspace_id') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    query = (request.args.get('q') or '').strip().lower()
    sort_key = (request.args.get('sort') or 'deleted_newest').strip().lower()
    order_by_map = {
        'deleted_newest': "deleted_at DESC, id DESC",
        'deleted_oldest': "deleted_at ASC, id ASC",
        'title_asc': "LOWER(COALESCE(title, '')) ASC, id ASC",
        'title_desc': "LOWER(COALESCE(title, '')) DESC, id DESC",
    }
    order_by_sql = order_by_map.get(sort_key, order_by_map['deleted_newest'])

    limit = parse_int(request.args.get('limit'), 100, 1, 300)
    offset = parse_int(request.args.get('offset'), 0, 0)

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500

    try:
        if workspace_id and not workspace_belongs_to_user(conn, workspace_id, username):
            return jsonify({'error': 'No access to this workspace'}), 403

        purge_result = purge_expired_trashed_documents(conn, username=username, workspace_id=workspace_id)

        where_parts = [
            'username = ?',
            "COALESCE(deleted_at, '') <> ''",
        ]
        params = [username]
        if workspace_id:
            where_parts.append('workspace_id = ?')
            params.append(workspace_id)
        if query:
            where_parts.append(
                "("
                "LOWER(COALESCE(title, '')) LIKE ? OR "
                "LOWER(COALESCE(filename, '')) LIKE ? OR "
                "LOWER(COALESCE(category, '')) LIKE ? OR "
                "LOWER(COALESCE(content, '')) LIKE ? OR "
                "LOWER(COALESCE(tags, '')) LIKE ?"
                ")"
            )
            like_query = f'%{query}%'
            params.extend([like_query, like_query, like_query, like_query, like_query])
        where_sql = ' AND '.join(where_parts)

        total_cursor = conn.execute(
            f'''
            SELECT COUNT(1) AS total
            FROM documents
            WHERE {where_sql}
            ''',
            tuple(params)
        )
        total_row = row_to_dict(total_cursor.fetchone()) or {}
        total = parse_int(total_row.get('total', 0), 0, 0)

        list_params = [*params, limit, offset]
        cursor = conn.execute(
            f'''
            SELECT *
            FROM documents
            WHERE {where_sql}
            ORDER BY {order_by_sql}
            LIMIT ? OFFSET ?
            ''',
            tuple(list_params)
        )
        items = [row_to_dict(item) for item in cursor.fetchall()]
        return jsonify({
            'items': items,
            'total': total,
            'limit': limit,
            'offset': offset,
            'q': query,
            'sort': sort_key,
            'retention_days': TRASH_RETENTION_DAYS,
            'purged_count': parse_int((purge_result or {}).get('purged_count', 0), 0, 0),
            'warnings': (purge_result or {}).get('warnings') if isinstance((purge_result or {}).get('warnings'), list) else [],
        }), 200
    finally:
        conn.close()


@app.route('/api/workspaces/<workspace_id>/documents', methods=['DELETE'])
def clear_workspace_documents(workspace_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        workspace_row = get_workspace_record(conn, workspace_id)
        if not workspace_row:
            return jsonify({'error': 'Workspace not found'}), 404
        if workspace_row.get('owner_username') != username:
            return jsonify({'error': 'Only workspace owner can clear workspace documents'}), 403

        docs_cursor = conn.execute(
            'SELECT id, filename FROM documents WHERE workspace_id = ?',
            (workspace_id,)
        )
        docs = [row_to_dict(item) for item in docs_cursor.fetchall()]
        if not docs:
            return jsonify({'deleted_count': 0, 'warnings': []}), 200

        doc_ids = [parse_int(item.get('id'), 0, 0) for item in docs]
        doc_ids = [item for item in doc_ids if item > 0]
        if doc_ids:
            placeholders = ','.join(['?'] * len(doc_ids))
            conn.execute(
                f'DELETE FROM document_share_links WHERE document_id IN ({placeholders})',
                tuple(doc_ids)
            )
            conn.execute(
                f'DELETE FROM document_summary_cache WHERE document_id IN ({placeholders})',
                tuple(doc_ids)
            )

        conn.execute('DELETE FROM documents WHERE workspace_id = ?', (workspace_id,))
        conn.commit()
    finally:
        conn.close()

    warnings = []
    for doc in docs:
        filename = str(doc.get('filename') or '').strip()
        if not filename:
            continue
        try:
            if S3_BUCKET and s3_client:
                s3_client.delete_object(Bucket=S3_BUCKET, Key=filename)
            else:
                local_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                if os.path.exists(local_path):
                    os.remove(local_path)
        except Exception as e:
            warnings.append(f'{filename}: {e}')

    return jsonify({
        'workspace_id': workspace_id,
        'deleted_count': len(docs),
        'warnings': warnings,
    }), 200

# ================= 修改后的上传接口 (支持 S3) =================
@app.route('/api/documents/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    username = (request.form.get('username') or 'Anonymous').strip()
    requested_workspace_id = (request.form.get('workspace_id') or '').strip()
    requested_category = normalize_document_category(request.form.get('category', ''))

    if requested_workspace_id and username:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        try:
            if not workspace_belongs_to_user(conn, requested_workspace_id, username):
                return jsonify({'error': 'No access to this workspace'}), 403
        finally:
            conn.close()
    
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        original_filename = file.filename
        try:
            ext = original_filename.rsplit('.', 1)[1].lower()
        except IndexError:
            return jsonify({'error': 'Filename must have an extension'}), 400
        
        unique_filename = f"{uuid.uuid4().hex}.{ext}"
        # 1. 先保存到本地临时文件夹 (为了提取文字)
        local_filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(local_filepath)

        # 2. 提取内容 (这一步需要本地文件)
        extracted_text, extracted_html = extract_document_content(local_filepath, ext)

        # 3. 【关键步骤】上传到 AWS S3
        try:
            if S3_BUCKET and s3_client: # 只有配置了 S3 且客户端初始化成功才上传
                print(f"🚀 Uploading to S3: {S3_BUCKET}")
                s3_client.upload_file(
                    local_filepath, 
                    S3_BUCKET, 
                    unique_filename,
                    ExtraArgs={'ContentType': file.content_type} # 设置文件类型
                )
                print("✅ Upload to S3 successful")
                
                # 4. 上传成功后，删除本地文件 (节省 Render 空间)
                os.remove(local_filepath)
                print("🗑️ Local file removed")
            else:
                print("⚠️ S3_BUCKET not set or client failed, keeping local file")

        except Exception as e:
            print(f"❌ S3 Upload Error: {e}")
            # 注意：即使 S3 上传失败，我们可能还是想保留数据库记录（或者报错，取决于你的需求）
            return jsonify({'error': f'Failed to upload to S3: {str(e)}'}), 500

        # 5. 存入数据库
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        try:
            workspace_id = ''
            workspace_settings = dict(DEFAULT_WORKSPACE_SETTINGS)
            if username:
                if requested_workspace_id:
                    if not workspace_belongs_to_user(conn, requested_workspace_id, username):
                        return jsonify({'error': 'No access to this workspace'}), 403
                    workspace_id = requested_workspace_id
                else:
                    workspace_id = get_or_create_default_workspace_id(conn, username)

                workspace_row = get_workspace_record(conn, workspace_id)
                workspace_settings = normalize_workspace_settings(
                    (workspace_row or {}).get('settings_json')
                )
                if not workspace_settings.get('allow_uploads', True):
                    return jsonify({'error': 'Uploads are disabled in this workspace settings'}), 403

            if requested_category:
                final_category = requested_category
            elif workspace_settings.get('auto_categorize', True):
                final_category = infer_document_category(original_filename, extracted_text)
            else:
                final_category = normalize_document_category(workspace_settings.get('default_category'))
                if not final_category:
                    final_category = DEFAULT_DOCUMENT_CATEGORY

            conn.execute(
                '''
                INSERT INTO documents (
                    filename, title, uploaded_at, file_type, content, content_html, username, tags, category, workspace_id
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    unique_filename,
                    original_filename,
                    datetime.utcnow().isoformat(),
                    ext,
                    extracted_text,
                    extracted_html if ext in ('docx', 'txt') else '',
                    username,
                    '',
                    final_category,
                    workspace_id
                )
            )
            conn.commit()
            return jsonify({'message': 'File uploaded successfully'}), 201
        except Exception as e:
            print(f"Database Error: {e}")
            return jsonify({'error': 'Database save failed'}), 500
        finally:
            conn.close()
    
    return jsonify({'error': 'File type not allowed'}), 400


@app.route('/api/canvas/files', methods=['POST'])
def get_canvas_files():
    data = request.get_json(silent=True) or {}
    token = str(data.get('token') or '').strip()
    domain = normalize_canvas_domain(data.get('domain') or CANVAS_DEFAULT_DOMAIN)
    username = str(data.get('username') or '').strip()
    requested_workspace_id = str(data.get('workspace_id') or '').strip()

    if not token:
        return jsonify({'error': 'Missing Canvas token'}), 400
    if not domain:
        return jsonify({'error': 'Invalid Canvas domain'}), 400
    if not username:
        return jsonify({'error': 'username is required'}), 400

    workspace_id = requested_workspace_id
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        if workspace_id:
            if not workspace_belongs_to_user(conn, workspace_id, username):
                return jsonify({'error': 'No access to this workspace'}), 403
        else:
            workspace_id = get_or_create_default_workspace_id(conn, username)
    finally:
        conn.close()

    try:
        raw_files = list_canvas_user_files(
            domain,
            canvas_headers(token),
            max_pages=CANVAS_MAX_LIST_PAGES,
            per_page=100,
        )
        supported_files = []
        for raw in raw_files:
            item = build_canvas_file_item(raw)
            if item:
                supported_files.append(item)

        supported_files.sort(
            key=lambda item: str(item.get('updated_at') or ''),
            reverse=True,
        )
        return jsonify({
            'files': supported_files,
            'count': len(supported_files),
            'domain': domain,
            'workspace_id': workspace_id,
        }), 200
    except requests.exceptions.HTTPError as e:
        status = e.response.status_code if e.response is not None else 400
        if status in (401, 403):
            return jsonify({'error': 'Canvas token is invalid or expired'}), 401
        detail = ''
        try:
            detail = str(e.response.text or '')[:300] if e.response is not None else ''
        except Exception:
            detail = ''
        return jsonify({'error': f'Failed to fetch Canvas files ({status}). {detail}'.strip()}), 400
    except requests.exceptions.RequestException as e:
        return jsonify({'error': f'Failed to connect to Canvas: {e}'}), 400
    except Exception as e:
        return jsonify({'error': f'Failed to fetch Canvas files: {e}'}), 500


@app.route('/api/canvas/import', methods=['POST'])
def import_canvas_file():
    data = request.get_json(silent=True) or {}
    token = str(data.get('token') or '').strip()
    domain = normalize_canvas_domain(data.get('domain') or CANVAS_DEFAULT_DOMAIN)
    username = str(data.get('username') or '').strip()
    requested_workspace_id = str(data.get('workspace_id') or '').strip()
    requested_category = normalize_document_category(data.get('category') or '')
    file_id = parse_int(data.get('file_id'), 0, 1)

    if not token:
        return jsonify({'error': 'Missing Canvas token'}), 400
    if not domain:
        return jsonify({'error': 'Invalid Canvas domain'}), 400
    if not username:
        return jsonify({'error': 'username is required'}), 400
    if not file_id:
        return jsonify({'error': 'Missing or invalid file ID'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500

    workspace_id = ''
    workspace_settings = dict(DEFAULT_WORKSPACE_SETTINGS)
    try:
        if requested_workspace_id:
            if not workspace_belongs_to_user(conn, requested_workspace_id, username):
                return jsonify({'error': 'No access to this workspace'}), 403
            workspace_id = requested_workspace_id
        else:
            workspace_id = get_or_create_default_workspace_id(conn, username)

        workspace_row = get_workspace_record(conn, workspace_id)
        workspace_settings = normalize_workspace_settings((workspace_row or {}).get('settings_json'))
        if not workspace_settings.get('allow_uploads', True):
            return jsonify({'error': 'Uploads are disabled in this workspace settings'}), 403
    finally:
        conn.close()

    headers = canvas_headers(token)
    local_filepath = ''
    unique_filename = ''
    original_filename = ''
    ext = ''
    extracted_text = ''
    extracted_html = ''
    final_category = ''
    imported_ok = False
    s3_uploaded = False

    try:
        info_url = f'https://{domain}/api/v1/files/{file_id}'
        info_res = requests.get(info_url, headers=headers, timeout=20)
        info_res.raise_for_status()
        info_payload = info_res.json()
        file_info = info_payload if isinstance(info_payload, dict) else {}

        original_filename = str(
            file_info.get('filename')
            or file_info.get('display_name')
            or file_info.get('name')
            or ''
        ).strip()
        if not original_filename:
            return jsonify({'error': 'Canvas file has no valid filename'}), 400
        if not allowed_file(original_filename):
            return jsonify({'error': 'File type is not supported by this project'}), 400

        ext = original_filename.rsplit('.', 1)[1].lower().strip('.')
        unique_filename = f'{uuid.uuid4().hex}.{ext}'

        max_size = int(app.config.get('MAX_CONTENT_LENGTH') or 0)
        canvas_file_size = parse_int(file_info.get('size'), 0, 0)
        if max_size and canvas_file_size and canvas_file_size > max_size:
            return jsonify({
                'error': f'Canvas file is too large ({canvas_file_size} bytes). Max allowed is {max_size} bytes.'
            }), 400

        download_url = str(file_info.get('url') or file_info.get('download_url') or '').strip()
        if not download_url:
            return jsonify({'error': 'Canvas file has no download URL'}), 400

        file_res = requests.get(download_url, headers=headers, timeout=60, stream=True)
        file_res.raise_for_status()
        chunks = []
        total_bytes = 0
        for chunk in file_res.iter_content(chunk_size=262144):
            if not chunk:
                continue
            total_bytes += len(chunk)
            if max_size and total_bytes > max_size:
                return jsonify({
                    'error': f'Downloaded file exceeded size limit ({max_size} bytes).'
                }), 400
            chunks.append(chunk)

        file_bytes = b''.join(chunks)
        if not file_bytes:
            return jsonify({'error': 'Failed to download file from Canvas'}), 400

        local_filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        with open(local_filepath, 'wb') as f:
            f.write(file_bytes)

        extracted_text, extracted_html = extract_document_content(local_filepath, ext)

        if S3_BUCKET and s3_client:
            s3_client.upload_file(
                local_filepath,
                S3_BUCKET,
                unique_filename,
                ExtraArgs={'ContentType': detect_mimetype(original_filename, ext)}
            )
            s3_uploaded = True
            try:
                os.remove(local_filepath)
                local_filepath = ''
            except Exception:
                pass

        if requested_category:
            final_category = requested_category
        elif workspace_settings.get('auto_categorize', True):
            final_category = infer_document_category(original_filename, extracted_text)
        else:
            final_category = normalize_document_category(workspace_settings.get('default_category'))
            if not final_category:
                final_category = DEFAULT_DOCUMENT_CATEGORY

        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        try:
            if workspace_id and not workspace_belongs_to_user(conn, workspace_id, username):
                return jsonify({'error': 'No access to this workspace'}), 403

            conn.execute(
                '''
                INSERT INTO documents (
                    filename, title, uploaded_at, file_type, content, content_html, username, tags, category, workspace_id
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    unique_filename,
                    original_filename,
                    datetime.utcnow().isoformat(),
                    ext,
                    extracted_text,
                    extracted_html if ext in ('docx', 'txt') else '',
                    username,
                    'Canvas Import',
                    final_category,
                    workspace_id
                )
            )
            conn.commit()

            cursor = conn.execute(
                '''
                SELECT *
                FROM documents
                WHERE filename = ? AND username = ?
                ORDER BY id DESC
                LIMIT 1
                ''',
                (unique_filename, username)
            )
            doc_row = cursor.fetchone()
            doc_payload = row_to_dict(doc_row) or {}
            imported_ok = True
            return jsonify({
                'message': 'Successfully imported from Canvas',
                'document': doc_payload,
                'workspace_id': workspace_id,
            }), 201
        finally:
            conn.close()
    except requests.exceptions.HTTPError as e:
        status = e.response.status_code if e.response is not None else 400
        if status in (401, 403):
            return jsonify({'error': 'Canvas token is invalid or expired'}), 401
        return jsonify({'error': f'Canvas API request failed with status {status}'}), 400
    except requests.exceptions.RequestException as e:
        return jsonify({'error': f'Failed to connect to Canvas: {e}'}), 400
    except Exception as e:
        print(f"Canvas Import Error: {e}")
        return jsonify({'error': f'Import failed: {e}'}), 500
    finally:
        if not imported_ok and unique_filename:
            if s3_uploaded:
                remove_document_file_from_storage(unique_filename)
            elif local_filepath and os.path.exists(local_filepath):
                try:
                    os.remove(local_filepath)
                except Exception:
                    pass


@app.route('/api/documents/<int:doc_id>', methods=['GET'])
def get_document(doc_id):
    username = (request.args.get('username') or '').strip()
    share_token = (request.args.get('share_token') or '').strip()
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        
        if doc:
            if is_document_soft_deleted(doc):
                return jsonify({'error': 'Document is in Trash'}), 404
            allowed, reason = check_document_access(conn, doc, username, share_token)
            if not allowed:
                return jsonify({'error': reason}), 403

            conn.execute('UPDATE documents SET last_access_at = ? WHERE id = ?', 
                         (datetime.utcnow().isoformat(), doc_id))
            conn.commit()
            doc_data = dict(doc)
            workspace_id = str(doc_data.get('workspace_id') or '').strip()
            workspace_settings = get_workspace_settings(conn, workspace_id)
            doc_data['link_sharing_mode'] = get_document_link_sharing_mode(conn, doc)
            doc_data['can_manage_share_links'] = user_can_manage_document_share_links(conn, doc, username)
            doc_data['allow_ai_tools'] = parse_bool(workspace_settings.get('allow_ai_tools', True), True)
            doc_data['allow_ocr'] = parse_bool(workspace_settings.get('allow_ocr', True), True)
            doc_data['allow_export'] = parse_bool(workspace_settings.get('allow_export', True), True)
            doc_data['summary_length'] = str(
                workspace_settings.get('summary_length', DEFAULT_WORKSPACE_SETTINGS.get('summary_length', 'medium'))
                or 'medium'
            ).strip().lower()
            doc_data['keyword_limit'] = parse_int(
                workspace_settings.get('keyword_limit', DEFAULT_WORKSPACE_SETTINGS.get('keyword_limit', 5)),
                5,
                3,
                12
            )
            doc_data['default_share_expiry_days'] = parse_int(
                workspace_settings.get(
                    'default_share_expiry_days',
                    DEFAULT_WORKSPACE_SETTINGS.get('default_share_expiry_days', 7)
                ),
                7,
                1,
                30
            )
            ext = str(doc_data.get('file_type') or '').lower().strip('.')
            if ext in ('docx', 'txt') and not (doc_data.get('content_html') or '').strip():
                doc_data['content_html'] = plaintext_to_html(doc_data.get('content') or '')
            return jsonify(doc_data)
        else:
            return jsonify({'error': 'Document not found'}), 404
    finally:
        conn.close()


@app.route('/api/documents/<int:doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or request.args.get('username') or '').strip()
    permanent = parse_bool(data.get('permanent') or request.args.get('permanent'), False)
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.execute(
            'SELECT id, filename, username, deleted_at FROM documents WHERE id = ?',
            (doc_id,)
        )
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        owner = (doc.get('username') if hasattr(doc, 'get') else doc['username']) or ''
        if owner and username != owner:
            return jsonify({'error': 'You can only delete your own documents'}), 403

        was_deleted = is_document_soft_deleted(doc)
        if permanent:
            deleted = hard_delete_document_record(conn, doc_id)
            if not deleted:
                return jsonify({'error': 'Document not found'}), 404
            conn.commit()
        else:
            if was_deleted:
                return jsonify({
                    'message': 'Document is already in Trash',
                    'id': doc_id,
                    'moved_to_trash': True,
                    'already_deleted': True,
                }), 200
            now_iso = utcnow_iso()
            conn.execute(
                'UPDATE documents SET deleted_at = ?, last_access_at = ? WHERE id = ?',
                (now_iso, now_iso, doc_id)
            )
            conn.execute('DELETE FROM document_share_links WHERE document_id = ?', (doc_id,))
            conn.commit()
    finally:
        conn.close()

    cleanup_warning = ''
    if permanent:
        filename = (deleted or {}).get('filename', '')
        cleanup_warning = remove_document_file_from_storage(filename)

    response = {
        'id': doc_id,
        'message': 'Document deleted permanently' if permanent else 'Document moved to Trash',
        'moved_to_trash': not permanent,
        'permanent': permanent,
        'trash_retention_days': TRASH_RETENTION_DAYS,
    }
    if cleanup_warning:
        response['warning'] = cleanup_warning
    return jsonify(response), 200


@app.route('/api/documents/<int:doc_id>/restore', methods=['POST'])
def restore_document(doc_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or request.args.get('username') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404
        doc_data = row_to_dict(doc) or {}
        owner = str(doc_data.get('username') or '').strip()
        if owner and owner != username:
            return jsonify({'error': 'You can only restore your own documents'}), 403
        if not is_document_soft_deleted(doc_data):
            return jsonify({
                'message': 'Document is already active',
                'id': doc_id,
                'restored': False,
            }), 200

        conn.execute(
            "UPDATE documents SET deleted_at = '' WHERE id = ?",
            (doc_id,)
        )
        conn.commit()

        refreshed_cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        refreshed = row_to_dict(refreshed_cursor.fetchone()) or {}
        return jsonify({
            'message': 'Document restored successfully',
            'id': doc_id,
            'restored': True,
            'document': refreshed,
        }), 200
    finally:
        conn.close()


@app.route('/api/documents/<int:doc_id>/share-links', methods=['POST'])
def create_document_share_link(doc_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500

    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        doc_data = row_to_dict(doc) or {}
        workspace_id = str(doc_data.get('workspace_id') or '').strip()

        if workspace_id:
            workspace_settings = get_workspace_settings(conn, workspace_id)
        else:
            workspace_settings = dict(DEFAULT_WORKSPACE_SETTINGS)

        if not user_can_manage_document_share_links(conn, doc, username):
            return jsonify({'error': 'Only owner (or allowed members) can create share links'}), 403

        link_mode = workspace_settings.get('link_sharing_mode', DEFAULT_WORKSPACE_SETTINGS['link_sharing_mode'])
        if link_mode == 'restricted':
            return jsonify({'error': 'Link sharing is restricted in this workspace settings'}), 403

        requested_expiry = data.get('expiry_days', None)
        if requested_expiry is None or str(requested_expiry).strip() == '':
            expiry_days = workspace_settings.get('default_share_expiry_days', 7)
        else:
            expiry_days = requested_expiry
        expiry_days = parse_int(expiry_days, 7, 1, 30)
        expires_at = expires_at_for_days(expiry_days)

        max_active_share_links = parse_int(
            workspace_settings.get('max_active_share_links_per_document', 5),
            5,
            1,
            20
        )
        auto_revoke_previous = parse_bool(
            workspace_settings.get('auto_revoke_previous_share_links', False),
            False
        )

        active_count = count_active_document_share_links(conn, doc_id)
        revoked_before_create = 0
        if auto_revoke_previous and active_count > 0:
            revoked_before_create = active_count
            conn.execute(
                '''
                UPDATE document_share_links
                SET status = 'revoked'
                WHERE document_id = ? AND status = 'active'
                ''',
                (doc_id,)
            )
            conn.commit()
            active_count = 0

        if active_count >= max_active_share_links:
            return jsonify({
                'error': (
                    f'Active share links reached limit ({max_active_share_links}). '
                    'Revoke existing links or enable auto-revoke in workspace settings.'
                ),
                'active_count': active_count,
                'max_active_share_links_per_document': max_active_share_links,
            }), 409

        token = create_document_share_token()
        try:
            conn.execute(
                '''
                INSERT INTO document_share_links (
                    document_id, workspace_id, token, created_by, status, expires_at, created_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    doc_id,
                    workspace_id,
                    token,
                    username,
                    'active',
                    expires_at,
                    utcnow_iso(),
                )
            )
        except Exception:
            return jsonify({'error': 'Failed to generate share token'}), 500

        conn.commit()
        share_cursor = conn.execute(
            'SELECT * FROM document_share_links WHERE token = ? LIMIT 1',
            (token,)
        )
        share_row = row_to_dict(share_cursor.fetchone())
        payload = serialize_document_share_link_row(share_row)
        payload['expiry_days'] = expiry_days
        payload['link_sharing_mode'] = link_mode
        payload['max_active_share_links_per_document'] = max_active_share_links
        payload['auto_revoke_previous_share_links'] = auto_revoke_previous
        payload['revoked_before_create'] = revoked_before_create
        return jsonify(payload), 201
    finally:
        conn.close()


@app.route('/api/documents/<int:doc_id>/share-links', methods=['GET'])
def list_document_share_links(doc_id):
    username = (request.args.get('username') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        if not user_can_manage_document_share_links(conn, doc, username):
            return jsonify({'error': 'Only owner (or allowed members) can manage share links'}), 403

        doc_data = row_to_dict(doc) or {}
        workspace_id = str(doc_data.get('workspace_id') or '').strip()
        link_mode = get_document_link_sharing_mode(conn, doc)
        items = list_document_share_link_payloads(conn, doc_id, limit=30)
        return jsonify({
            'document_id': doc_id,
            'workspace_id': workspace_id,
            'link_sharing_mode': link_mode,
            'items': items,
        }), 200
    finally:
        conn.close()


@app.route('/api/documents/<int:doc_id>/share-links', methods=['DELETE'])
def revoke_all_document_share_links(doc_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or request.args.get('username') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        if not user_can_manage_document_share_links(conn, doc, username):
            return jsonify({'error': 'Only owner (or allowed members) can manage share links'}), 403

        count_cursor = conn.execute(
            '''
            SELECT COUNT(1) AS total
            FROM document_share_links
            WHERE document_id = ? AND status != 'revoked'
            ''',
            (doc_id,)
        )
        count_row = row_to_dict(count_cursor.fetchone()) or {}
        revoke_count = parse_int(count_row.get('total', 0), 0, 0)

        conn.execute(
            '''
            UPDATE document_share_links
            SET status = 'revoked'
            WHERE document_id = ? AND status != 'revoked'
            ''',
            (doc_id,)
        )
        conn.commit()

        items = list_document_share_link_payloads(conn, doc_id, limit=30)
        return jsonify({
            'message': 'All share links revoked',
            'document_id': doc_id,
            'revoked_count': revoke_count,
            'items': items,
        }), 200
    finally:
        conn.close()


@app.route('/api/documents/<int:doc_id>/share-links/<int:share_link_id>', methods=['DELETE'])
def revoke_document_share_link(doc_id, share_link_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or request.args.get('username') or '').strip()
    if not username:
        return jsonify({'error': 'username is required'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        if not user_can_manage_document_share_links(conn, doc, username):
            return jsonify({'error': 'Only owner (or allowed members) can manage share links'}), 403

        link_cursor = conn.execute(
            'SELECT * FROM document_share_links WHERE id = ? AND document_id = ?',
            (share_link_id, doc_id)
        )
        link_row = row_to_dict(link_cursor.fetchone())
        if not link_row:
            return jsonify({'error': 'Share link not found'}), 404

        current_status = str(link_row.get('status') or '').strip().lower()
        if current_status == 'revoked':
            payload = to_document_share_link_payload(link_row)
            payload['message'] = 'Share link already revoked'
            return jsonify(payload), 200

        conn.execute(
            "UPDATE document_share_links SET status = 'revoked' WHERE id = ?",
            (share_link_id,)
        )
        conn.commit()
        refreshed_cursor = conn.execute(
            'SELECT * FROM document_share_links WHERE id = ? LIMIT 1',
            (share_link_id,)
        )
        refreshed = row_to_dict(refreshed_cursor.fetchone()) or link_row
        payload = to_document_share_link_payload(refreshed)
        payload['message'] = 'Share link revoked'
        return jsonify(payload), 200
    finally:
        conn.close()


@app.route('/api/share-links/<token>', methods=['GET'])
def get_document_by_share_token(token):
    safe_token = str(token or '').strip()
    username = (request.args.get('username') or '').strip()
    if not safe_token:
        return jsonify({'error': 'Missing share token'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500

    try:
        share_cursor = conn.execute(
            '''
            SELECT *
            FROM document_share_links
            WHERE token = ?
            ORDER BY id DESC
            LIMIT 1
            ''',
            (safe_token,)
        )
        share_row = row_to_dict(share_cursor.fetchone())
        if not share_row:
            return jsonify({'error': 'Share link not found'}), 404

        doc_id = parse_int(share_row.get('document_id'), 0, 0)
        if doc_id <= 0:
            return jsonify({'error': 'Share link is invalid'}), 404

        token_ok, validated_share_row, token_reason = validate_document_share_token(
            conn,
            doc_id,
            safe_token,
            mark_access=False
        )
        if not token_ok:
            return jsonify({'error': token_reason or 'Share link is invalid'}), 403
        if validated_share_row:
            share_row = validated_share_row

        doc_cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = doc_cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        allowed, reason = check_document_access(conn, doc, username, safe_token)
        if not allowed:
            return jsonify({'error': reason}), 403

        refreshed_share_cursor = conn.execute(
            'SELECT * FROM document_share_links WHERE token = ? ORDER BY id DESC LIMIT 1',
            (safe_token,)
        )
        refreshed_share_row = row_to_dict(refreshed_share_cursor.fetchone()) or share_row

        conn.execute(
            'UPDATE documents SET last_access_at = ? WHERE id = ?',
            (utcnow_iso(), doc_id)
        )
        conn.commit()

        doc_data = dict(doc)
        workspace_id = str(doc_data.get('workspace_id') or '').strip()
        workspace_settings = get_workspace_settings(conn, workspace_id)
        doc_data['link_sharing_mode'] = get_document_link_sharing_mode(conn, doc)
        doc_data['can_manage_share_links'] = user_can_manage_document_share_links(conn, doc, username)
        doc_data['share'] = serialize_document_share_link_row(refreshed_share_row)
        doc_data['allow_ai_tools'] = parse_bool(workspace_settings.get('allow_ai_tools', True), True)
        doc_data['allow_ocr'] = parse_bool(workspace_settings.get('allow_ocr', True), True)
        doc_data['allow_export'] = parse_bool(workspace_settings.get('allow_export', True), True)
        doc_data['summary_length'] = str(
            workspace_settings.get('summary_length', DEFAULT_WORKSPACE_SETTINGS.get('summary_length', 'medium'))
            or 'medium'
        ).strip().lower()
        doc_data['keyword_limit'] = parse_int(
            workspace_settings.get('keyword_limit', DEFAULT_WORKSPACE_SETTINGS.get('keyword_limit', 5)),
            5,
            3,
            12
        )
        doc_data['default_share_expiry_days'] = parse_int(
            workspace_settings.get(
                'default_share_expiry_days',
                DEFAULT_WORKSPACE_SETTINGS.get('default_share_expiry_days', 7)
            ),
            7,
            1,
            30
        )
        ext = str(doc_data.get('file_type') or '').lower().strip('.')
        if ext in ('docx', 'txt') and not (doc_data.get('content_html') or '').strip():
            doc_data['content_html'] = plaintext_to_html(doc_data.get('content') or '')
        return jsonify(doc_data), 200
    finally:
        conn.close()


@app.route('/api/documents/<int:doc_id>/tags', methods=['PUT'])
def update_document_tags(doc_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or request.args.get('username') or '').strip()
    raw_tags = data.get('tags', [])

    if isinstance(raw_tags, list):
        tags_list = [str(tag).strip() for tag in raw_tags if str(tag).strip()]
    elif isinstance(raw_tags, str):
        tags_list = [tag.strip() for tag in raw_tags.split(',') if tag.strip()]
    else:
        return jsonify({'error': 'tags must be a list or comma-separated string'}), 400

    tags_value = ','.join(tags_list)

    conn = get_db_connection()
    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        allowed, reason = check_document_access(conn, doc, username)
        if not allowed:
            return jsonify({'error': reason}), 403

        if not user_can_edit_document(conn, doc, username):
            return jsonify({'error': 'Only workspace members can edit this document'}), 403

        workspace_id = str((doc.get('workspace_id') if hasattr(doc, 'get') else doc['workspace_id']) or '').strip()
        workspace_settings = get_workspace_settings(conn, workspace_id)
        if not workspace_settings.get('allow_note_editing', True):
            return jsonify({'error': 'Editing is disabled in this workspace settings'}), 403

        conn.execute('UPDATE documents SET tags = ? WHERE id = ?', (tags_value, doc_id))
        conn.commit()

        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        return jsonify(dict(doc)), 200
    finally:
        conn.close()


@app.route('/api/documents/<int:doc_id>/category', methods=['PUT'])
def update_document_category(doc_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or request.args.get('username') or '').strip()
    next_category = normalize_document_category(data.get('category', ''))
    if not next_category:
        next_category = DEFAULT_DOCUMENT_CATEGORY

    conn = get_db_connection()
    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        allowed, reason = check_document_access(conn, doc, username)
        if not allowed:
            return jsonify({'error': reason}), 403

        if not user_can_edit_document(conn, doc, username):
            return jsonify({'error': 'Only workspace members can edit this document'}), 403

        workspace_id = str((doc.get('workspace_id') if hasattr(doc, 'get') else doc['workspace_id']) or '').strip()
        workspace_settings = get_workspace_settings(conn, workspace_id)
        if not workspace_settings.get('allow_note_editing', True):
            return jsonify({'error': 'Editing is disabled in this workspace settings'}), 403

        conn.execute('UPDATE documents SET category = ? WHERE id = ?', (next_category, doc_id))
        conn.commit()

        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        return jsonify(dict(doc)), 200
    finally:
        conn.close()


def extract_key_sentences(text_content, keywords=None, limit=3):
    normalized = normalize_newlines(text_content or '')
    fragments = [part.strip() for part in re.split(r'(?<=[.!?。！？])\s+', normalized) if part.strip()]
    if not fragments:
        return []

    keyword_list = [
        str(item).strip().lower()
        for item in (keywords or [])
        if str(item).strip()
    ]
    keyword_list = [item for item in keyword_list if item != 'not enough text']

    scored = []
    for idx, sentence in enumerate(fragments[:150]):
        lower_sentence = sentence.lower()
        score = 0
        for keyword in keyword_list:
            score += lower_sentence.count(keyword) * 2
        score += min(len(sentence.split()) / 8.0, 1.5)
        scored.append((score, -idx, sentence))

    scored.sort(reverse=True)
    top = []
    for _, _, sentence in scored:
        if sentence in top:
            continue
        top.append(sentence)
        if len(top) >= max(1, int(limit or 3)):
            break

    if not top:
        top = fragments[:max(1, int(limit or 3))]
    return top


@app.route('/api/documents/<int:doc_id>/content', methods=['PUT'])
def update_document_content(doc_id):
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or request.args.get('username') or '').strip()
    content = data.get('content', '')
    content_html = data.get('content_html')

    if content is None:
        content = ''
    if not isinstance(content, str):
        return jsonify({'error': 'content must be a string'}), 400
    if content_html is not None and not isinstance(content_html, str):
        return jsonify({'error': 'content_html must be a string'}), 400

    conn = get_db_connection()
    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        allowed, reason = check_document_access(conn, doc, username)
        if not allowed:
            return jsonify({'error': reason}), 403

        if not user_can_edit_document(conn, doc, username):
            return jsonify({'error': 'Only workspace members can edit this document'}), 403

        workspace_id = str((doc.get('workspace_id') if hasattr(doc, 'get') else doc['workspace_id']) or '').strip()
        workspace_settings = get_workspace_settings(conn, workspace_id)
        if not workspace_settings.get('allow_note_editing', True):
            return jsonify({'error': 'Editing is disabled in this workspace settings'}), 403

        file_type = (doc.get('file_type') if hasattr(doc, 'get') else doc['file_type']) or ''
        file_type = str(file_type).lower().strip('.')
        existing_html = (doc.get('content_html') if hasattr(doc, 'get') else doc['content_html']) or ''

        if file_type in ('docx', 'txt'):
            if content_html is None:
                content_html = plaintext_to_html(content)
            if not content_html.strip() and existing_html.strip():
                content_html = existing_html
            content_html = sanitize_editor_html(content_html)
            content = html_to_plaintext(content_html)
        else:
            content_html = ''

        try:
            file_bytes, mimetype = build_editable_file_bytes(file_type, content, content_html)
            filename = doc.get('filename') if hasattr(doc, 'get') else doc['filename']
            write_file_bytes_to_storage(filename, file_bytes, mimetype)
        except ValueError as e:
            return jsonify({'error': str(e)}), 400
        except RuntimeError as e:
            return jsonify({'error': str(e)}), 500
        except Exception as e:
            print(f"File update failed: {e}")
            return jsonify({'error': 'Failed to update source file'}), 500

        conn.execute(
            'UPDATE documents SET content = ?, content_html = ? WHERE id = ?',
            (content, content_html, doc_id)
        )
        conn.commit()

        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        updated_doc = cursor.fetchone()
        return jsonify(dict(updated_doc)), 200
    finally:
        conn.close()


@app.route('/api/documents/<int:doc_id>/pdf', methods=['PUT'])
def update_document_pdf_file(doc_id):
    username = (
        (request.args.get('username') or request.form.get('username') or '').strip()
    )
    if request.files and 'file' in request.files:
        file_bytes = request.files['file'].read()
    else:
        file_bytes = request.get_data(cache=False) or b''

    if not file_bytes:
        return jsonify({'error': 'No PDF data provided'}), 400

    if not file_bytes.lstrip().startswith(b'%PDF'):
        return jsonify({'error': 'Invalid PDF payload'}), 400

    conn = get_db_connection()
    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        allowed, reason = check_document_access(conn, doc, username)
        if not allowed:
            return jsonify({'error': reason}), 403

        if not user_can_edit_document(conn, doc, username):
            return jsonify({'error': 'Only workspace members can edit this document'}), 403

        workspace_id = str((doc.get('workspace_id') if hasattr(doc, 'get') else doc['workspace_id']) or '').strip()
        workspace_settings = get_workspace_settings(conn, workspace_id)
        if not workspace_settings.get('allow_note_editing', True):
            return jsonify({'error': 'Editing is disabled in this workspace settings'}), 403

        file_type = (doc.get('file_type') if hasattr(doc, 'get') else doc['file_type']) or ''
        if str(file_type).lower() != 'pdf':
            return jsonify({'error': 'This endpoint only supports PDF documents'}), 400

        filename = doc.get('filename') if hasattr(doc, 'get') else doc['filename']
        try:
            write_file_bytes_to_storage(filename, file_bytes, MIME_BY_EXT['pdf'])
        except Exception as e:
            print(f"PDF file update failed: {e}")
            return jsonify({'error': 'Failed to update source PDF file'}), 500

        extracted_text = extract_text_from_pdf_bytes(file_bytes)
        if not extracted_text.strip():
            extracted_text = (doc.get('content') if hasattr(doc, 'get') else doc['content']) or ''
        conn.execute(
            'UPDATE documents SET content = ?, content_html = ? WHERE id = ?',
            (extracted_text, '', doc_id)
        )
        conn.commit()

        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        updated_doc = cursor.fetchone()
        return jsonify(dict(updated_doc)), 200
    finally:
        conn.close()


@app.route('/api/documents/<int:doc_id>/file', methods=['GET'])
def get_document_file(doc_id):
    username = (request.args.get('username') or '').strip()
    share_token = (request.args.get('share_token') or '').strip()
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500
    try:
        cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc = cursor.fetchone()
        if not doc:
            return jsonify({'error': 'Document not found'}), 404

        allowed, reason = check_document_access(conn, doc, username, share_token)
        if not allowed:
            return jsonify({'error': reason}), 403

        doc_data = row_to_dict(doc) or {}
        filename = doc_data.get('filename', '')
        title = doc_data.get('title', '')
        file_ext = doc_data.get('file_type', '')
        mimetype = detect_mimetype(filename, file_ext)
    finally:
        conn.close()

    if S3_BUCKET and s3_client:
        try:
            s3_obj = s3_client.get_object(Bucket=S3_BUCKET, Key=filename)
            return send_file(
                io.BytesIO(s3_obj['Body'].read()),
                mimetype=mimetype,
                download_name=title or filename,
                as_attachment=False
            )
        except Exception as e:
            print(f"S3 stream error: {e}")
            return jsonify({'error': 'Could not read file from S3'}), 500

    local_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(local_path):
        return jsonify({'error': 'File not found'}), 404

    return send_file(
        local_path,
        mimetype=mimetype,
        download_name=title or filename,
        as_attachment=False
    )


def get_hf_headers(content_type=None):
    if not HF_TOKEN:
        return None
    result = {"Authorization": f"Bearer {HF_TOKEN}"}
    if content_type:
        result["Content-Type"] = content_type
    return result


def hf_model_url(model_id):
    return f"{HF_MODEL_BASE_URL}/{model_id}"


def hf_error_message(response):
    try:
        body = response.json()
        if isinstance(body, dict):
            if body.get('error'):
                return str(body['error'])
            if body.get('message'):
                return str(body['message'])
    except Exception:
        pass
    return (response.text or '').strip()[:240] or 'Unknown error'


def split_text_for_summary(text_content, max_chars=3600, min_chars=1200, overlap_chars=220):
    normalized = normalize_newlines(text_content or '')
    normalized = re.sub(r'[ \t]+', ' ', normalized).strip()
    if not normalized:
        return []
    if len(normalized) <= max_chars:
        return [normalized]

    markers = (
        ('\n\n', 2),
        ('. ', 2),
        ('! ', 2),
        ('? ', 2),
        ('。', 1),
        ('！', 1),
        ('？', 1),
        ('; ', 2),
        ('；', 1),
    )
    chunks = []
    start = 0
    total_len = len(normalized)
    guard = 0

    while start < total_len and guard < 10000:
        guard += 1
        hard_end = min(total_len, start + max_chars)
        end = hard_end

        if hard_end < total_len:
            window = normalized[start:hard_end]
            best_pos = -1
            best_tail = 0
            for marker, tail in markers:
                marker_pos = window.rfind(marker)
                if marker_pos > best_pos:
                    best_pos = marker_pos
                    best_tail = tail
            if best_pos >= min_chars:
                end = start + best_pos + best_tail

        if end <= start:
            end = hard_end

        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= total_len:
            break

        next_start = max(0, end - max(0, overlap_chars))
        if next_start <= start:
            next_start = end
        start = next_start

    if not chunks:
        return [normalized[:max_chars]]
    return chunks


def build_fallback_summary(text_content, sentence_limit=3, max_chars=560):
    raw_text = normalize_newlines(text_content or '')
    raw_text = re.sub(r'(?im)^\s*part\s+\d+\s*:\s*', '', raw_text)
    raw_text = re.sub(r'[ \t]+', ' ', raw_text).strip()
    if not raw_text:
        return ''

    safe_limit = max(1, int(sentence_limit or 3))
    fragments = [
        part.strip()
        for part in re.split(r'(?<=[.!?。！？])\s+|\n+', raw_text)
        if part.strip()
    ]

    if len(fragments) < safe_limit:
        compact_lines = [
            part.strip()
            for part in re.split(r'\n+', normalize_newlines(text_content or ''))
            if part.strip()
        ]
        for item in compact_lines:
            candidate = re.sub(r'\s+', ' ', item)
            if candidate and candidate not in fragments:
                fragments.append(candidate)

    if fragments:
        if len(fragments) <= safe_limit:
            picked = fragments
        else:
            picked = []
            key_indexes = [0, len(fragments) // 2, len(fragments) - 1]
            for idx in key_indexes:
                sentence = fragments[idx]
                if sentence not in picked:
                    picked.append(sentence)
                if len(picked) >= safe_limit:
                    break
            if len(picked) < safe_limit:
                for sentence in fragments:
                    if sentence in picked:
                        continue
                    picked.append(sentence)
                    if len(picked) >= safe_limit:
                        break
        summary = ' '.join(picked[:safe_limit]).strip()
    else:
        summary = raw_text

    if len(summary) > max_chars:
        text_len = len(summary)
        slice_len = max(90, max_chars // max(1, safe_limit))
        points = [0, text_len // 2, max(0, text_len - slice_len)]
        excerpts = []
        for point in points:
            start = max(0, min(point, max(0, text_len - slice_len)))
            end = min(text_len, start + slice_len)
            snippet = summary[start:end].strip()
            if snippet and snippet not in excerpts:
                excerpts.append(snippet)
        summary = ' ... '.join(excerpts).strip() or summary[:max_chars]
        if len(summary) > max_chars:
            clipped = summary[:max_chars]
            summary = clipped.rsplit(' ', 1)[0].strip() or clipped

    return summary


def call_hf_summarizer_once(text_content, length_options):
    safe_text = str(text_content or '').strip()
    if not safe_text:
        return {'ok': False, 'summary': '', 'error': 'Empty input text'}

    hf_headers = get_hf_headers('application/json')
    if not hf_headers:
        return {'ok': False, 'summary': '', 'error': 'HF_API_TOKEN is not configured on server.'}

    payload = {
        "inputs": safe_text,
        "parameters": {
            "max_new_tokens": length_options['max_new_tokens'],
            "min_new_tokens": length_options['min_new_tokens'],
            "do_sample": False
        },
        "options": {"wait_for_model": True}
    }

    try:
        response = requests.post(
            hf_model_url(SUMMARIZER_MODEL_ID),
            headers=hf_headers,
            json=payload,
            timeout=90
        )
    except Exception as e:
        return {'ok': False, 'summary': '', 'error': str(e) or 'AI service busy.'}

    if response.status_code >= 400:
        return {
            'ok': False,
            'summary': '',
            'error': f"Summary service failed ({response.status_code}): {hf_error_message(response)}"
        }

    summary = ''
    try:
        summary_res = response.json()
        if isinstance(summary_res, list) and summary_res and isinstance(summary_res[0], dict):
            summary = str(
                summary_res[0].get('summary_text')
                or summary_res[0].get('generated_text')
                or ''
            ).strip()
        elif isinstance(summary_res, dict):
            summary = str(
                summary_res.get('summary_text')
                or summary_res.get('generated_text')
                or ''
            ).strip()
    except Exception:
        summary = ''

    if summary:
        return {'ok': True, 'summary': summary, 'error': ''}

    return {'ok': False, 'summary': '', 'error': 'Summary service returned empty output.'}


def summarize_text_with_chunk_merge(text_content, length_options):
    safe_text = str(text_content or '').strip()
    sentence_limit = max(1, int((length_options or {}).get('sentence_limit', 3) or 3))
    hf_available = bool(get_hf_headers('application/json'))
    chunks = split_text_for_summary(
        safe_text,
        max_chars=3600,
        min_chars=1200,
        overlap_chars=220
    )

    if not chunks:
        return {
            'summary': '',
            'summary_source': 'fallback',
            'summary_note': 'No text provided.',
            'meta': {'chunk_count': 0, 'merge_rounds': 0}
        }

    hf_success_count = 0
    fallback_count = 0
    error_samples = []

    def summarize_unit(unit_text):
        nonlocal hf_success_count, fallback_count, error_samples
        result = call_hf_summarizer_once(unit_text, length_options)
        if result.get('ok'):
            hf_success_count += 1
            return str(result.get('summary') or '').strip()

        fallback_count += 1
        err = str(result.get('error') or '').strip()
        if err and err not in error_samples and len(error_samples) < 2:
            error_samples.append(err)
        return build_fallback_summary(
            unit_text,
            sentence_limit=max(2, sentence_limit),
            max_chars=560
        )

    layer = [summarize_unit(chunk) for chunk in chunks]
    merge_rounds = 0
    if not hf_available and len(layer) > 1:
        merge_rounds = 1
        pick_idx = [0, len(layer) // 2, len(layer) - 1]
        picked = []
        for idx in pick_idx:
            unit = str(layer[idx] or '').strip()
            if unit and unit not in picked:
                picked.append(unit)
        combined = '\n'.join(picked).strip()
        layer = [build_fallback_summary(combined or safe_text, sentence_limit=max(3, sentence_limit), max_chars=780)]
    else:
        while len(layer) > 1 and merge_rounds < 5:
            merge_rounds += 1
            combined_text = '\n\n'.join(
                part
                for part in layer
                if str(part).strip()
            ).strip()
            if not combined_text:
                break
            merge_chunks = split_text_for_summary(
                combined_text,
                max_chars=3400,
                min_chars=900,
                overlap_chars=120
            )
            if not merge_chunks:
                break
            layer = [summarize_unit(chunk) for chunk in merge_chunks]

    summary = str(layer[0] if layer else '').strip()
    if not summary:
        summary = build_fallback_summary(
            safe_text,
            sentence_limit=max(3, sentence_limit),
            max_chars=560
        )
        fallback_count += 1

    summary_source = 'huggingface' if hf_success_count > 0 else 'fallback'
    note_parts = []
    if len(chunks) > 1:
        note_parts.append(f"Chunked summary used {len(chunks)} sections")
    if merge_rounds > 0:
        note_parts.append(f"{merge_rounds} merge round(s)")
    if fallback_count > 0:
        note_parts.append(f"{fallback_count} section(s) used fallback")
    if error_samples:
        note_parts.append("HF note: " + ' | '.join(error_samples))
    summary_note = '; '.join(note_parts)

    return {
        'summary': summary,
        'summary_source': summary_source,
        'summary_note': summary_note,
        'meta': {
            'chunk_count': len(chunks),
            'merge_rounds': merge_rounds,
            'hf_success_count': hf_success_count,
            'fallback_count': fallback_count,
        }
    }


def normalize_ocr_text(payload, depth=0):
    if depth > 5 or payload is None:
        return ''
    if isinstance(payload, str):
        return payload.strip()
    if isinstance(payload, (int, float, bool)):
        return str(payload).strip()
    if isinstance(payload, bytes):
        try:
            return payload.decode('utf-8', errors='ignore').strip()
        except Exception:
            return ''

    if isinstance(payload, list):
        parts = []
        for item in payload:
            text = normalize_ocr_text(item, depth + 1)
            if text:
                parts.append(text)
        return '\n'.join(parts).strip() if parts else ''

    if isinstance(payload, dict):
        preferred_keys = (
            'text',
            'ocr_text',
            'extracted_text',
            'result',
            'content',
            'generated_text',
            'output_text',
            'prediction',
            'predictions',
            'value',
            'data',
            'lines',
            'texts',
        )
        for key in preferred_keys:
            if key not in payload:
                continue
            text = normalize_ocr_text(payload.get(key), depth + 1)
            if text:
                return text

        choices = payload.get('choices')
        if isinstance(choices, list):
            parts = []
            for choice in choices:
                if not isinstance(choice, dict):
                    continue
                message = choice.get('message')
                if isinstance(message, dict):
                    text = normalize_ocr_text(message.get('content'), depth + 1)
                else:
                    text = normalize_ocr_text(choice.get('text'), depth + 1)
                if text:
                    parts.append(text)
            if parts:
                return '\n'.join(parts).strip()

    return ''


def get_local_ocr_engine():
    global _rapid_ocr_engine
    if RapidOCR is None or cv2 is None or np is None:
        return None
    if _rapid_ocr_engine is None:
        _rapid_ocr_engine = RapidOCR()
    return _rapid_ocr_engine


def run_local_ocr(img_bytes):
    try:
        engine = get_local_ocr_engine()
    except Exception as e:
        return '', f'RapidOCR engine init failed: {e}'

    if engine is None:
        reasons = []
        if cv2 is None:
            reasons.append(f'cv2 unavailable: {CV2_IMPORT_ERROR or "not installed"}')
        if np is None:
            reasons.append('numpy unavailable')
        if RapidOCR is None:
            reasons.append(f'rapidocr unavailable: {RAPIDOCR_IMPORT_ERROR or "not installed"}')
        reason_text = '; '.join(reasons) if reasons else 'unknown reason'
        return '', f'RapidOCR dependencies are missing ({reason_text})'

    try:
        image_arr = np.frombuffer(img_bytes, dtype=np.uint8)
        image = cv2.imdecode(image_arr, cv2.IMREAD_COLOR)
        if image is None:
            return '', 'Failed to decode image for local OCR'

        result, _ = engine(image)
        if not result:
            return '', 'Local OCR returned empty text'

        lines = []
        for item in result:
            if not isinstance(item, (list, tuple)) or len(item) < 2:
                continue
            line = str(item[1]).strip()
            if line:
                lines.append(line)

        text = '\n'.join(lines).strip()
        if not text:
            return '', 'Local OCR returned empty text'

        return text, ''
    except Exception as e:
        return '', f'Local OCR error: {e}'


def get_ocr_runtime_status():
    ocrmypdf_path = shutil.which(OCRMYPDF_BINARY)
    status = {
        'external_ocr_configured': bool(EXTERNAL_OCR_SERVICE_URL),
        'external_ocr_url': EXTERNAL_OCR_SERVICE_URL or '',
        'hf_token_configured': bool(HF_TOKEN),
        'hf_ocr_model': OCR_MODEL_ID,
        'hf_model_base_url': HF_MODEL_BASE_URL,
        'pdf_ocr_fallback_enabled': ENABLE_PDF_OCR_FALLBACK,
        'ocrmypdf_binary': OCRMYPDF_BINARY,
        'ocrmypdf_available': bool(ocrmypdf_path),
        'ocrmypdf_path': ocrmypdf_path or '',
        'ocrmypdf_language': OCRMYPDF_LANGUAGE,
        'cv2_available': cv2 is not None,
        'numpy_available': np is not None,
        'rapidocr_available': RapidOCR is not None,
        'cv2_import_error': CV2_IMPORT_ERROR,
        'rapidocr_import_error': RAPIDOCR_IMPORT_ERROR,
        'local_engine_ready': False,
        'local_engine_error': '',
        'hints': [],
    }

    if status['rapidocr_available'] and status['cv2_available'] and status['numpy_available']:
        try:
            get_local_ocr_engine()
            status['local_engine_ready'] = True
        except Exception as e:
            status['local_engine_error'] = str(e)
    else:
        missing = []
        if not status['rapidocr_available']:
            missing.append('rapidocr_onnxruntime')
        if not status['cv2_available']:
            missing.append('opencv-python-headless')
        if not status['numpy_available']:
            missing.append('numpy')
        if missing:
            status['local_engine_error'] = f"Missing local OCR dependencies: {', '.join(missing)}"

    if not status['hf_token_configured']:
        status['hints'].append('Set HF_API_TOKEN in environment variables to enable Hugging Face OCR fallback.')
    if not status['external_ocr_configured']:
        status['hints'].append('Set EXTERNAL_OCR_SERVICE_URL to enable external OCR service routing.')
    if ENABLE_PDF_OCR_FALLBACK and not status['ocrmypdf_available']:
        status['hints'].append('Install ocrmypdf binary to enable automatic PDF OCR fallback for low-quality text extraction.')

    local_error_lower = str(status['local_engine_error'] or '').lower()
    import_error_lower = str(status['rapidocr_import_error'] or '').lower()
    cv2_error_lower = str(status['cv2_import_error'] or '').lower()
    joined_errors = ' | '.join([local_error_lower, import_error_lower, cv2_error_lower])

    if 'libgomp.so.1' in joined_errors:
        status['hints'].append('Install system package libgomp1 in Docker image for onnxruntime.')
    if 'libgl.so.1' in joined_errors:
        status['hints'].append('Install libgl1-mesa-glx or libgl1 in Docker image.')
    if 'libglib-2.0.so.0' in joined_errors:
        status['hints'].append('Install libglib2.0-0 in Docker image.')

    return status


@app.route('/api/ocr/health', methods=['GET'])
def ocr_health():
    status = get_ocr_runtime_status()
    ok = bool(
        status.get('local_engine_ready')
        or status.get('hf_token_configured')
        or status.get('external_ocr_configured')
    )
    status['ok'] = ok
    status['checked_at'] = utcnow_iso()
    if not ok:
        status['hints'].append(
            'No OCR provider is ready. Configure EXTERNAL_OCR_SERVICE_URL or HF_API_TOKEN, or enable local RapidOCR.'
        )
    return jsonify(status), (200 if ok else 503)

# ==========================================
# 专家 1 号：视觉专家 (负责看图识字)
# 对应前端的【按钮 1】
# ==========================================
@app.route('/api/extract-text', methods=['POST'])
@app.route('/api/extract-text/<int:doc_id>', methods=['POST'])
def extract_text_from_image(doc_id=None):
    username = (request.values.get('username') or '').strip()
    share_token = (request.values.get('share_token') or '').strip()
    requested_workspace_id = (request.values.get('workspace_id') or '').strip()
    img_bytes = b''
    mimetype = 'application/octet-stream'
    source_filename = 'image.jpg'

    if doc_id is not None:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        doc = None
        workspace_settings = dict(DEFAULT_WORKSPACE_SETTINGS)
        try:
            cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
            doc = cursor.fetchone()
            if doc:
                allowed, reason = check_document_access(conn, doc, username, share_token)
                if not allowed:
                    return jsonify({"error": reason}), 403

                workspace_id = str(
                    (doc.get('workspace_id') if hasattr(doc, 'get') else doc['workspace_id']) or ''
                ).strip()
                workspace_settings = get_workspace_settings(conn, workspace_id)
        finally:
            conn.close()

        if not doc:
            return jsonify({"error": "Document not found"}), 404
        if not workspace_settings.get('allow_ai_tools', True):
            return jsonify({"error": "AI tools are disabled in this workspace settings"}), 403
        if not workspace_settings.get('allow_ocr', True):
            return jsonify({"error": "OCR is disabled in this workspace settings"}), 403

        filename = doc.get('filename') if hasattr(doc, 'get') else doc['filename']
        file_type = doc.get('file_type') if hasattr(doc, 'get') else doc['file_type']
        source_filename = str(filename or source_filename)
        if str(file_type or '').lower() not in ('png', 'jpg', 'jpeg', 'webp', 'gif'):
            return jsonify({"error": "This endpoint only supports image documents"}), 400
        mimetype = detect_mimetype(filename, file_type)

        try:
            if S3_BUCKET and s3_client:
                s3_obj = s3_client.get_object(Bucket=S3_BUCKET, Key=filename)
                img_bytes = s3_obj['Body'].read()
            else:
                local_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                if not os.path.exists(local_path):
                    return jsonify({"error": "Source image not found"}), 404
                with open(local_path, 'rb') as f:
                    img_bytes = f.read()
        except Exception as e:
            return jsonify({"error": f"Failed to read source image: {e}"}), 500
    else:
        if username:
            conn = get_db_connection()
            if not conn:
                return jsonify({'error': 'Database connection failed'}), 500
            try:
                workspace_id = requested_workspace_id
                if not workspace_id:
                    default_cursor = conn.execute(
                        '''
                        SELECT id
                        FROM workspaces
                        WHERE owner_username = ?
                        ORDER BY created_at ASC, id ASC
                        LIMIT 1
                        ''',
                        (username,)
                    )
                    default_row = row_to_dict(default_cursor.fetchone())
                    workspace_id = str(default_row.get('id') or '').strip()
                if workspace_id and not workspace_belongs_to_user(conn, workspace_id, username):
                    return jsonify({'error': 'No access to this workspace'}), 403
                workspace_settings = get_workspace_settings(conn, workspace_id)
            finally:
                conn.close()

            if not workspace_settings.get('allow_ai_tools', True):
                return jsonify({"error": "AI tools are disabled in this workspace settings"}), 403
            if not workspace_settings.get('allow_ocr', True):
                return jsonify({"error": "OCR is disabled in this workspace settings"}), 403

        if 'image' not in request.files:
            return jsonify({"error": "No image provided"}), 400
        file = request.files['image']
        mimetype = file.mimetype or 'application/octet-stream'
        source_filename = str(file.filename or source_filename)
        img_bytes = file.read()

    if not img_bytes:
        return jsonify({"error": "Empty image file"}), 400

    external_error = ''
    # --- 优先尝试外部算力中心（仅在配置环境变量后启用） ---
    if EXTERNAL_OCR_SERVICE_URL:
        print(f"正在尝试调用算力中心: {EXTERNAL_OCR_SERVICE_URL}")
        try:
            files = {'file': (source_filename if doc_id else 'image.jpg', img_bytes, mimetype)}
            response = requests.post(EXTERNAL_OCR_SERVICE_URL, files=files, timeout=30)
            if response.status_code == 200:
                result = None
                extracted_external_text = ''
                try:
                    result = response.json()
                    extracted_external_text = normalize_ocr_text(result)
                except Exception:
                    extracted_external_text = normalize_ocr_text(response.text or '')

                if extracted_external_text:
                    print("算力中心识别成功。")
                    return jsonify({"text": extracted_external_text, "source": "external_ocr_service"})

                if isinstance(result, dict):
                    keys_preview = ', '.join([str(k) for k in list(result.keys())[:8]])
                    external_error = f"External OCR returned empty text (response keys: {keys_preview or 'none'})"
                elif isinstance(result, list):
                    external_error = f"External OCR returned empty text (response type: list[{len(result)}])"
                else:
                    external_error = "External OCR returned empty text"
                print(external_error)
            else:
                external_error = f"External OCR failed ({response.status_code}): {response.text[:220]}"
                print(f"外部算力中心返回错误码: {response.status_code}")
        except Exception as e:
            external_error = f"External OCR error: {e}"
            print(f"外部算力中心未响应，切换备用方案。错误: {e}")
    # --- 外部算力中心不可用时，继续使用现有 HF + RapidOCR 兜底 ---

    hf_error = ''
    hf_headers = get_hf_headers(mimetype or 'application/octet-stream')
    if hf_headers:
        try:
            response = requests.post(hf_model_url(OCR_MODEL_ID), headers=hf_headers, data=img_bytes, timeout=90)
            if response.status_code < 400:
                try:
                    ocr_result = response.json()
                    extracted_text = ''
                    if isinstance(ocr_result, list) and ocr_result and isinstance(ocr_result[0], dict):
                        extracted_text = str(ocr_result[0].get('generated_text', '')).strip()
                    elif isinstance(ocr_result, dict):
                        extracted_text = str(ocr_result.get('generated_text', '')).strip()

                    if extracted_text:
                        return jsonify({"text": extracted_text, "source": "huggingface"})
                    hf_error = "HF OCR returned empty text"
                except Exception:
                    hf_error = f"HF OCR returned non-JSON response: {hf_error_message(response)}"
            else:
                hf_error = f"HF OCR failed ({response.status_code}): {hf_error_message(response)}"
        except Exception as e:
            hf_error = f"HF OCR error: {e}"
    else:
        hf_error = "HF_API_TOKEN is not configured on server"

    local_text, local_error = run_local_ocr(img_bytes)
    if local_text:
        return jsonify({"text": local_text, "source": "rapidocr", "note": hf_error or None})

    runtime_status = get_ocr_runtime_status()

    if '404' in hf_error:
        hf_error = (
            hf_error
            + ". Hugging Face hf-inference currently has no OCR image endpoint for this model/account."
        )

    error_parts = [external_error, hf_error, local_error]
    error_text = ' | '.join([part for part in error_parts if part])
    return jsonify({
        "error": f"OCR failed: {error_text}" if error_text else "OCR failed",
        "details": {
            "external": external_error,
            "huggingface": hf_error,
            "local": local_error,
            "runtime": runtime_status,
            "hint": "Configure EXTERNAL_OCR_SERVICE_URL, HF_API_TOKEN, or enable local RapidOCR fallback."
        }
    }), 502


# ==========================================
# 专家 2 号：语言专家 (负责摘要和提取关键词)
# 对应前端的【按钮 2】
# ==========================================
@app.route('/api/analyze-text', methods=['POST'])
def analyze_text():
    data = request.get_json(silent=True) or {}
    username = str(data.get('username') or '').strip()
    share_token = str(data.get('share_token') or '').strip()
    requested_workspace_id = str(data.get('workspace_id') or '').strip()
    requested_doc_id = parse_int(data.get('doc_id', 0), 0, 0)
    force_refresh = parse_bool(data.get('force_refresh'), False)
    workspace_settings = dict(DEFAULT_WORKSPACE_SETTINGS)
    workspace_id = requested_workspace_id
    doc_text_content = ''
    document_owner_username = ''
    text_source = 'request_text'
    refreshed_from_file = False
    pdf_refresh_meta = {}

    if requested_doc_id > 0:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        try:
            cursor = conn.execute('SELECT * FROM documents WHERE id = ?', (requested_doc_id,))
            doc = cursor.fetchone()
            if not doc:
                return jsonify({'error': 'Document not found'}), 404

            allowed, reason = check_document_access(conn, doc, username, share_token)
            if not allowed:
                return jsonify({'error': reason}), 403

            workspace_id = str(
                (doc.get('workspace_id') if hasattr(doc, 'get') else doc['workspace_id']) or ''
            ).strip()
            workspace_settings = get_workspace_settings(conn, workspace_id)
            doc_text_content = str(
                (doc.get('content') if hasattr(doc, 'get') else doc['content']) or ''
            ).strip()
            document_owner_username = str(
                (doc.get('username') if hasattr(doc, 'get') else doc['username']) or ''
            ).strip()

            # On explicit rebuild, refresh PDF text from source file so summary
            # uses latest/full extraction quality instead of stale db content.
            doc_file_type = str(
                (doc.get('file_type') if hasattr(doc, 'get') else doc['file_type']) or ''
            ).strip().lower()
            doc_filename = str(
                (doc.get('filename') if hasattr(doc, 'get') else doc['filename']) or ''
            ).strip()
            if force_refresh and doc_file_type == 'pdf' and doc_filename:
                try:
                    source_bytes = read_file_bytes_from_storage(doc_filename)
                    refreshed_text, refresh_meta = extract_text_from_pdf_bytes_with_meta(source_bytes)
                    pdf_refresh_meta = refresh_meta if isinstance(refresh_meta, dict) else {}
                    refreshed_text = str(refreshed_text or '').strip()
                    if refreshed_text:
                        if refreshed_text != doc_text_content:
                            conn.execute(
                                'UPDATE documents SET content = ?, content_html = ? WHERE id = ?',
                                (refreshed_text, '', requested_doc_id)
                            )
                            conn.commit()
                        doc_text_content = refreshed_text
                        refreshed_from_file = True
                except Exception as e:
                    print(f"PDF re-extraction on summary refresh failed: {e}")
        finally:
            conn.close()

        if not workspace_settings.get('allow_ai_tools', True):
            return jsonify({"error": "AI tools are disabled in this workspace settings"}), 403
    elif username:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        try:
            workspace_id = requested_workspace_id
            if not workspace_id:
                default_cursor = conn.execute(
                    '''
                    SELECT id
                    FROM workspaces
                    WHERE owner_username = ?
                    ORDER BY created_at ASC, id ASC
                    LIMIT 1
                    ''',
                    (username,)
                )
                default_row = row_to_dict(default_cursor.fetchone())
                workspace_id = str(default_row.get('id') or '').strip()
            if workspace_id and not workspace_belongs_to_user(conn, workspace_id, username):
                return jsonify({'error': 'No access to this workspace'}), 403
            workspace_settings = get_workspace_settings(conn, workspace_id)
        finally:
            conn.close()

        if not workspace_settings.get('allow_ai_tools', True):
            return jsonify({"error": "AI tools are disabled in this workspace settings"}), 403

    text_content = (data.get('text') or '').strip()
    if not text_content and doc_text_content:
        text_content = doc_text_content
        text_source = 'document_content'
    elif not text_content:
        text_source = 'empty'
    summary_length = str(
        data.get('summary_length')
        or workspace_settings.get('summary_length')
        or DEFAULT_WORKSPACE_SETTINGS['summary_length']
    ).strip().lower()
    if summary_length not in WORKSPACE_SUMMARY_LENGTH_LEVELS:
        summary_length = DEFAULT_WORKSPACE_SETTINGS['summary_length']
    keyword_limit = parse_int(
        data.get('keyword_limit', workspace_settings.get('keyword_limit', DEFAULT_WORKSPACE_SETTINGS['keyword_limit'])),
        5,
        3,
        12
    )

    token_map = {
        'short': {'max_new_tokens': 80, 'min_new_tokens': 18, 'sentence_limit': 2},
        'medium': {'max_new_tokens': 120, 'min_new_tokens': 24, 'sentence_limit': 3},
        'long': {'max_new_tokens': 200, 'min_new_tokens': 48, 'sentence_limit': 5},
    }
    length_options = token_map.get(summary_length, token_map['medium'])

    if not text_content:
        if requested_doc_id > 0:
            return jsonify({
                "error": "No text available in this document. Open the note and add/edit content first.",
                "details": {"doc_id": requested_doc_id}
            }), 400
        return jsonify({"error": "No text provided"}), 400

    use_document_cache = requested_doc_id > 0 and text_source == 'document_content'
    text_hash = build_summary_cache_text_hash(text_content)
    text_char_count = len(text_content)
    text_word_count = len(re.findall(r'\S+', text_content))
    base_options_used = {
        "summary_length": summary_length,
        "keyword_limit": keyword_limit,
        "sentence_limit": length_options['sentence_limit'],
        "chunk_count": 1,
        "merge_rounds": 0,
        "refreshed_from_file": refreshed_from_file,
        "pdf_extractor": str(pdf_refresh_meta.get('extractor') or ''),
        "pdf_ocr_attempted": bool(pdf_refresh_meta.get('ocr_attempted')),
        "pdf_ocr_used": bool(pdf_refresh_meta.get('ocr_used')),
        "pdf_quality_score_before": parse_float(pdf_refresh_meta.get('quality_score_before'), 0.0),
        "pdf_quality_score_after": parse_float(pdf_refresh_meta.get('quality_score_after'), 0.0),
        "text_char_count": text_char_count,
        "text_word_count": text_word_count,
        "summarizer_model": SUMMARIZER_MODEL_ID,
    }
    if use_document_cache and text_hash and not force_refresh:
        conn = get_db_connection()
        if conn:
            try:
                cached_payload = load_document_summary_cache(
                    conn,
                    requested_doc_id,
                    text_hash,
                    summary_length,
                    keyword_limit
                )
            finally:
                conn.close()
            if cached_payload:
                cached_options_raw = cached_payload.get("options_used")
                cached_options = cached_options_raw if isinstance(cached_options_raw, dict) else {}
                options_used = dict(base_options_used)
                if cached_options:
                    options_used["summary_length"] = str(
                        cached_options.get("summary_length") or summary_length
                    ).strip().lower() or summary_length
                    options_used["keyword_limit"] = parse_int(
                        cached_options.get("keyword_limit"),
                        keyword_limit,
                        3,
                        12
                    )
                    options_used["sentence_limit"] = parse_int(
                        cached_options.get("sentence_limit"),
                        length_options['sentence_limit'],
                        1,
                        20
                    )
                    options_used["chunk_count"] = parse_int(
                        cached_options.get("chunk_count"),
                        1,
                        1
                    )
                    options_used["merge_rounds"] = parse_int(
                        cached_options.get("merge_rounds"),
                        0,
                        0
                    )
                    options_used["refreshed_from_file"] = parse_bool(
                        cached_options.get("refreshed_from_file"),
                        refreshed_from_file
                    )
                    options_used["pdf_extractor"] = str(
                        cached_options.get("pdf_extractor") or options_used["pdf_extractor"]
                    ).strip()
                    options_used["pdf_ocr_attempted"] = parse_bool(
                        cached_options.get("pdf_ocr_attempted"),
                        options_used["pdf_ocr_attempted"]
                    )
                    options_used["pdf_ocr_used"] = parse_bool(
                        cached_options.get("pdf_ocr_used"),
                        options_used["pdf_ocr_used"]
                    )
                    options_used["pdf_quality_score_before"] = parse_float(
                        cached_options.get("pdf_quality_score_before"),
                        options_used["pdf_quality_score_before"]
                    )
                    options_used["pdf_quality_score_after"] = parse_float(
                        cached_options.get("pdf_quality_score_after"),
                        options_used["pdf_quality_score_after"]
                    )
                    options_used["text_char_count"] = parse_int(
                        cached_options.get("text_char_count"),
                        text_char_count,
                        0
                    )
                    options_used["text_word_count"] = parse_int(
                        cached_options.get("text_word_count"),
                        text_word_count,
                        0
                    )
                    options_used["summarizer_model"] = str(
                        cached_options.get("summarizer_model") or SUMMARIZER_MODEL_ID
                    ).strip() or SUMMARIZER_MODEL_ID
                return jsonify({
                    "summary": str(cached_payload.get("summary") or '').strip(),
                    "keywords": cached_payload.get("keywords") if isinstance(cached_payload.get("keywords"), list) else [],
                    "key_sentences": (
                        cached_payload.get("key_sentences")
                        if isinstance(cached_payload.get("key_sentences"), list)
                        else []
                    ),
                    "summary_source": str(
                        cached_payload.get("summary_source") or "cache"
                    ).strip().lower() or "cache",
                    "summary_note": str(cached_payload.get("summary_note") or '').strip(),
                    "text_source": text_source,
                    "document_id": requested_doc_id,
                    "cache_hit": True,
                    "cached_at": cached_payload.get("cached_at"),
                    "options_used": options_used,
                })

    summary_result = summarize_text_with_chunk_merge(text_content, length_options)
    summary = str(summary_result.get('summary') or '').strip()
    summary_source = str(summary_result.get('summary_source') or 'fallback').strip().lower() or 'fallback'
    summary_note = str(summary_result.get('summary_note') or '').strip()
    summary_meta = summary_result.get('meta') if isinstance(summary_result.get('meta'), dict) else {}
    pdf_refresh_note = str(pdf_refresh_meta.get('note') or '').strip()

    if pdf_refresh_note and force_refresh and requested_doc_id > 0:
        summary_note = f"{summary_note}; PDF refresh: {pdf_refresh_note}" if summary_note else f"PDF refresh: {pdf_refresh_note}"

    if not summary:
        summary = build_fallback_summary(text_content, sentence_limit=length_options['sentence_limit'], max_chars=560)
        summary_source = 'fallback'
        if not summary_note:
            summary_note = "Summary service returned empty output."

    keywords = []
    try:
        if len(text_content.split()) > 5:
            vectorizer = TfidfVectorizer(stop_words='english', max_features=keyword_limit)
            vectorizer.fit_transform([text_content])
            keywords = vectorizer.get_feature_names_out().tolist()
    except Exception:
        keywords = ["Not enough text"]

    key_sentences = extract_key_sentences(text_content, keywords, limit=length_options['sentence_limit'])

    response_payload = {
        "summary": summary,
        "keywords": keywords,
        "key_sentences": key_sentences,
        "summary_source": summary_source,
        "summary_note": summary_note,
        "text_source": text_source,
        "document_id": requested_doc_id if requested_doc_id > 0 else None,
        "cache_hit": False,
        "options_used": {
            **base_options_used,
            "chunk_count": parse_int(summary_meta.get('chunk_count'), 1, 1),
            "merge_rounds": parse_int(summary_meta.get('merge_rounds'), 0, 0),
            "refreshed_from_file": refreshed_from_file,
            "pdf_extractor": str(pdf_refresh_meta.get('extractor') or base_options_used.get('pdf_extractor') or ''),
            "pdf_ocr_attempted": bool(pdf_refresh_meta.get('ocr_attempted')) or bool(base_options_used.get('pdf_ocr_attempted')),
            "pdf_ocr_used": bool(pdf_refresh_meta.get('ocr_used')),
            "pdf_quality_score_before": parse_float(
                pdf_refresh_meta.get('quality_score_before'),
                parse_float(base_options_used.get('pdf_quality_score_before'), 0.0)
            ),
            "pdf_quality_score_after": parse_float(
                pdf_refresh_meta.get('quality_score_after'),
                parse_float(base_options_used.get('pdf_quality_score_after'), 0.0)
            ),
        },
    }

    if use_document_cache and text_hash:
        cache_conn = get_db_connection()
        if cache_conn:
            try:
                save_document_summary_cache(
                    cache_conn,
                    requested_doc_id,
                    workspace_id,
                    username or document_owner_username,
                    text_hash,
                    summary_length,
                    keyword_limit,
                    {
                        "summary": summary,
                        "keywords": keywords,
                        "key_sentences": key_sentences,
                        "summary_source": summary_source,
                        "summary_note": summary_note,
                        "options_used": response_payload.get("options_used") if isinstance(response_payload.get("options_used"), dict) else {},
                    }
                )
            finally:
                cache_conn.close()

    return jsonify(response_payload)

# ================= 修改后的下载/访问接口 (支持 S3) =================
@app.route('/uploads/<filename>')
def uploaded_file(filename):
    username = (request.args.get('username') or '').strip()
    share_token = (request.args.get('share_token') or '').strip()
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.execute(
                'SELECT * FROM documents WHERE filename = ? ORDER BY id DESC LIMIT 1',
                (filename,)
            )
            doc = cursor.fetchone()
            if doc:
                allowed, reason = check_document_access(conn, doc, username, share_token)
                if not allowed:
                    return jsonify({'error': reason}), 403
        finally:
            conn.close()

    # 如果配置了 S3，直接生成一个 S3 的链接跳转过去
    if S3_BUCKET and s3_client:
        try:
            # 生成一个“预签名 URL”，有效期 1 小时 (3600秒)
            presigned_url = s3_client.generate_presigned_url(
                'get_object',
                Params={'Bucket': S3_BUCKET, 'Key': filename},
                ExpiresIn=3600
            )
            # 让浏览器直接跳转到 AWS S3 下载
            return redirect(presigned_url, code=302)
        except Exception as e:
            print(f"S3 Link Generation Error: {e}")
            return jsonify({'error': 'Could not generate file link'}), 500
    else:
        # 如果没配 S3 (比如本地测试)，还是从本地文件夹读
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# ================= 前端路由 =================
@app.route('/')
def serve_index():
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/<path:path>')
def catch_all(path):
    if path.startswith('api/') or path.startswith('uploads/'):
        return jsonify({'error': 'Not found'}), 404
    
    if os.path.exists(os.path.join(app.static_folder, path)):
        return send_from_directory(app.static_folder, path)
    
    return send_from_directory(app.static_folder, 'index.html')

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    app.run(debug=True, port=port, host='0.0.0.0')
