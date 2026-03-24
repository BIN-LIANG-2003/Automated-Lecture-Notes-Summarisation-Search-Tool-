import io
import os
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime, timedelta
from html import escape as html_escape

import docx
import PyPDF2
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from lxml import etree, html as lxml_html

try:
    import fitz  # PyMuPDF
except Exception as e:
    fitz = None
    print(f"⚠️ PyMuPDF unavailable: {e}")

from .config import (
    BLOCK_TAGS,
    CATEGORY_KEYWORDS,
    DEFAULT_DOCUMENT_CATEGORY,
    EDITOR_ALLOWED_STYLE_PROPS,
    EDITOR_ALLOWED_TAGS,
    ENABLE_PDF_OCR_FALLBACK,
    HIGHLIGHT_RGB_BY_INDEX,
    MIME_BY_EXT,
    NAMED_COLORS,
    OCRMYPDF_BINARY,
    OCRMYPDF_LANGUAGE,
    OCRMYPDF_TIMEOUT_SECONDS,
    TRASH_RETENTION_DAYS,
)
from .storage import remove_document_file_from_storage
from .utils import parse_int, row_to_dict
from .workspace_domain import workspace_belongs_to_user


def hard_delete_document_record(conn, doc_id):
    safe_doc_id = parse_int(doc_id, 0, 0)
    if safe_doc_id <= 0:
        return None
    cursor = conn.execute(
        'SELECT id, filename, username, workspace_id, deleted_at FROM documents WHERE id = ?',
        (safe_doc_id,),
    )
    doc_row = cursor.fetchone()
    if not doc_row:
        return None
    doc = row_to_dict(doc_row) or {}
    conn.execute('DELETE FROM document_share_links WHERE document_id = ?', (safe_doc_id,))
    conn.execute('DELETE FROM document_summary_cache WHERE document_id = ?', (safe_doc_id,))
    conn.execute('DELETE FROM documents WHERE id = ?', (safe_doc_id,))
    return doc


def purge_expired_trashed_documents(conn, username='', workspace_id=''):
    safe_username = str(username or '').strip()
    safe_workspace_id = str(workspace_id or '').strip()
    cutoff = (datetime.utcnow() - timedelta(days=TRASH_RETENTION_DAYS)).isoformat()
    where_parts = [
        "COALESCE(deleted_at, '') <> ''",
        "COALESCE(deleted_at, '') <= ?",
    ]
    params = [cutoff]
    if safe_username:
        where_parts.append('username = ?')
        params.append(safe_username)
    if safe_workspace_id:
        where_parts.append('workspace_id = ?')
        params.append(safe_workspace_id)

    where_sql = ' AND '.join(where_parts)
    cursor = conn.execute(
        f'''
        SELECT id, filename
        FROM documents
        WHERE {where_sql}
        ORDER BY deleted_at ASC, id ASC
        ''',
        tuple(params),
    )
    stale_rows = [row_to_dict(item) for item in cursor.fetchall()]
    if not stale_rows:
        return {'purged_count': 0, 'warnings': []}

    purged_files = []
    for row in stale_rows:
        doc_id = parse_int((row or {}).get('id'), 0, 0)
        if doc_id <= 0:
            continue
        deleted = hard_delete_document_record(conn, doc_id)
        if deleted:
            purged_files.append(str((deleted or {}).get('filename') or '').strip())
    conn.commit()

    warnings = []
    for filename in purged_files:
        warning = remove_document_file_from_storage(filename)
        if warning:
            warnings.append(f'{filename}: {warning}')
    return {'purged_count': len(purged_files), 'warnings': warnings}


def user_can_edit_document(conn, doc_row, username=''):
    doc = row_to_dict(doc_row) or {}
    editor = str(username or '').strip()
    if not editor:
        return False

    workspace_id = str(doc.get('workspace_id') or '').strip()
    owner_username = str(doc.get('username') or '').strip()
    if workspace_id:
        return workspace_belongs_to_user(conn, workspace_id, editor)
    if owner_username:
        return owner_username == editor
    return False


def normalize_newlines(value):
    text = value if isinstance(value, str) else str(value or '')
    return text.replace('\r\n', '\n').replace('\r', '\n')


def infer_document_category(title, text_content=''):
    title_text = str(title or '')
    body_text = str(text_content or '')[:5000]
    source = f"{title_text}\n{body_text}".lower()
    if not source.strip():
        return DEFAULT_DOCUMENT_CATEGORY

    for category, keywords in CATEGORY_KEYWORDS.items():
        if any(keyword in source for keyword in keywords):
            return category
    return DEFAULT_DOCUMENT_CATEGORY


def sanitize_style_declarations(style_text):
    style_map = {}
    if not isinstance(style_text, str):
        return style_map

    for declaration in style_text.split(';'):
        if ':' not in declaration:
            continue
        prop, val = declaration.split(':', 1)
        prop = prop.strip().lower()
        val = val.strip()
        if prop not in EDITOR_ALLOWED_STYLE_PROPS or not val:
            continue

        lower_val = val.lower()
        if 'expression(' in lower_val or 'javascript:' in lower_val or 'url(' in lower_val:
            continue

        if prop == 'font-family':
            cleaned_parts = [
                part.strip().strip('"').strip("'")
                for part in val.split(',')
                if part.strip()
            ]
            if not cleaned_parts:
                continue
            val = ', '.join(cleaned_parts[:3])
        elif prop in ('width', 'height', 'margin-left'):
            if not re.fullmatch(r'\d+(?:\.\d+)?(px|pt|em|rem|%)', lower_val):
                continue
        elif prop == 'border-collapse':
            if lower_val not in ('collapse', 'separate'):
                continue
        elif prop == 'border':
            if not re.fullmatch(r'[\w\s.#()-]+', val):
                continue

        style_map[prop] = val

    return style_map


def style_map_to_inline(style_map):
    if not style_map:
        return ''
    return '; '.join(f'{k}: {v}' for k, v in style_map.items())


def parse_css_color(color_value):
    if not isinstance(color_value, str):
        return None
    value = color_value.strip().lower()
    if not value:
        return None

    if value in NAMED_COLORS:
        return NAMED_COLORS[value]

    if re.fullmatch(r'#?[0-9a-f]{6}', value):
        hex_color = value.lstrip('#')
        return tuple(int(hex_color[idx:idx + 2], 16) for idx in (0, 2, 4))

    rgb_match = re.fullmatch(
        r'rgb\(\s*(\d{1,3})\s*,\s*(\d{1,3})\s*,\s*(\d{1,3})\s*\)',
        value,
    )
    if rgb_match:
        channels = [max(0, min(255, int(item))) for item in rgb_match.groups()]
        return tuple(channels)

    return None


def sanitize_int_attr(value, min_value=1, max_value=10000):
    raw = str(value or '').strip()
    if not raw or not raw.isdigit():
        return None
    number = int(raw)
    if number < min_value:
        return None
    return min(number, max_value)


def is_safe_link_href(href):
    value = str(href or '').strip()
    if not value:
        return False
    lower = value.lower()
    if lower.startswith(('javascript:', 'data:', 'vbscript:')):
        return False
    return True


def is_safe_image_src(src):
    value = str(src or '').strip()
    if not value:
        return False
    lower = value.lower()
    if lower.startswith(('javascript:', 'vbscript:')):
        return False
    if lower.startswith('data:'):
        return lower.startswith('data:image/')
    return True


def sanitize_colwidth_attr(value):
    raw = str(value or '').strip()
    if not raw:
        return None
    if re.fullmatch(r'\d+(,\d+)*', raw):
        return raw
    return None


def pick_highlight_index_from_css(color_value):
    rgb = parse_css_color(color_value)
    if not rgb:
        return None

    closest_index = None
    closest_distance = None
    for index, target_rgb in HIGHLIGHT_RGB_BY_INDEX.items():
        distance = (
            (rgb[0] - target_rgb[0]) ** 2
            + (rgb[1] - target_rgb[1]) ** 2
            + (rgb[2] - target_rgb[2]) ** 2
        )
        if closest_distance is None or distance < closest_distance:
            closest_distance = distance
            closest_index = index

    return closest_index


def parse_css_font_size_pt(size_value):
    if not isinstance(size_value, str):
        return None
    value = size_value.strip().lower()
    if not value:
        return None
    match = re.fullmatch(r'([0-9]+(?:\.[0-9]+)?)(pt|px|em|rem)?', value)
    if not match:
        return None

    amount = float(match.group(1))
    unit = (match.group(2) or 'pt')
    if unit == 'pt':
        return amount
    if unit == 'px':
        return amount * 0.75
    if unit in ('em', 'rem'):
        return amount * 12.0
    return amount


def css_alignment_from_docx_alignment(alignment):
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return 'center'
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return 'right'
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return 'justify'
    return ''


def docx_alignment_from_css(style_map):
    alignment = (style_map.get('text-align') or '').strip().lower()
    if alignment == 'center':
        return WD_ALIGN_PARAGRAPH.CENTER
    if alignment == 'right':
        return WD_ALIGN_PARAGRAPH.RIGHT
    if alignment == 'justify':
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def apply_block_style_to_paragraph(paragraph, style_map):
    paragraph.alignment = docx_alignment_from_css(style_map)
    margin_left = parse_css_font_size_pt(style_map.get('margin-left'))
    if isinstance(margin_left, (int, float)) and margin_left > 0:
        paragraph.paragraph_format.left_indent = Pt(margin_left)


def plaintext_to_html(content):
    text = normalize_newlines(content)
    lines = text.split('\n')
    if not lines:
        return '<p><br></p>'

    blocks = []
    for line in lines:
        if line == '':
            blocks.append('<p><br></p>')
        else:
            blocks.append(f'<p>{html_escape(line)}</p>')
    return ''.join(blocks) or '<p><br></p>'


def html_to_plaintext(content_html):
    if not isinstance(content_html, str) or not content_html.strip():
        return ''

    normalized_html = re.sub(r'(?i)<br\s*/?>', '\n', content_html)
    normalized_html = re.sub(
        r'(?i)</(p|div|li|h[1-6]|blockquote|pre|ul|ol|table|thead|tbody|tr|th|td)>',
        '\n',
        normalized_html,
    )
    normalized_html = re.sub(r'(?i)<hr\s*/?>', '\n', normalized_html)

    try:
        root = lxml_html.fragment_fromstring(normalized_html, create_parent='div')
        text = root.text_content()
    except Exception:
        text = re.sub(r'<[^>]+>', '', normalized_html)

    text = text.replace('\xa0', ' ')
    text = normalize_newlines(text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def sanitize_editor_html(raw_html):
    if raw_html is None:
        return '<p><br></p>'

    source_html = str(raw_html).strip()
    if not source_html:
        return '<p><br></p>'

    source_html = re.sub(r'(?is)<(script|style)[^>]*>.*?</\1>', '', source_html)
    source_html = re.sub(r'(?is)<!--.*?-->', '', source_html)

    try:
        root = lxml_html.fragment_fromstring(source_html, create_parent='div')
    except Exception:
        return plaintext_to_html(source_html)

    present_tags = {el.tag.lower() for el in root.iter() if isinstance(el.tag, str)}
    strip_tags = [tag for tag in present_tags if tag not in EDITOR_ALLOWED_TAGS]
    if strip_tags:
        etree.strip_tags(root, *strip_tags)

    for el in root.iter():
        if not isinstance(el.tag, str):
            continue
        tag = el.tag.lower()
        style_map = sanitize_style_declarations(el.attrib.get('style', ''))

        sanitized_attrs = {}
        if style_map:
            sanitized_attrs['style'] = style_map_to_inline(style_map)

        if tag == 'a':
            href = (el.attrib.get('href') or '').strip()
            if is_safe_link_href(href):
                sanitized_attrs['href'] = href
                sanitized_attrs['target'] = '_blank'
                sanitized_attrs['rel'] = 'noopener noreferrer'
        elif tag == 'img':
            src = (el.attrib.get('src') or '').strip()
            if is_safe_image_src(src):
                sanitized_attrs['src'] = src

            alt = str(el.attrib.get('alt') or '').strip()
            if alt:
                sanitized_attrs['alt'] = alt[:200]

            title = str(el.attrib.get('title') or '').strip()
            if title:
                sanitized_attrs['title'] = title[:200]

            width = sanitize_int_attr(el.attrib.get('width'), 1, 4000)
            if width:
                sanitized_attrs['width'] = str(width)

            height = sanitize_int_attr(el.attrib.get('height'), 1, 4000)
            if height:
                sanitized_attrs['height'] = str(height)
        elif tag in ('th', 'td'):
            colspan = sanitize_int_attr(el.attrib.get('colspan'), 1, 20)
            if colspan and colspan > 1:
                sanitized_attrs['colspan'] = str(colspan)

            rowspan = sanitize_int_attr(el.attrib.get('rowspan'), 1, 100)
            if rowspan and rowspan > 1:
                sanitized_attrs['rowspan'] = str(rowspan)

            colwidth = sanitize_colwidth_attr(el.attrib.get('colwidth'))
            if colwidth:
                sanitized_attrs['colwidth'] = colwidth
        elif tag == 'col':
            span = sanitize_int_attr(el.attrib.get('span'), 1, 20)
            if span and span > 1:
                sanitized_attrs['span'] = str(span)

            width = sanitize_int_attr(el.attrib.get('width'), 1, 4000)
            if width:
                sanitized_attrs['width'] = str(width)

        el.attrib.clear()
        if sanitized_attrs:
            el.attrib.update(sanitized_attrs)

    serialized_parts = []
    if root.text and root.text.strip():
        serialized_parts.append(f'<p>{html_escape(root.text)}</p>')
    for child in root:
        serialized_parts.append(lxml_html.tostring(child, encoding='unicode', method='html'))

    sanitized_html = ''.join(serialized_parts).strip()
    return sanitized_html or '<p><br></p>'


def apply_run_style(run, style_ctx):
    run.bold = bool(style_ctx.get('bold'))
    run.italic = bool(style_ctx.get('italic'))
    run.underline = bool(style_ctx.get('underline'))
    run.font.strike = bool(style_ctx.get('strike'))
    has_subscript = bool(style_ctx.get('subscript'))
    has_superscript = bool(style_ctx.get('superscript'))
    if has_subscript and has_superscript:
        has_subscript = False
    run.font.subscript = has_subscript
    run.font.superscript = has_superscript

    font_size_pt = style_ctx.get('font_size_pt')
    if isinstance(font_size_pt, (int, float)) and font_size_pt > 0:
        run.font.size = Pt(font_size_pt)

    font_name = style_ctx.get('font_name')
    if isinstance(font_name, str) and font_name.strip():
        run.font.name = font_name.strip()

    rgb = style_ctx.get('color_rgb')
    if isinstance(rgb, tuple) and len(rgb) == 3:
        run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])

    highlight_index = style_ctx.get('highlight_index')
    if highlight_index in HIGHLIGHT_RGB_BY_INDEX:
        run.font.highlight_color = highlight_index


def add_text_to_paragraph(paragraph, text, style_ctx):
    if not text:
        return
    parts = text.split('\n')
    for idx, part in enumerate(parts):
        run = paragraph.add_run(part)
        apply_run_style(run, style_ctx)
        if idx < len(parts) - 1:
            run.add_break()


def merge_inline_style(base_style, node):
    style = dict(base_style or {})
    tag = node.tag.lower() if isinstance(node.tag, str) else ''

    if tag in ('strong', 'b'):
        style['bold'] = True
    elif tag in ('em', 'i'):
        style['italic'] = True
    elif tag == 'u':
        style['underline'] = True
    elif tag in ('s', 'strike', 'del'):
        style['strike'] = True
    elif tag == 'sub':
        style['subscript'] = True
        style['superscript'] = False
    elif tag == 'sup':
        style['superscript'] = True
        style['subscript'] = False
    elif tag == 'mark':
        style['highlight_index'] = WD_COLOR_INDEX.YELLOW
    elif tag == 'code':
        style['font_name'] = 'Courier New'
    elif tag == 'a':
        style['underline'] = True
        style['color_rgb'] = (29, 78, 216)

    style_map = sanitize_style_declarations(node.attrib.get('style', ''))
    font_weight = (style_map.get('font-weight') or '').lower()
    if font_weight == 'bold':
        style['bold'] = True
    elif font_weight.isdigit():
        style['bold'] = int(font_weight) >= 600

    if (style_map.get('font-style') or '').lower() == 'italic':
        style['italic'] = True

    decoration = (style_map.get('text-decoration') or '').lower()
    if 'underline' in decoration:
        style['underline'] = True
    if 'line-through' in decoration:
        style['strike'] = True

    color_rgb = parse_css_color(style_map.get('color'))
    if color_rgb:
        style['color_rgb'] = color_rgb

    highlight_index = pick_highlight_index_from_css(style_map.get('background-color'))
    if highlight_index:
        style['highlight_index'] = highlight_index

    vertical_align = (style_map.get('vertical-align') or '').strip().lower()
    if vertical_align == 'sub':
        style['subscript'] = True
        style['superscript'] = False
    elif vertical_align in ('super', 'sup'):
        style['superscript'] = True
        style['subscript'] = False

    font_size_pt = parse_css_font_size_pt(style_map.get('font-size'))
    if font_size_pt:
        style['font_size_pt'] = font_size_pt

    font_family = style_map.get('font-family')
    if font_family:
        style['font_name'] = font_family.split(',')[0].strip()

    return style


def add_inline_node_to_paragraph(paragraph, node, inherited_style=None):
    current_style = merge_inline_style(inherited_style or {}, node)

    if node.text:
        add_text_to_paragraph(paragraph, node.text, current_style)

    for child in node:
        child_tag = child.tag.lower() if isinstance(child.tag, str) else ''
        if child_tag == 'br':
            paragraph.add_run().add_break()
        elif child_tag == 'img':
            alt = (child.attrib.get('alt') or '').strip()
            src = (child.attrib.get('src') or '').strip()
            placeholder = alt or src or 'Image'
            add_text_to_paragraph(paragraph, f'[Image] {placeholder}', current_style)
        elif child_tag == 'hr':
            add_text_to_paragraph(paragraph, '------------------------------', current_style)
        elif child_tag in BLOCK_TAGS.union({'ul', 'ol'}):
            nested_text = html_to_plaintext(lxml_html.tostring(child, encoding='unicode', method='html'))
            if nested_text:
                paragraph.add_run().add_break()
                add_text_to_paragraph(paragraph, nested_text, current_style)
        else:
            add_inline_node_to_paragraph(paragraph, child, current_style)

        if child.tail:
            add_text_to_paragraph(paragraph, child.tail, current_style)


def append_html_element_to_docx(document, element):
    if not isinstance(element.tag, str):
        return

    tag = element.tag.lower()
    style_map = sanitize_style_declarations(element.attrib.get('style', ''))

    if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(tag[1])
        paragraph = document.add_paragraph(style=f'Heading {min(level, 6)}')
        apply_block_style_to_paragraph(paragraph, style_map)
        add_inline_node_to_paragraph(paragraph, element, {})
        return

    if tag in ('ul', 'ol'):
        list_style = 'List Number' if tag == 'ol' else 'List Bullet'
        items = [child for child in element if isinstance(child.tag, str) and child.tag.lower() == 'li']
        if not items and (element.text or '').strip():
            items = [element]
        for item in items:
            paragraph = document.add_paragraph(style=list_style)
            add_inline_node_to_paragraph(paragraph, item, {})
        return

    if tag == 'blockquote':
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.space_after = Pt(6)
        apply_block_style_to_paragraph(paragraph, style_map)
        add_inline_node_to_paragraph(paragraph, element, {})
        return

    if tag == 'pre':
        text = normalize_newlines(element.text_content())
        lines = text.split('\n')
        if not lines:
            lines = ['']
        for line in lines:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = 'Courier New'
        return

    if tag == 'li':
        paragraph = document.add_paragraph(style='List Bullet')
        add_inline_node_to_paragraph(paragraph, element, {})
        return

    if tag == 'table':
        row_elements = []
        for child in element.iter():
            child_tag = child.tag.lower() if isinstance(child.tag, str) else ''
            if child_tag == 'tr':
                row_elements.append(child)

        if not row_elements:
            return

        max_cols = 1
        parsed_rows = []
        for row_el in row_elements:
            cells = [
                cell for cell in row_el
                if isinstance(cell.tag, str) and cell.tag.lower() in ('td', 'th')
            ]
            if not cells:
                continue
            parsed_rows.append(cells)
            max_cols = max(max_cols, len(cells))

        if not parsed_rows:
            return

        table = document.add_table(rows=len(parsed_rows), cols=max_cols)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass

        for row_idx, row_cells in enumerate(parsed_rows):
            for col_idx in range(max_cols):
                target_cell = table.cell(row_idx, col_idx)
                if col_idx >= len(row_cells):
                    target_cell.text = ''
                    continue
                cell_text = normalize_newlines(row_cells[col_idx].text_content()).strip()
                target_cell.text = cell_text
        return

    if tag == 'img':
        alt = (element.attrib.get('alt') or '').strip()
        src = (element.attrib.get('src') or '').strip()
        placeholder = alt or src or 'Image'
        document.add_paragraph(f'[Image] {placeholder}')
        return

    if tag == 'hr':
        document.add_paragraph('------------------------------')
        return

    paragraph = document.add_paragraph()
    apply_block_style_to_paragraph(paragraph, style_map)
    add_inline_node_to_paragraph(paragraph, element, {})


def create_docx_bytes_from_html(content_html, fallback_text=''):
    html_content = sanitize_editor_html(content_html)
    document = docx.Document()

    try:
        root = lxml_html.fragment_fromstring(html_content, create_parent='div')
    except Exception:
        root = None

    if root is None:
        lines = normalize_newlines(fallback_text).split('\n')
        for line in lines or ['']:
            document.add_paragraph(line)
    else:
        if root.text and root.text.strip():
            document.add_paragraph(root.text.strip())

        for child in root:
            append_html_element_to_docx(document, child)
            if child.tail and child.tail.strip():
                document.add_paragraph(child.tail.strip())

        if not document.paragraphs:
            fallback_lines = normalize_newlines(fallback_text).split('\n')
            for line in fallback_lines or ['']:
                document.add_paragraph(line)

    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream.read()


def run_to_html(run):
    raw_text = run.text or ''
    if raw_text == '':
        return ''

    chunk = html_escape(raw_text).replace('\n', '<br/>')
    style_parts = []

    font_size = run.font.size.pt if run.font and run.font.size else None
    if font_size:
        style_parts.append(f'font-size: {round(font_size, 2)}pt')

    font_name = run.font.name if run.font else ''
    if font_name:
        style_parts.append(f'font-family: {html_escape(font_name)}')

    font_color = run.font.color.rgb if run.font and run.font.color else None
    if font_color:
        color_hex = str(font_color)
        if re.fullmatch(r'[0-9A-Fa-f]{6}', color_hex):
            style_parts.append(f'color: #{color_hex}')

    highlight_color = run.font.highlight_color if run.font else None
    if highlight_color in HIGHLIGHT_RGB_BY_INDEX:
        rgb = HIGHLIGHT_RGB_BY_INDEX[highlight_color]
        style_parts.append(f'background-color: #{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}')

    if style_parts:
        chunk = f'<span style="{"; ".join(style_parts)}">{chunk}</span>'

    if run.bold:
        chunk = f'<strong>{chunk}</strong>'
    if run.italic:
        chunk = f'<em>{chunk}</em>'
    if run.underline:
        chunk = f'<u>{chunk}</u>'
    if run.font and run.font.strike:
        chunk = f'<s>{chunk}</s>'
    if run.font and run.font.subscript:
        chunk = f'<sub>{chunk}</sub>'
    if run.font and run.font.superscript:
        chunk = f'<sup>{chunk}</sup>'

    return chunk


def paragraph_to_html_block(paragraph):
    style_name = (paragraph.style.name if paragraph.style else '').strip().lower()
    alignment = css_alignment_from_docx_alignment(paragraph.alignment)
    style_attr = f' style="text-align: {alignment}"' if alignment else ''

    inline_html = ''.join(run_to_html(run) for run in paragraph.runs)
    if not inline_html:
        inline_html = html_escape(paragraph.text or '').replace('\n', '<br/>')
    if not inline_html:
        inline_html = '<br/>'

    if 'list bullet' in style_name:
        return 'ul', f'<li{style_attr}>{inline_html}</li>'
    if 'list number' in style_name:
        return 'ol', f'<li{style_attr}>{inline_html}</li>'

    heading_match = re.search(r'heading\s*([1-6])', style_name)
    if heading_match:
        level = heading_match.group(1)
        return '', f'<h{level}{style_attr}>{inline_html}</h{level}>'

    return '', f'<p{style_attr}>{inline_html}</p>'


def extract_docx_content(filepath):
    document = docx.Document(filepath)
    plain_lines = []
    html_parts = []
    list_type = ''
    list_items = []

    def flush_list():
        nonlocal list_type, list_items
        if list_type and list_items:
            html_parts.append(f'<{list_type}>{"".join(list_items)}</{list_type}>')
        list_type = ''
        list_items = []

    for paragraph in document.paragraphs:
        plain_lines.append(paragraph.text or '')
        block_type, block_html = paragraph_to_html_block(paragraph)
        if block_type in ('ul', 'ol'):
            if list_type and list_type != block_type:
                flush_list()
            list_type = block_type
            list_items.append(block_html)
        else:
            flush_list()
            html_parts.append(block_html)

    flush_list()

    for table in document.tables:
        row_html = []
        for row in table.rows:
            cell_html = []
            cell_text_parts = []
            for cell in row.cells:
                text_value = normalize_newlines(cell.text or '').strip()
                cell_text_parts.append(text_value)
                cell_html.append(f'<td>{html_escape(text_value) if text_value else "<br/>"}</td>')
            plain_lines.append(' | '.join(cell_text_parts).strip())
            row_html.append(f'<tr>{"".join(cell_html)}</tr>')
        if row_html:
            html_parts.append(f'<table><tbody>{"".join(row_html)}</tbody></table>')

    plain_text = '\n'.join(plain_lines)
    html_content = ''.join(html_parts).strip() or plaintext_to_html(plain_text)
    return plain_text, sanitize_editor_html(html_content)


def extract_text_content(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def extract_document_content(filepath, ext):
    file_ext = (ext or '').lower().strip('.')
    text = ''
    content_html = ''

    try:
        if file_ext == 'docx':
            text, content_html = extract_docx_content(filepath)
        elif file_ext == 'pdf':
            with open(filepath, 'rb') as f:
                file_bytes = f.read()
            text = extract_text_from_pdf_bytes(file_bytes)
        elif file_ext == 'txt':
            text = extract_text_content(filepath)
            content_html = plaintext_to_html(text)
    except Exception as e:
        print(f"Error extracting content: {e}")
        text = 'Text extraction failed.'
        content_html = plaintext_to_html(text)

    if file_ext in ('docx', 'txt'):
        content_html = sanitize_editor_html(content_html or plaintext_to_html(text))

    return text, content_html


def normalize_pdf_text(text):
    value = normalize_newlines(text or '')
    value = value.replace('\x00', ' ')
    value = re.sub(r'-\n(?=[A-Za-z])', '', value)
    value = re.sub(r'[ \t]+', ' ', value)
    value = re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', value)
    value = re.sub(r'(?<=[A-Za-z])(?=\d)', ' ', value)
    value = re.sub(r'(?<=\d)(?=[A-Za-z])', ' ', value)
    value = re.sub(r'([。！？.!?])(?=[A-Za-z0-9])', r'\1 ', value)
    value = re.sub(r'\n{3,}', '\n\n', value)
    return value.strip()


def compute_pdf_text_quality_metrics(text):
    normalized = normalize_pdf_text(text or '')
    tokens = re.findall(r'\S+', normalized)
    token_count = len(tokens)
    char_count = len(normalized)
    if token_count <= 0:
        return {
            'char_count': char_count,
            'token_count': 0,
            'avg_token_len': 0.0,
            'long_token_ratio': 0.0,
            'cjk_ratio': 0.0,
            'line_count': 0,
        }

    avg_token_len = sum(len(item) for item in tokens) / max(1, token_count)
    long_token_count = sum(1 for item in tokens if len(item) >= 18)
    long_token_ratio = long_token_count / max(1, token_count)
    cjk_chars = re.findall(r'[\u3400-\u9fff]', normalized)
    cjk_ratio = len(cjk_chars) / max(1, char_count)
    line_count = len([line for line in normalized.split('\n') if line.strip()])
    return {
        'char_count': char_count,
        'token_count': token_count,
        'avg_token_len': avg_token_len,
        'long_token_ratio': long_token_ratio,
        'cjk_ratio': cjk_ratio,
        'line_count': line_count,
    }


def score_pdf_text_quality(text):
    metrics = compute_pdf_text_quality_metrics(text)
    char_count = metrics['char_count']
    token_count = metrics['token_count']
    avg_token_len = metrics['avg_token_len']
    long_token_ratio = metrics['long_token_ratio']
    cjk_ratio = metrics['cjk_ratio']
    line_count = metrics['line_count']

    if char_count <= 0 or token_count <= 0:
        return 0.0, metrics

    score = 0.0
    score += min(char_count / 650.0, 28.0)
    score += min(token_count / 90.0, 20.0)
    score += min(line_count / 80.0, 6.0)
    if cjk_ratio < 0.2:
        score -= max(0.0, (avg_token_len - 8.8)) * 2.8
        score -= long_token_ratio * 36.0
    else:
        score -= max(0.0, (avg_token_len - 12.0)) * 1.2
        score -= long_token_ratio * 10.0

    return round(score, 3), metrics


def should_try_pdf_ocr_fallback(text):
    score, metrics = score_pdf_text_quality(text)
    char_count = metrics['char_count']
    token_count = metrics['token_count']
    avg_token_len = metrics['avg_token_len']
    long_token_ratio = metrics['long_token_ratio']
    cjk_ratio = metrics['cjk_ratio']

    if char_count < 320 or token_count < 60:
        return True, score, metrics
    if cjk_ratio < 0.2:
        if avg_token_len >= 10.5:
            return True, score, metrics
        if long_token_ratio >= 0.12:
            return True, score, metrics
    if score < 16.0:
        return True, score, metrics
    return False, score, metrics


def extract_text_from_pdf_bytes_pymupdf(file_bytes):
    if fitz is None:
        return ''

    try:
        doc = fitz.open(stream=file_bytes, filetype='pdf')
    except Exception as e:
        print(f"PyMuPDF open failed: {e}")
        return ''

    page_outputs = []
    try:
        for page in doc:
            page_lines = []
            words = page.get_text('words') or []
            if words:
                grouped = {}
                for item in words:
                    if not isinstance(item, (list, tuple)) or len(item) < 8:
                        continue
                    token = str(item[4] or '').strip()
                    if not token:
                        continue
                    block_no = int(item[5])
                    line_no = int(item[6])
                    word_no = int(item[7])
                    key = (block_no, line_no)
                    grouped.setdefault(key, []).append((word_no, token))

                for key in sorted(grouped.keys()):
                    tokens = [token for _, token in sorted(grouped[key], key=lambda pair: pair[0])]
                    line = ' '.join(tokens).strip()
                    if line:
                        page_lines.append(line)

            if not page_lines:
                raw_text = page.get_text('text') or ''
                raw_text = normalize_newlines(raw_text).strip()
                if raw_text:
                    page_lines.append(raw_text)

            page_outputs.append('\n'.join(page_lines).strip())
    except Exception as e:
        print(f"PyMuPDF extraction failed: {e}")
        return ''
    finally:
        doc.close()

    result = '\n\n'.join([chunk for chunk in page_outputs if chunk]).strip()
    return normalize_pdf_text(result)


def extract_text_from_pdf_bytes_pypdf2(file_bytes):
    text = ''
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text() or ''
            text += page_text + '\n'
    except Exception as e:
        print(f"PyPDF2 extraction failed: {e}")
        return ''
    return normalize_pdf_text(text)


def run_ocrmypdf_on_pdf_bytes(file_bytes):
    if not ENABLE_PDF_OCR_FALLBACK:
        return b'', 'ocrmypdf fallback disabled by ENABLE_PDF_OCR_FALLBACK'
    if not file_bytes:
        return b'', 'Empty PDF bytes'

    ocrmypdf_path = shutil.which(OCRMYPDF_BINARY)
    if not ocrmypdf_path:
        return b'', f'ocrmypdf binary not found: {OCRMYPDF_BINARY}'

    try:
        with tempfile.TemporaryDirectory(prefix='studyhub-pdfocr-') as tmp_dir:
            input_path = os.path.join(tmp_dir, 'input.pdf')
            output_path = os.path.join(tmp_dir, 'output.pdf')
            with open(input_path, 'wb') as f:
                f.write(file_bytes)

            cmd = [
                ocrmypdf_path,
                '--force-ocr',
                '--output-type', 'pdf',
                '--optimize', '0',
                '--quiet',
                '-l', OCRMYPDF_LANGUAGE,
                input_path,
                output_path,
            ]
            proc = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=OCRMYPDF_TIMEOUT_SECONDS,
                check=False,
            )
            if proc.returncode != 0:
                stderr = (proc.stderr or b'').decode('utf-8', errors='ignore').strip()
                stdout = (proc.stdout or b'').decode('utf-8', errors='ignore').strip()
                details = stderr or stdout or f'return code {proc.returncode}'
                return b'', f'ocrmypdf failed: {details[:260]}'

            if not os.path.exists(output_path):
                return b'', 'ocrmypdf did not produce output file'
            with open(output_path, 'rb') as f:
                return f.read(), ''
    except subprocess.TimeoutExpired:
        return b'', f'ocrmypdf timeout after {OCRMYPDF_TIMEOUT_SECONDS}s'
    except Exception as e:
        return b'', f'ocrmypdf runtime error: {e}'


def extract_text_from_pdf_bytes_with_meta(file_bytes):
    meta = {
        'extractor': 'none',
        'ocr_attempted': False,
        'ocr_used': False,
        'quality_score_before': 0.0,
        'quality_score_after': 0.0,
        'quality_metrics_before': {},
        'quality_metrics_after': {},
        'note': '',
    }
    if not file_bytes:
        meta['note'] = 'Empty PDF bytes'
        return '', meta

    primary_text = extract_text_from_pdf_bytes_pymupdf(file_bytes)
    extractor = 'pymupdf'
    if not primary_text:
        primary_text = extract_text_from_pdf_bytes_pypdf2(file_bytes)
        extractor = 'pypdf2'

    if not primary_text:
        meta['note'] = 'Primary PDF text extractors returned empty text'
        return 'Text extraction failed.', meta

    before_score, before_metrics = score_pdf_text_quality(primary_text)
    meta['extractor'] = extractor
    meta['quality_score_before'] = before_score
    meta['quality_metrics_before'] = before_metrics

    should_ocr, _, _ = should_try_pdf_ocr_fallback(primary_text)
    if not should_ocr:
        meta['quality_score_after'] = before_score
        meta['quality_metrics_after'] = before_metrics
        return primary_text, meta

    ocr_pdf_bytes, ocr_error = run_ocrmypdf_on_pdf_bytes(file_bytes)
    meta['ocr_attempted'] = True
    if not ocr_pdf_bytes:
        meta['quality_score_after'] = before_score
        meta['quality_metrics_after'] = before_metrics
        meta['note'] = ocr_error
        return primary_text, meta

    ocr_text = extract_text_from_pdf_bytes_pymupdf(ocr_pdf_bytes) or extract_text_from_pdf_bytes_pypdf2(ocr_pdf_bytes)
    if not ocr_text:
        meta['quality_score_after'] = before_score
        meta['quality_metrics_after'] = before_metrics
        meta['note'] = 'ocrmypdf produced file but text extraction stayed empty'
        return primary_text, meta

    after_score, after_metrics = score_pdf_text_quality(ocr_text)
    meta['quality_score_after'] = after_score
    meta['quality_metrics_after'] = after_metrics
    if after_score >= before_score + 1.0:
        meta['ocr_used'] = True
        if ocr_error:
            meta['note'] = ocr_error
        return ocr_text, meta

    meta['note'] = 'OCR output not better than primary extraction'
    return primary_text, meta


def extract_text_from_pdf_bytes(file_bytes):
    text, _ = extract_text_from_pdf_bytes_with_meta(file_bytes)
    return text


def create_pdf_bytes_from_text(content):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas
    except Exception as e:
        raise RuntimeError(
            f'PDF editing requires reportlab in current interpreter: {sys.executable}. '
            f'Install it with "{sys.executable} -m pip install reportlab".'
        ) from e

    text = content if isinstance(content, str) else str(content or '')
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    stream = io.BytesIO()
    pdf_canvas = canvas.Canvas(stream, pagesize=A4)
    page_width, page_height = A4

    font_name = 'Helvetica'
    font_size = 11
    try:
        pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
        font_name = 'STSong-Light'
    except Exception:
        font_name = 'Helvetica'

    pdf_canvas.setFont(font_name, font_size)
    left_margin = 48
    top_margin = 56
    bottom_margin = 56
    line_height = 16
    max_line_width = page_width - (left_margin * 2)

    def wrap_paragraph(paragraph):
        if paragraph == '':
            return ['']

        wrapped = []
        current = ''
        for char in paragraph:
            candidate = current + char
            try:
                text_width = pdfmetrics.stringWidth(candidate, font_name, font_size)
            except Exception:
                text_width = len(candidate) * font_size * 0.6

            if text_width <= max_line_width or not current:
                current = candidate
            else:
                wrapped.append(current)
                current = char

        if current or not wrapped:
            wrapped.append(current)
        return wrapped

    y = page_height - top_margin
    paragraphs = text.split('\n')
    if not paragraphs:
        paragraphs = ['']

    for paragraph in paragraphs:
        for line in wrap_paragraph(paragraph):
            if y < bottom_margin:
                pdf_canvas.showPage()
                pdf_canvas.setFont(font_name, font_size)
                y = page_height - top_margin

            draw_text = line
            if font_name == 'Helvetica':
                draw_text = line.encode('latin-1', 'replace').decode('latin-1')

            pdf_canvas.drawString(left_margin, y, draw_text)
            y -= line_height

    pdf_canvas.save()
    stream.seek(0)
    return stream.read()


def build_editable_file_bytes(file_ext, content, content_html=''):
    ext = (file_ext or '').lower().strip('.')
    text = normalize_newlines(content if isinstance(content, str) else str(content or ''))

    if ext == 'txt':
        return text.encode('utf-8'), 'text/plain'

    if ext == 'docx':
        return create_docx_bytes_from_html(content_html, text), MIME_BY_EXT['docx']

    if ext == 'pdf':
        return create_pdf_bytes_from_text(text), MIME_BY_EXT['pdf']

    raise ValueError('Only txt, docx and pdf support direct source-file update right now.')


__all__ = [
    'build_editable_file_bytes',
    'extract_document_content',
    'extract_text_from_pdf_bytes',
    'extract_text_from_pdf_bytes_with_meta',
    'hard_delete_document_record',
    'html_to_plaintext',
    'infer_document_category',
    'normalize_newlines',
    'plaintext_to_html',
    'purge_expired_trashed_documents',
    'sanitize_editor_html',
    'user_can_edit_document',
]
